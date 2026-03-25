# -*- coding: utf-8 -*-
import io
import numpy as np
import pandas as pd
import streamlit as st
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from reports.docx_helpers import (
    _estilo_celda_report, _agregar_titulo_report,
    _tabla_info_report, _tabla_datos_report,
)
from utils.formatters import (
    _fmt_precio, _fmt_entero, _fmt_monto, _fmt_iv, _fmt_lado,
    determinar_sentimiento,
)
from utils.helpers import _enriquecer_datos_opcion

def _generar_reporte_live_scanning():
    """Genera reporte DOCX con todos los datos del Live Scanning."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — LIVE SCANNING", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    ticker_name = st.session_state.get("ticker_anterior", "N/A")
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run(f"Ticker: {ticker_name}")
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run_sub.font.name = "Calibri"
    run_sub.bold = True

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_fecha.font.name = "Calibri"

    doc.add_paragraph("")

    # Resumen Ejecutivo
    n_opciones = len(st.session_state.datos_completos)
    n_alertas = len(st.session_state.alertas_actuales)
    n_clusters = len(st.session_state.clusters_detectados)
    precio_subyacente = st.session_state.get('precio_subyacente', 0)

    datos_calls = [d for d in st.session_state.datos_completos if d.get('Tipo', '') == 'CALL']
    datos_puts = [d for d in st.session_state.datos_completos if d.get('Tipo', '') == 'PUT']

    vol_promedio = np.mean([d.get('Volumen', 0) for d in st.session_state.datos_completos]) if st.session_state.datos_completos else 0
    oi_promedio = np.mean([d.get('OI', 0) for d in st.session_state.datos_completos]) if st.session_state.datos_completos else 0
    iv_promedio = np.mean([d.get('IV', 0) for d in st.session_state.datos_completos if d.get('IV', 0) > 0]) if st.session_state.datos_completos else 0

    principales = [a for a in st.session_state.alertas_actuales if a.get("Tipo_Alerta") == "PRINCIPAL"]
    prima_alta = [a for a in st.session_state.alertas_actuales if a.get("Tipo_Alerta") == "PRIMA_ALTA"]

    _agregar_titulo_report(doc, "RESUMEN EJECUTIVO", level=1)
    _tabla_info_report(doc, {
        "Ticker Analizado": ticker_name,
        "Precio Subyacente": f"${precio_subyacente:,.2f}" if precio_subyacente > 0 else "N/D",
        "Fecha del Reporte": fecha_legible,
        "Total Opciones Escaneadas": f"{n_opciones:,}",
        "Calls vs Puts": f"{len(datos_calls):,} calls / {len(datos_puts):,} puts",
        "Alertas Detectadas": f"{n_alertas} ({len(principales)} institucionales, {len(prima_alta)} prima alta)",
        "Clusters de Compra": f"{n_clusters}",
        "Volumen Promedio": f"{vol_promedio:,.0f}",
        "OI Promedio": f"{oi_promedio:,.0f}",
        "IV Promedio": f"{iv_promedio:.1f}%" if iv_promedio > 0 else "N/D",
    })

    # Alertas Institucionales
    if principales:
        _agregar_titulo_report(doc, f"ALERTAS INSTITUCIONALES ({len(principales)})", level=1)
        p_desc = doc.add_paragraph()
        run_d = p_desc.add_run(
            "Operaciones con prima significativa que sugieren actividad institucional."
        )
        run_d.font.size = Pt(10)
        run_d.font.italic = True
        run_d.font.name = "Calibri"

        headers = ["#", "Tipo", "Strike", "Vencimiento", "Volumen", "OI",
                   "Ask", "Bid", "Último", "IV", "Sentimiento", "Lado", "Prima Total", "Contrato"]
        rows = []
        principales_enriq = _enriquecer_datos_opcion(principales, precio_subyacente)

        for i, a in enumerate(principales_enriq, 1):
            sent_txt, sent_emoji, _ = determinar_sentimiento(a["Tipo_Opcion"], a.get("Lado", "N/A"))
            rows.append([
                i, a["Tipo_Opcion"], _fmt_precio(a['Strike']), a["Vencimiento"],
                _fmt_entero(a['Volumen']), _fmt_entero(a['OI']),
                _fmt_precio(a['Ask']), _fmt_precio(a['Bid']), _fmt_precio(a['Ultimo']),
                _fmt_iv(a.get('IV', 0)),
                f"{sent_emoji} {sent_txt}", _fmt_lado(a.get('Lado', 'N/A')),
                _fmt_monto(a['Prima_Volumen']),
                a.get("Contrato", "N/A"),
            ])
        _tabla_datos_report(doc, headers, rows)

    # Alertas Prima Alta
    if prima_alta:
        _agregar_titulo_report(doc, f"ALERTAS PRIMA ALTA ({len(prima_alta)})", level=1)
        p_desc = doc.add_paragraph()
        run_d = p_desc.add_run(
            "Opciones con volumen y open interest por encima de los umbrales configurados."
        )
        run_d.font.size = Pt(10)
        run_d.font.italic = True
        run_d.font.name = "Calibri"

        headers = ["#", "Tipo", "Strike", "Vencimiento", "Volumen", "OI",
                   "Ask", "Bid", "Último", "IV", "Sentimiento", "Lado", "Prima Total"]
        rows = []
        prima_alta_enriq = _enriquecer_datos_opcion(prima_alta, precio_subyacente)

        for i, a in enumerate(prima_alta_enriq, 1):
            sent_txt, sent_emoji, _ = determinar_sentimiento(a["Tipo_Opcion"], a.get("Lado", "N/A"))
            rows.append([
                i, a["Tipo_Opcion"], _fmt_precio(a['Strike']), a["Vencimiento"],
                _fmt_entero(a['Volumen']), _fmt_entero(a['OI']),
                _fmt_precio(a['Ask']), _fmt_precio(a['Bid']), _fmt_precio(a['Ultimo']),
                _fmt_iv(a.get('IV', 0)),
                f"{sent_emoji} {sent_txt}", _fmt_lado(a.get('Lado', 'N/A')),
                _fmt_monto(a['Prima_Volumen']),
            ])
        _tabla_datos_report(doc, headers, rows)

    # Clusters
    if st.session_state.clusters_detectados:
        _agregar_titulo_report(doc, f"CLUSTERS DE COMPRA CONTINUA ({n_clusters})", level=1)
        p_desc = doc.add_paragraph()
        run_d = p_desc.add_run(
            "Grupos de contratos con strikes cercanos y primas similares en la misma expiración."
        )
        run_d.font.size = Pt(10)
        run_d.font.italic = True
        run_d.font.name = "Calibri"

        headers_cl = ["#", "Tipo", "Vencimiento", "Contratos", "Rango Strikes",
                      "Prima Total", "Prima Prom.", "Vol Total", "OI Total"]
        rows_cl = []
        for i, c in enumerate(st.session_state.clusters_detectados, 1):
            rows_cl.append([
                i, c["Tipo_Opcion"], c["Vencimiento"], c["Contratos"],
                f"${c['Strike_Min']} — ${c['Strike_Max']}",
                _fmt_monto(c['Prima_Total']), _fmt_monto(c['Prima_Promedio']),
                _fmt_entero(c['Vol_Total']), _fmt_entero(c['OI_Total']),
            ])
        _tabla_datos_report(doc, headers_cl, rows_cl)

        # Detalle de cada cluster
        for i, c in enumerate(st.session_state.clusters_detectados, 1):
            if c.get("Detalle"):
                p_cl = doc.add_paragraph()
                run_cl = p_cl.add_run(f"Detalle Cluster #{i} — {c['Tipo_Opcion']} Venc. {c['Vencimiento']}")
                run_cl.bold = True
                run_cl.font.size = Pt(10)
                run_cl.font.name = "Calibri"

                headers_det = ["#", "Strike", "Volumen", "OI", "Prima Total"]
                rows_det = []
                for j, d in enumerate(c["Detalle"], 1):
                    rows_det.append([
                        j, _fmt_precio(d['Strike']),
                        _fmt_entero(d['Volumen']), _fmt_entero(d['OI']),
                        _fmt_monto(d['Prima_Volumen']),
                    ])
                _tabla_datos_report(doc, headers_det, rows_det)

    # Todas las opciones escaneadas
    if st.session_state.datos_completos:
        _agregar_titulo_report(doc, "TODAS LAS OPCIONES ESCANEADAS", level=1)

        datos_sorted = sorted(
            st.session_state.datos_completos,
            key=lambda x: x.get("Prima_Volumen", 0), reverse=True,
        )

        p_info = doc.add_paragraph()
        run_info = p_info.add_run(f"Total de opciones: {len(datos_sorted):,}")
        run_info.font.size = Pt(10)
        run_info.font.italic = True
        run_info.font.name = "Calibri"

        headers_opt = ["Tipo", "Vencimiento", "Strike", "Volumen", "OI",
                       "Ask", "Bid", "Último", "Lado", "IV", "Prima Total"]
        rows_opt = []
        datos_enriquecidos = _enriquecer_datos_opcion(datos_sorted, precio_subyacente)

        for d in datos_enriquecidos:
            rows_opt.append([
                d["Tipo"], d["Vencimiento"], _fmt_precio(d['Strike']),
                _fmt_entero(d['Volumen']), _fmt_entero(d['OI']),
                _fmt_precio(d['Ask']), _fmt_precio(d['Bid']), _fmt_precio(d['Ultimo']),
                _fmt_lado(d.get('Lado', 'N/A')), _fmt_iv(d.get('IV', 0)),
                _fmt_monto(d.get('Prima_Volumen', 0)),
            ])
        _tabla_datos_report(doc, headers_opt, rows_opt)

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Live Scanning — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


