# -*- coding: utf-8 -*-
import io
import numpy as np
import pandas as pd
import streamlit as st
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from reports.docx_helpers import (
    _estilo_celda_report, _agregar_titulo_report,
    _tabla_info_report, _tabla_datos_report,
)

def _generar_reporte_open_interest():
    """Genera reporte DOCX con análisis de Open Interest."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — OPEN INTEREST", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    ticker_name = st.session_state.get("ticker_anterior", "N/A")
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run(f"Ticker: {ticker_name}")
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run_sub.font.name = "Calibri"
    run_sub.bold = True

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_fecha.font.name = "Calibri"

    doc.add_paragraph("")

    # Datos de Barchart
    if st.session_state.barchart_data is not None and not st.session_state.barchart_data.empty:
        df_bc = st.session_state.barchart_data.copy()

        df_positivos = df_bc[df_bc["OI_Chg"] > 0].sort_values("OI_Chg", ascending=False)
        df_negativos = df_bc[df_bc["OI_Chg"] < 0].sort_values("OI_Chg", ascending=True)

        n_total = len(df_bc)
        n_pos = len(df_positivos)
        n_neg = len(df_negativos)
        n_calls = len(df_bc[df_bc["Tipo"] == "CALL"]) if "Tipo" in df_bc.columns else 0
        n_puts = len(df_bc[df_bc["Tipo"] == "PUT"]) if "Tipo" in df_bc.columns else 0

        contratos_abiertos = int(df_positivos["OI_Chg"].sum()) if n_pos > 0 else 0
        contratos_cerrados = int(df_negativos["OI_Chg"].sum()) if n_neg > 0 else 0

        # Resumen
        _agregar_titulo_report(doc, "RESUMEN DE CAMBIOS EN OPEN INTEREST", level=1)
        _tabla_info_report(doc, {
            "Ticker": ticker_name,
            "Fecha": fecha_legible,
            "Total Contratos Analizados": f"{n_total:,}",
            "Calls": f"{n_calls:,}",
            "Puts": f"{n_puts:,}",
            "Contratos Abiertos (OI Positivo)": f"{contratos_abiertos:,}",
            "Contratos Cerrados (OI Negativo)": f"{contratos_cerrados:,}",
            "Señales Positivas": f"{n_pos:,}",
            "Señales Negativas": f"{n_neg:,}",
        })

        # Tabla OI Positivo
        if n_pos > 0:
            _agregar_titulo_report(doc, f"OI POSITIVO — ABRIENDO POSICIONES ({n_pos})", level=1)
            p_desc = doc.add_paragraph()
            run_d = p_desc.add_run(
                "Contratos donde el Open Interest aumentó, indicando nuevas posiciones abiertas."
            )
            run_d.font.size = Pt(10)
            run_d.font.italic = True
            run_d.font.name = "Calibri"

            headers_pos = ["#", "Tipo", "Strike", "Vencimiento", "DTE", "Volumen", "OI", "OI Chg", "IV", "Delta", "Último"]
            rows_pos = []
            for i, row in enumerate(df_positivos.head(100).itertuples(), 1):
                rows_pos.append([
                    i,
                    row.Tipo if hasattr(row, 'Tipo') else "N/A",
                    f"${row.Strike:,.1f}" if hasattr(row, 'Strike') else "N/A",
                    row.Vencimiento if hasattr(row, 'Vencimiento') else "N/A",
                    f"{row.DTE}d" if hasattr(row, 'DTE') else "N/A",
                    f"{int(row.Volumen):,}" if hasattr(row, 'Volumen') else "N/A",
                    f"{int(row.OI):,}" if hasattr(row, 'OI') else "N/A",
                    f"+{int(row.OI_Chg):,}" if hasattr(row, 'OI_Chg') else "N/A",
                    f"{row.IV:.1f}%" if hasattr(row, 'IV') and row.IV > 0 else "N/A",
                    f"{row.Delta:.3f}" if hasattr(row, 'Delta') and row.Delta != 0 else "N/A",
                    f"${row.Último:.2f}" if hasattr(row, 'Último') and row.Último > 0 else "N/A",
                ])
            _tabla_datos_report(doc, headers_pos, rows_pos)

        # Tabla OI Negativo
        if n_neg > 0:
            _agregar_titulo_report(doc, f"OI NEGATIVO — CERRANDO POSICIONES ({n_neg})", level=1)
            p_desc = doc.add_paragraph()
            run_d = p_desc.add_run(
                "Contratos donde el Open Interest disminuyó, indicando posiciones cerradas o ejercidas."
            )
            run_d.font.size = Pt(10)
            run_d.font.italic = True
            run_d.font.name = "Calibri"

            headers_neg = ["#", "Tipo", "Strike", "Vencimiento", "DTE", "Volumen", "OI", "OI Chg", "IV", "Delta", "Último"]
            rows_neg = []
            for i, row in enumerate(df_negativos.head(100).itertuples(), 1):
                rows_neg.append([
                    i,
                    row.Tipo if hasattr(row, 'Tipo') else "N/A",
                    f"${row.Strike:,.1f}" if hasattr(row, 'Strike') else "N/A",
                    row.Vencimiento if hasattr(row, 'Vencimiento') else "N/A",
                    f"{row.DTE}d" if hasattr(row, 'DTE') else "N/A",
                    f"{int(row.Volumen):,}" if hasattr(row, 'Volumen') else "N/A",
                    f"{int(row.OI):,}" if hasattr(row, 'OI') else "N/A",
                    f"{int(row.OI_Chg):,}" if hasattr(row, 'OI_Chg') else "N/A",
                    f"{row.IV:.1f}%" if hasattr(row, 'IV') and row.IV > 0 else "N/A",
                    f"{row.Delta:.3f}" if hasattr(row, 'Delta') and row.Delta != 0 else "N/A",
                    f"${row.Último:.2f}" if hasattr(row, 'Último') and row.Último > 0 else "N/A",
                ])
            _tabla_datos_report(doc, headers_neg, rows_neg)

    else:
        # Sin datos
        p_sin = doc.add_paragraph()
        run_sin = p_sin.add_run("No hay datos de Open Interest disponibles. Ejecuta un escaneo primero.")
        run_sin.font.size = Pt(11)
        run_sin.font.italic = True
        run_sin.font.name = "Calibri"

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Open Interest — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
