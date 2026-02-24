# -*- coding: utf-8 -*-
"""
Generadores de reportes DOCX profesionales.
Extraídos de app_web.py — cero cambios de lógica.
"""
import io
import numpy as np
import pandas as pd
import streamlit as st
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from utils.formatters import (
    _fmt_precio, _fmt_entero, _fmt_monto, _fmt_iv, _fmt_lado,
    determinar_sentimiento,
)
from utils.helpers import _enriquecer_datos_opcion
from ui.components import format_market_cap


# ============================================================================
#                    HELPERS PARA GENERAR REPORTES DOCX
# ============================================================================
def _estilo_celda_report(cell, texto, negrita=False, color_fondo=None, color_texto=None, size=9, align="left"):
    """Aplica formato a una celda de tabla Word."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(str(texto))
    run.bold = negrita
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color_texto:
        run.font.color.rgb = color_texto
    if color_fondo:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(qn("w:shd"), {
            qn("w:fill"): color_fondo,
            qn("w:val"): "clear",
        })
        tcPr.append(shading)


def _agregar_titulo_report(doc, texto, level=2):
    """Agrega un título de sección con formato."""
    heading = doc.add_heading(texto, level=level)
    for run in heading.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)


def _tabla_info_report(doc, datos_dict, titulo=None):
    """Tabla de 2 columnas Campo/Valor para info resumida."""
    if titulo:
        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for campo, valor in datos_dict.items():
        row = table.add_row()
        _estilo_celda_report(row.cells[0], campo, negrita=True, size=10)
        _estilo_celda_report(row.cells[1], str(valor), size=10)
    doc.add_paragraph("")


def _tabla_datos_report(doc, headers, rows_data):
    """Tabla con encabezados y filas de datos."""
    if not rows_data:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Encabezados
    for i, h in enumerate(headers):
        _estilo_celda_report(
            table.rows[0].cells[i], h,
            negrita=True, size=9,
            color_fondo="1E3A5F",
            color_texto=RGBColor(0xFF, 0xFF, 0xFF),
            align="center",
        )
    # Datos
    for row_idx, row_data in enumerate(rows_data):
        row = table.add_row()
        bg = "F0F4F8" if row_idx % 2 == 0 else None
        for i, val in enumerate(row_data):
            _estilo_celda_report(row.cells[i], str(val), size=9, color_fondo=bg)
    doc.add_paragraph("")


# ============================================================================
#   FUNCIÓN 1: REPORTE LIVE SCANNING
# ============================================================================
def _generar_reporte_live_scanning():
    """Genera reporte DOCX con todos los datos del Live Scanning."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — LIVE SCANNING", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    ticker_name = st.session_state.get("ticker_anterior", "N/A")
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run(f"Ticker: {ticker_name}")
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run_sub.font.name = "Calibri"
    run_sub.bold = True

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_fecha.font.name = "Calibri"

    doc.add_paragraph("")

    # Resumen Ejecutivo
    n_opciones = len(st.session_state.datos_completos)
    n_alertas = len(st.session_state.alertas_actuales)
    n_clusters = len(st.session_state.clusters_detectados)
    precio_subyacente = st.session_state.get('precio_subyacente', 0)

    datos_calls = [d for d in st.session_state.datos_completos if d.get('Tipo', '') == 'CALL']
    datos_puts = [d for d in st.session_state.datos_completos if d.get('Tipo', '') == 'PUT']

    vol_promedio = np.mean([d.get('Volumen', 0) for d in st.session_state.datos_completos]) if st.session_state.datos_completos else 0
    oi_promedio = np.mean([d.get('OI', 0) for d in st.session_state.datos_completos]) if st.session_state.datos_completos else 0
    iv_promedio = np.mean([d.get('IV', 0) for d in st.session_state.datos_completos if d.get('IV', 0) > 0]) if st.session_state.datos_completos else 0

    principales = [a for a in st.session_state.alertas_actuales if a.get("Tipo_Alerta") == "PRINCIPAL"]
    prima_alta = [a for a in st.session_state.alertas_actuales if a.get("Tipo_Alerta") == "PRIMA_ALTA"]

    _agregar_titulo_report(doc, "RESUMEN EJECUTIVO", level=1)
    _tabla_info_report(doc, {
        "Ticker Analizado": ticker_name,
        "Precio Subyacente": f"${precio_subyacente:,.2f}" if precio_subyacente > 0 else "N/D",
        "Fecha del Reporte": fecha_legible,
        "Total Opciones Escaneadas": f"{n_opciones:,}",
        "Calls vs Puts": f"{len(datos_calls):,} calls / {len(datos_puts):,} puts",
        "Alertas Detectadas": f"{n_alertas} ({len(principales)} institucionales, {len(prima_alta)} prima alta)",
        "Clusters de Compra": f"{n_clusters}",
        "Volumen Promedio": f"{vol_promedio:,.0f}",
        "OI Promedio": f"{oi_promedio:,.0f}",
        "IV Promedio": f"{iv_promedio:.1f}%" if iv_promedio > 0 else "N/D",
    })

    # Alertas Institucionales
    if principales:
        _agregar_titulo_report(doc, f"ALERTAS INSTITUCIONALES ({len(principales)})", level=1)
        p_desc = doc.add_paragraph()
        run_d = p_desc.add_run(
            "Operaciones con prima significativa que sugieren actividad institucional."
        )
        run_d.font.size = Pt(10)
        run_d.font.italic = True
        run_d.font.name = "Calibri"

        headers = ["#", "Tipo", "Strike", "Vencimiento", "Volumen", "OI",
                   "Ask", "Bid", "Último", "IV", "Sentimiento", "Lado", "Prima Total", "Contrato"]
        rows = []
        principales_enriq = _enriquecer_datos_opcion(principales, precio_subyacente)

        for i, a in enumerate(principales_enriq, 1):
            sent_txt, sent_emoji, _ = determinar_sentimiento(a["Tipo_Opcion"], a.get("Lado", "N/A"))
            rows.append([
                i, a["Tipo_Opcion"], _fmt_precio(a['Strike']), a["Vencimiento"],
                _fmt_entero(a['Volumen']), _fmt_entero(a['OI']),
                _fmt_precio(a['Ask']), _fmt_precio(a['Bid']), _fmt_precio(a['Ultimo']),
                _fmt_iv(a.get('IV', 0)),
                f"{sent_emoji} {sent_txt}", _fmt_lado(a.get('Lado', 'N/A')),
                _fmt_monto(a['Prima_Volumen']),
                a.get("Contrato", "N/A"),
            ])
        _tabla_datos_report(doc, headers, rows)

    # Alertas Prima Alta
    if prima_alta:
        _agregar_titulo_report(doc, f"ALERTAS PRIMA ALTA ({len(prima_alta)})", level=1)
        p_desc = doc.add_paragraph()
        run_d = p_desc.add_run(
            "Opciones con volumen y open interest por encima de los umbrales configurados."
        )
        run_d.font.size = Pt(10)
        run_d.font.italic = True
        run_d.font.name = "Calibri"

        headers = ["#", "Tipo", "Strike", "Vencimiento", "Volumen", "OI",
                   "Ask", "Bid", "Último", "IV", "Sentimiento", "Lado", "Prima Total"]
        rows = []
        prima_alta_enriq = _enriquecer_datos_opcion(prima_alta, precio_subyacente)

        for i, a in enumerate(prima_alta_enriq, 1):
            sent_txt, sent_emoji, _ = determinar_sentimiento(a["Tipo_Opcion"], a.get("Lado", "N/A"))
            rows.append([
                i, a["Tipo_Opcion"], _fmt_precio(a['Strike']), a["Vencimiento"],
                _fmt_entero(a['Volumen']), _fmt_entero(a['OI']),
                _fmt_precio(a['Ask']), _fmt_precio(a['Bid']), _fmt_precio(a['Ultimo']),
                _fmt_iv(a.get('IV', 0)),
                f"{sent_emoji} {sent_txt}", _fmt_lado(a.get('Lado', 'N/A')),
                _fmt_monto(a['Prima_Volumen']),
            ])
        _tabla_datos_report(doc, headers, rows)

    # Clusters
    if st.session_state.clusters_detectados:
        _agregar_titulo_report(doc, f"CLUSTERS DE COMPRA CONTINUA ({n_clusters})", level=1)
        p_desc = doc.add_paragraph()
        run_d = p_desc.add_run(
            "Grupos de contratos con strikes cercanos y primas similares en la misma expiración."
        )
        run_d.font.size = Pt(10)
        run_d.font.italic = True
        run_d.font.name = "Calibri"

        headers_cl = ["#", "Tipo", "Vencimiento", "Contratos", "Rango Strikes",
                      "Prima Total", "Prima Prom.", "Vol Total", "OI Total"]
        rows_cl = []
        for i, c in enumerate(st.session_state.clusters_detectados, 1):
            rows_cl.append([
                i, c["Tipo_Opcion"], c["Vencimiento"], c["Contratos"],
                f"${c['Strike_Min']} — ${c['Strike_Max']}",
                _fmt_monto(c['Prima_Total']), _fmt_monto(c['Prima_Promedio']),
                _fmt_entero(c['Vol_Total']), _fmt_entero(c['OI_Total']),
            ])
        _tabla_datos_report(doc, headers_cl, rows_cl)

        # Detalle de cada cluster
        for i, c in enumerate(st.session_state.clusters_detectados, 1):
            if c.get("Detalle"):
                p_cl = doc.add_paragraph()
                run_cl = p_cl.add_run(f"Detalle Cluster #{i} — {c['Tipo_Opcion']} Venc. {c['Vencimiento']}")
                run_cl.bold = True
                run_cl.font.size = Pt(10)
                run_cl.font.name = "Calibri"

                headers_det = ["#", "Strike", "Volumen", "OI", "Prima Total"]
                rows_det = []
                for j, d in enumerate(c["Detalle"], 1):
                    rows_det.append([
                        j, _fmt_precio(d['Strike']),
                        _fmt_entero(d['Volumen']), _fmt_entero(d['OI']),
                        _fmt_monto(d['Prima_Volumen']),
                    ])
                _tabla_datos_report(doc, headers_det, rows_det)

    # Todas las opciones escaneadas
    if st.session_state.datos_completos:
        _agregar_titulo_report(doc, "TODAS LAS OPCIONES ESCANEADAS", level=1)

        datos_sorted = sorted(
            st.session_state.datos_completos,
            key=lambda x: x.get("Prima_Volumen", 0), reverse=True,
        )

        p_info = doc.add_paragraph()
        run_info = p_info.add_run(f"Total de opciones: {len(datos_sorted):,}")
        run_info.font.size = Pt(10)
        run_info.font.italic = True
        run_info.font.name = "Calibri"

        headers_opt = ["Tipo", "Vencimiento", "Strike", "Volumen", "OI",
                       "Ask", "Bid", "Último", "Lado", "IV", "Prima Total"]
        rows_opt = []
        datos_enriquecidos = _enriquecer_datos_opcion(datos_sorted, precio_subyacente)

        for d in datos_enriquecidos:
            rows_opt.append([
                d["Tipo"], d["Vencimiento"], _fmt_precio(d['Strike']),
                _fmt_entero(d['Volumen']), _fmt_entero(d['OI']),
                _fmt_precio(d['Ask']), _fmt_precio(d['Bid']), _fmt_precio(d['Ultimo']),
                _fmt_lado(d.get('Lado', 'N/A')), _fmt_iv(d.get('IV', 0)),
                _fmt_monto(d.get('Prima_Volumen', 0)),
            ])
        _tabla_datos_report(doc, headers_opt, rows_opt)

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Live Scanning — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
#   FUNCIÓN 2: REPORTE OPEN INTEREST
# ============================================================================
def _generar_reporte_open_interest():
    """Genera reporte DOCX con análisis de Open Interest."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — OPEN INTEREST", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    ticker_name = st.session_state.get("ticker_anterior", "N/A")
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run(f"Ticker: {ticker_name}")
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run_sub.font.name = "Calibri"
    run_sub.bold = True

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_fecha.font.name = "Calibri"

    doc.add_paragraph("")

    # Datos de Barchart
    if st.session_state.barchart_data is not None and not st.session_state.barchart_data.empty:
        df_bc = st.session_state.barchart_data.copy()

        df_positivos = df_bc[df_bc["OI_Chg"] > 0].sort_values("OI_Chg", ascending=False)
        df_negativos = df_bc[df_bc["OI_Chg"] < 0].sort_values("OI_Chg", ascending=True)

        n_total = len(df_bc)
        n_pos = len(df_positivos)
        n_neg = len(df_negativos)
        n_calls = len(df_bc[df_bc["Tipo"] == "CALL"]) if "Tipo" in df_bc.columns else 0
        n_puts = len(df_bc[df_bc["Tipo"] == "PUT"]) if "Tipo" in df_bc.columns else 0

        contratos_abiertos = int(df_positivos["OI_Chg"].sum()) if n_pos > 0 else 0
        contratos_cerrados = int(df_negativos["OI_Chg"].sum()) if n_neg > 0 else 0

        # Resumen
        _agregar_titulo_report(doc, "RESUMEN DE CAMBIOS EN OPEN INTEREST", level=1)
        _tabla_info_report(doc, {
            "Ticker": ticker_name,
            "Fecha": fecha_legible,
            "Total Contratos Analizados": f"{n_total:,}",
            "Calls": f"{n_calls:,}",
            "Puts": f"{n_puts:,}",
            "Contratos Abiertos (OI Positivo)": f"{contratos_abiertos:,}",
            "Contratos Cerrados (OI Negativo)": f"{contratos_cerrados:,}",
            "Señales Positivas": f"{n_pos:,}",
            "Señales Negativas": f"{n_neg:,}",
        })

        # Tabla OI Positivo
        if n_pos > 0:
            _agregar_titulo_report(doc, f"OI POSITIVO — ABRIENDO POSICIONES ({n_pos})", level=1)
            p_desc = doc.add_paragraph()
            run_d = p_desc.add_run(
                "Contratos donde el Open Interest aumentó, indicando nuevas posiciones abiertas."
            )
            run_d.font.size = Pt(10)
            run_d.font.italic = True
            run_d.font.name = "Calibri"

            headers_pos = ["#", "Tipo", "Strike", "Vencimiento", "DTE", "Volumen", "OI", "OI Chg", "IV", "Delta", "Último"]
            rows_pos = []
            for i, row in enumerate(df_positivos.head(100).itertuples(), 1):
                rows_pos.append([
                    i,
                    row.Tipo if hasattr(row, 'Tipo') else "N/A",
                    f"${row.Strike:,.1f}" if hasattr(row, 'Strike') else "N/A",
                    row.Vencimiento if hasattr(row, 'Vencimiento') else "N/A",
                    f"{row.DTE}d" if hasattr(row, 'DTE') else "N/A",
                    f"{int(row.Volumen):,}" if hasattr(row, 'Volumen') else "N/A",
                    f"{int(row.OI):,}" if hasattr(row, 'OI') else "N/A",
                    f"+{int(row.OI_Chg):,}" if hasattr(row, 'OI_Chg') else "N/A",
                    f"{row.IV:.1f}%" if hasattr(row, 'IV') and row.IV > 0 else "N/A",
                    f"{row.Delta:.3f}" if hasattr(row, 'Delta') and row.Delta != 0 else "N/A",
                    f"${row.Último:.2f}" if hasattr(row, 'Último') and row.Último > 0 else "N/A",
                ])
            _tabla_datos_report(doc, headers_pos, rows_pos)

        # Tabla OI Negativo
        if n_neg > 0:
            _agregar_titulo_report(doc, f"OI NEGATIVO — CERRANDO POSICIONES ({n_neg})", level=1)
            p_desc = doc.add_paragraph()
            run_d = p_desc.add_run(
                "Contratos donde el Open Interest disminuyó, indicando posiciones cerradas o ejercidas."
            )
            run_d.font.size = Pt(10)
            run_d.font.italic = True
            run_d.font.name = "Calibri"

            headers_neg = ["#", "Tipo", "Strike", "Vencimiento", "DTE", "Volumen", "OI", "OI Chg", "IV", "Delta", "Último"]
            rows_neg = []
            for i, row in enumerate(df_negativos.head(100).itertuples(), 1):
                rows_neg.append([
                    i,
                    row.Tipo if hasattr(row, 'Tipo') else "N/A",
                    f"${row.Strike:,.1f}" if hasattr(row, 'Strike') else "N/A",
                    row.Vencimiento if hasattr(row, 'Vencimiento') else "N/A",
                    f"{row.DTE}d" if hasattr(row, 'DTE') else "N/A",
                    f"{int(row.Volumen):,}" if hasattr(row, 'Volumen') else "N/A",
                    f"{int(row.OI):,}" if hasattr(row, 'OI') else "N/A",
                    f"{int(row.OI_Chg):,}" if hasattr(row, 'OI_Chg') else "N/A",
                    f"{row.IV:.1f}%" if hasattr(row, 'IV') and row.IV > 0 else "N/A",
                    f"{row.Delta:.3f}" if hasattr(row, 'Delta') and row.Delta != 0 else "N/A",
                    f"${row.Último:.2f}" if hasattr(row, 'Último') and row.Último > 0 else "N/A",
                ])
            _tabla_datos_report(doc, headers_neg, rows_neg)

    else:
        # Sin datos
        p_sin = doc.add_paragraph()
        run_sin = p_sin.add_run("No hay datos de Open Interest disponibles. Ejecuta un escaneo primero.")
        run_sin.font.size = Pt(11)
        run_sin.font.italic = True
        run_sin.font.name = "Calibri"

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Open Interest — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
#   FUNCIÓN 3: REPORTE IMPORTANT COMPANIES
# ============================================================================
def _generar_reporte_important_companies():
    """Genera reporte DOCX con análisis detallado de Important Companies."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — IMPORTANT COMPANIES", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run("Análisis de Empresas Consolidadas y Emergentes")
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run_sub.font.name = "Calibri"
    run_sub.bold = True

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_fecha.font.name = "Calibri"

    doc.add_paragraph("")

    # EMPRESAS CONSOLIDADAS
    if "proyecciones_resultados" in st.session_state and st.session_state.proyecciones_resultados:
        resultados = st.session_state.proyecciones_resultados

        _agregar_titulo_report(doc, f"EMPRESAS CONSOLIDADAS ({len(resultados)})", level=1)
        p_desc = doc.add_paragraph()
        run_d = p_desc.add_run(
            "Grandes corporaciones con historial probado y proyección de crecimiento sostenido. "
            "Análisis fundamental + técnico + sentimiento."
        )
        run_d.font.size = Pt(10)
        run_d.font.italic = True
        run_d.font.name = "Calibri"

        # Resumen métricas
        alta = sum(1 for r in resultados if r.get("veredicto", "").startswith("OPORTUNIDAD"))
        considerar = sum(1 for r in resultados if "CONSIDERAR" in r.get("veredicto", ""))
        mantener = sum(1 for r in resultados if "MANTENER" in r.get("veredicto", ""))
        precaucion = sum(1 for r in resultados if "PRECAUCIÓN" in r.get("veredicto", "") or "PRECAU" in r.get("veredicto", ""))

        _tabla_info_report(doc, {
            "Total Empresas": len(resultados),
            "Oportunidad de Compra": alta,
            "Considerar": considerar,
            "Mantener": mantener,
            "Precaución": precaucion,
        })

        # Tabla comparativa
        headers_comp = ["#", "Ticker", "Empresa", "Precio", "Score Fund.", "Score Téc.", "Score Comb.", "Veredicto",
                        "Crec. Ingresos", "Margen Op.", "P/E Fwd", "PEG", "Tendencia", "RSI", "Target", "Upside"]
        rows_comp = []
        for i, r in enumerate(resultados, 1):
            tecnico = r.get("tecnico", {})
            rows_comp.append([
                i,
                r["symbol"],
                r["nombre"][:30],
                f"${r['precio']:,.2f}",
                f"{r.get('score', 0)}/100",
                f"{r.get('score_tecnico', 0)}/100",
                f"{r.get('score_combinado', 0)}/100",
                r.get("veredicto", "N/A")[:25],
                f"{r['revenue_growth']*100:.1f}%",
                f"{r['operating_margins']*100:.1f}%",
                f"{r['forward_pe']:.1f}x" if r['forward_pe'] > 0 else "N/A",
                f"{r['peg_ratio']:.2f}" if r['peg_ratio'] > 0 else "N/A",
                tecnico.get("tendencia", "N/A"),
                f"{tecnico.get('rsi', 0):.0f}" if tecnico else "N/A",
                f"${r['target_mean']:,.0f}" if r.get('target_mean', 0) > 0 else "N/A",
                f"{r['upside_pct']:.1f}%" if r.get('upside_pct') else "N/A",
            ])
        _tabla_datos_report(doc, headers_comp, rows_comp)

        # Detalle por empresa
        for r in resultados:
            doc.add_page_break()
            _agregar_titulo_report(doc, f"{r['symbol']} — {r['nombre']}", level=2)

            tecnico = r.get("tecnico", {})

            # Info básica
            _tabla_info_report(doc, {
                "Precio Actual": f"${r['precio']:,.2f}",
                "Market Cap": format_market_cap(r.get("market_cap", 0)),
                "Sector": r.get("sector", "N/A"),
                "Industria": r.get("industria", "N/A"),
                "Score Combinado": f"{r.get('score_combinado', 0)}/100",
                "Veredicto": r.get("veredicto", "N/A"),
            })

            # Fundamental
            _agregar_titulo_report(doc, "📊 Análisis Fundamental", level=3)
            _tabla_info_report(doc, {
                "Ingresos Totales": f"${r.get('revenue', 0)/1e9:.1f}B" if r.get('revenue', 0) > 0 else "N/A",
                "Crecimiento Ingresos": f"{r['revenue_growth']*100:.1f}%",
                "Margen Bruto": f"{r['gross_margins']*100:.1f}%",
                "Margen Operativo": f"{r['operating_margins']*100:.1f}%",
                "Margen Neto": f"{r['profit_margins']*100:.1f}%",
                "P/E Forward": f"{r['forward_pe']:.1f}x" if r['forward_pe'] > 0 else "N/A",
                "P/E Trailing": f"{r['trailing_pe']:.1f}x" if r['trailing_pe'] > 0 else "N/A",
                "PEG Ratio": f"{r['peg_ratio']:.2f}" if r['peg_ratio'] > 0 else "N/A",
                "Free Cash Flow": f"${r.get('free_cashflow', 0)/1e9:.1f}B" if r.get('free_cashflow', 0) > 0 else "N/A",
                "Crecimiento Beneficios": f"{r['earnings_growth']*100:.1f}%",
            })

            # Técnico
            if tecnico:
                _agregar_titulo_report(doc, "📈 Análisis Técnico", level=3)
                _tabla_info_report(doc, {
                    "Tendencia": tecnico.get("tendencia", "N/A"),
                    "RSI (14)": f"{tecnico.get('rsi', 0):.0f}",
                    "ADX (14)": f"{tecnico.get('adx', 0):.0f}",
                    "SMA 20": f"${tecnico.get('sma_20', 0):,.2f}",
                    "SMA 50": f"${tecnico.get('sma_50', 0):,.2f}",
                    "SMA 200": f"${tecnico.get('sma_200', 0):,.2f}" if tecnico.get('sma_200', 0) > 0 else "N/A",
                    "Volumen Ratio": f"{tecnico.get('vol_ratio', 0):.2f}x",
                    "Soporte 20d": f"${tecnico.get('soporte_20d', 0):,.2f}",
                    "Resistencia 20d": f"${tecnico.get('resistencia_20d', 0):,.2f}",
                    "Rango 52sem": f"{tecnico.get('rango_52w_pct', 0):.0f}%",
                })

            # Sentimiento
            _agregar_titulo_report(doc, "🎯 Sentimiento", level=3)
            _tabla_info_report(doc, {
                "Recomendación": r.get("recommendation", "N/A").upper(),
                "Número Analistas": r.get("num_analysts", 0),
                "Target Medio": f"${r['target_mean']:,.2f}" if r.get('target_mean', 0) > 0 else "N/A",
                "Target Alto": f"${r['target_high']:,.2f}" if r.get('target_high', 0) > 0 else "N/A",
                "Target Bajo": f"${r['target_low']:,.2f}" if r.get('target_low', 0) > 0 else "N/A",
                "Upside Potencial": f"{r['upside_pct']:.1f}%" if r.get('upside_pct') else "N/A",
                "Beta": f"{r['beta']:.2f}" if r.get('beta', 0) > 0 else "N/A",
                "52 Week Low": f"${r['fifty_two_low']:,.2f}" if r.get('fifty_two_low', 0) > 0 else "N/A",
                "52 Week High": f"${r['fifty_two_high']:,.2f}" if r.get('fifty_two_high', 0) > 0 else "N/A",
            })

            # Razones fundamentales
            if r.get("razones"):
                p_raz = doc.add_paragraph()
                run_raz = p_raz.add_run("Factores del Score Fundamental:")
                run_raz.bold = True
                run_raz.font.size = Pt(10)
                run_raz.font.name = "Calibri"
                for razon in r["razones"]:
                    p_item = doc.add_paragraph(razon, style='List Bullet')
                    p_item.paragraph_format.left_indent = Pt(20)

            # Señales técnicas
            if r.get("señales_tecnicas"):
                p_sen = doc.add_paragraph()
                run_sen = p_sen.add_run("Señales Técnicas:")
                run_sen.bold = True
                run_sen.font.size = Pt(10)
                run_sen.font.name = "Calibri"
                for senal in r["señales_tecnicas"]:
                    p_item = doc.add_paragraph(senal, style='List Bullet')
                    p_item.paragraph_format.left_indent = Pt(20)

    # EMPRESAS EMERGENTES
    if "emergentes_resultados" in st.session_state and st.session_state.emergentes_resultados:
        doc.add_page_break()
        resultados_em = st.session_state.emergentes_resultados

        _agregar_titulo_report(doc, f"EMPRESAS EMERGENTES ({len(resultados_em)})", level=1)
        p_desc_em = doc.add_paragraph()
        run_d_em = p_desc_em.add_run(
            "Empresas innovadoras con alto potencial de crecimiento disruptivo a 10 años."
        )
        run_d_em.font.size = Pt(10)
        run_d_em.font.italic = True
        run_d_em.font.name = "Calibri"

        # Tabla comparativa emergentes
        headers_em = ["#", "Ticker", "Empresa", "Precio", "Score Comb.", "Veredicto",
                      "Crec. Ingresos", "P/E Fwd", "Target", "Upside"]
        rows_em = []
        for i, r in enumerate(resultados_em, 1):
            rows_em.append([
                i,
                r["symbol"],
                r["nombre"][:30],
                f"${r['precio']:,.2f}",
                f"{r.get('score_combinado', 0)}/100",
                r.get("veredicto", "N/A")[:25],
                f"{r['revenue_growth']*100:.1f}%",
                f"{r['forward_pe']:.1f}x" if r['forward_pe'] > 0 else "N/A",
                f"${r['target_mean']:,.0f}" if r.get('target_mean', 0) > 0 else "N/A",
                f"{r['upside_pct']:.1f}%" if r.get('upside_pct') else "N/A",
            ])
        _tabla_datos_report(doc, headers_em, rows_em)

    # Sin datos
    if ("proyecciones_resultados" not in st.session_state or not st.session_state.proyecciones_resultados) and \
       ("emergentes_resultados" not in st.session_state or not st.session_state.emergentes_resultados):
        p_sin = doc.add_paragraph()
        run_sin = p_sin.add_run("No hay datos de análisis disponibles. Ejecuta el análisis en Important Companies primero.")
        run_sin.font.size = Pt(11)
        run_sin.font.italic = True
        run_sin.font.name = "Calibri"

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Important Companies — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
#   FUNCIÓN 3.5: REPORTE DATA ANALYSIS (REAL)
# ============================================================================
def _generar_reporte_data_analysis():
    """Genera reporte DOCX con análisis de sentimiento, soportes y resistencias del Live Scanning."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — DATA ANALYSIS", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run("Análisis de Sentimiento, Soportes y Resistencias")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.name = "Calibri"

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')

    # CONTENIDO
    if not st.session_state.datos_completos:
        p_sin = doc.add_paragraph()
        run_sin = p_sin.add_run("No hay datos de Live Scanning disponibles. Ejecuta el escaneo primero.")
        run_sin.font.size = Pt(11)
        run_sin.font.italic = True
        run_sin.font.name = "Calibri"
    else:
        df_analisis = pd.DataFrame(st.session_state.datos_completos)
        if "Prima_Volumen" in df_analisis.columns:
            df_analisis = df_analisis.rename(columns={"Prima_Volumen": "Prima_Vol"})

        ticker_symbol = st.session_state.get('ticker_actual') or st.session_state.get('ticker_anterior', 'N/A')
        precio_actual = st.session_state.get('precio_subyacente', None)

        # Header con ticker y precio
        _agregar_titulo_report(doc, f"TICKER: {ticker_symbol}", level=1)
        if precio_actual:
            p_precio = doc.add_paragraph()
            run_precio = p_precio.add_run(f"Precio Actual: ${precio_actual:,.2f}")
            run_precio.font.size = Pt(12)
            run_precio.font.bold = True
            run_precio.font.name = "Calibri"

        # ================================================================
        # SENTIMIENTO POR PRIMAS
        # ================================================================
        _agregar_titulo_report(doc, "💰 Desglose de Sentimiento por Primas", level=2)

        df_sent = df_analisis.copy()
        df_sent["_mid"] = (df_sent["Ask"] + df_sent["Bid"]) / 2

        mask_call = df_sent["Tipo"] == "CALL"
        mask_put = df_sent["Tipo"] == "PUT"
        mask_ask = df_sent["Ultimo"] >= df_sent["_mid"]
        mask_bid = df_sent["Ultimo"] < df_sent["_mid"]

        call_ask_val = df_sent.loc[mask_call & mask_ask, "Prima_Vol"].sum()
        call_bid_val = df_sent.loc[mask_call & mask_bid, "Prima_Vol"].sum()
        put_ask_val = df_sent.loc[mask_put & mask_ask, "Prima_Vol"].sum()
        put_bid_val = df_sent.loc[mask_put & mask_bid, "Prima_Vol"].sum()

        total_sent = call_ask_val + call_bid_val + put_ask_val + put_bid_val

        if total_sent > 0:
            bullish_total = call_ask_val + put_bid_val
            bearish_total = call_bid_val + put_ask_val
            net_pct = ((bullish_total - bearish_total) / total_sent) * 100

            _tabla_info_report(doc, {
                "📞 CALL Ask (Compra agresiva)": f"${call_ask_val:,.0f} (+{call_ask_val/total_sent*100:.1f}%)",
                "📞 CALL Bid (Venta agresiva)": f"${call_bid_val:,.0f} (-{call_bid_val/total_sent*100:.1f}%)",
                "📋 PUT Ask (Compra agresiva)": f"${put_ask_val:,.0f} (-{put_ask_val/total_sent*100:.1f}%)",
                "📋 PUT Bid (Venta agresiva)": f"${put_bid_val:,.0f} (+{put_bid_val/total_sent*100:.1f}%)",
                "Total Prima": f"${total_sent:,.0f}",
                "🟢 Alcista Total": f"${bullish_total:,.0f} ({bullish_total/total_sent*100:.1f}%)",
                "🔴 Bajista Total": f"${bearish_total:,.0f} ({bearish_total/total_sent*100:.1f}%)",
                "Sentimiento Neto": f"{'+' if net_pct >= 0 else ''}{net_pct:.1f}% ({'ALCISTA' if net_pct >= 0 else 'BAJISTA'})",
            })

        # ================================================================
        # SOPORTES Y RESISTENCIAS
        # ================================================================
        _agregar_titulo_report(doc, "🛡️ Soportes y Resistencias por Opciones", level=2)

        df_calls_sr = df_analisis[(df_analisis["Tipo"] == "CALL") & (df_analisis["Volumen"] > 0)].copy()
        df_puts_sr = df_analisis[(df_analisis["Tipo"] == "PUT") & (df_analisis["Volumen"] > 0)].copy()

        if not df_calls_sr.empty and not df_puts_sr.empty:
            # Top 5 CALL strikes → Soportes
            top_calls = df_calls_sr.groupby("Strike").agg(
                Vol_Total=("Volumen", "sum"),
                OI_Total=("OI", "sum"),
                Prima_Total=("Prima_Vol", "sum"),
            ).sort_values("Vol_Total", ascending=False).head(5).reset_index()

            # Top 5 PUT strikes → Resistencias
            top_puts = df_puts_sr.groupby("Strike").agg(
                Vol_Total=("Volumen", "sum"),
                OI_Total=("OI", "sum"),
                Prima_Total=("Prima_Vol", "sum"),
            ).sort_values("Vol_Total", ascending=False).head(5).reset_index()

            # Tabla de Soportes
            _agregar_titulo_report(doc, "🟢 Soportes (CALLs más tradeados)", level=3)
            headers_s = ["Nivel", "Strike", "Volumen", "OI", "Prima Total"]
            rows_s = []
            for idx, row in top_calls.iterrows():
                pct_str = ""
                if precio_actual and precio_actual > 0:
                    dist = ((row["Strike"] - precio_actual) / precio_actual) * 100
                    pct_str = f" ({'+' if dist >= 0 else ''}{dist:.1f}%)"
                rows_s.append([
                    f"S{idx+1}",
                    f"${row['Strike']:,.1f}{pct_str}",
                    f"{row['Vol_Total']:,.0f}",
                    f"{row['OI_Total']:,.0f}",
                    f"${row['Prima_Total']:,.0f}",
                ])
            _tabla_datos_report(doc, headers_s, rows_s)

            # Tabla de Resistencias
            _agregar_titulo_report(doc, "🔴 Resistencias (PUTs más tradeados)", level=3)
            headers_r = ["Nivel", "Strike", "Volumen", "OI", "Prima Total"]
            rows_r = []
            for idx, row in top_puts.iterrows():
                pct_str = ""
                if precio_actual and precio_actual > 0:
                    dist = ((row["Strike"] - precio_actual) / precio_actual) * 100
                    pct_str = f" ({'+' if dist >= 0 else ''}{dist:.1f}%)"
                rows_r.append([
                    f"R{idx+1}",
                    f"${row['Strike']:,.1f}{pct_str}",
                    f"{row['Vol_Total']:,.0f}",
                    f"{row['OI_Total']:,.0f}",
                    f"${row['Prima_Total']:,.0f}",
                ])
            _tabla_datos_report(doc, headers_r, rows_r)

        # ================================================================
        # DISTRIBUCIÓN CALL VS PUT
        # ================================================================
        _agregar_titulo_report(doc, "📊 Distribución CALL vs PUT", level=2)

        tipo_counts = df_analisis["Tipo"].value_counts()
        n_calls = tipo_counts.get("CALL", 0)
        n_puts = tipo_counts.get("PUT", 0)
        ratio_pc = n_puts / n_calls if n_calls > 0 else 0

        _tabla_info_report(doc, {
            "Total CALLs": f"{n_calls:,}",
            "Total PUTs": f"{n_puts:,}",
            "Put/Call Ratio": f"{ratio_pc:.3f}",
            "Interpretación": "Mayor actividad en CALLs (alcista)" if ratio_pc < 0.7 else "Ratio neutral",
        })

        # ================================================================
        # TOP 20 POR VOLUMEN
        # ================================================================
        _agregar_titulo_report(doc, "🎯 Top 20 Strikes por Volumen", level=2)

        vol_cols = ["Vencimiento", "Tipo", "Strike", "Volumen", "OI", "OI_Chg", "IV", "Ultimo", "Prima_Vol"]
        top_vol = df_analisis.nlargest(20, "Volumen")[[c for c in vol_cols if c in df_analisis.columns]].reset_index(drop=True)

        has_oi_chg = "OI_Chg" in top_vol.columns
        headers_vol = ["#", "Vencimiento", "Tipo", "Strike", "Volumen", "OI"]
        if has_oi_chg:
            headers_vol.append("OI Chg")
        headers_vol.extend(["IV", "Último", "Prima Total"])

        rows_vol = []
        for i, row in top_vol.iterrows():
            row_data = [
                i + 1,
                row.get("Vencimiento", "N/A"),
                row.get("Tipo", "N/A"),
                f"${row.get('Strike', 0):,.1f}",
                f"{row.get('Volumen', 0):,}",
                f"{row.get('OI', 0):,}",
            ]
            if has_oi_chg:
                oi_chg = row.get('OI_Chg', 0)
                row_data.append(f"+{int(oi_chg):,}" if oi_chg > 0 else f"{int(oi_chg):,}")
            row_data.extend([
                f"{row.get('IV', 0):.2f}%" if row.get('IV', 0) > 0 else "N/A",
                f"${row.get('Ultimo', 0):.2f}",
                f"${row.get('Prima_Vol', 0):,.0f}",
            ])
            rows_vol.append(row_data)
        _tabla_datos_report(doc, headers_vol, rows_vol)

        # ================================================================
        # TOP 20 POR OI
        # ================================================================
        _agregar_titulo_report(doc, "🏛️ Top 20 Strikes por Open Interest", level=2)

        oi_cols = ["Vencimiento", "Tipo", "Strike", "OI", "OI_Chg", "Volumen", "IV", "Ultimo", "Prima_Vol"]
        top_oi = df_analisis.nlargest(20, "OI")[[c for c in oi_cols if c in df_analisis.columns]].reset_index(drop=True)

        has_oi_chg_oi = "OI_Chg" in top_oi.columns
        headers_oi = ["#", "Vencimiento", "Tipo", "Strike", "OI"]
        if has_oi_chg_oi:
            headers_oi.append("OI Chg")
        headers_oi.extend(["Volumen", "IV", "Último", "Prima Total"])

        rows_oi = []
        for i, row in top_oi.iterrows():
            row_data = [
                i + 1,
                row.get("Vencimiento", "N/A"),
                row.get("Tipo", "N/A"),
                f"${row.get('Strike', 0):,.1f}",
                f"{row.get('OI', 0):,}",
            ]
            if has_oi_chg_oi:
                oi_chg = row.get('OI_Chg', 0)
                row_data.append(f"+{int(oi_chg):,}" if oi_chg > 0 else f"{int(oi_chg):,}")
            row_data.extend([
                f"{row.get('Volumen', 0):,}",
                f"{row.get('IV', 0):.2f}%" if row.get('IV', 0) > 0 else "N/A",
                f"${row.get('Ultimo', 0):.2f}",
                f"${row.get('Prima_Vol', 0):,.0f}",
            ])
            rows_oi.append(row_data)
        _tabla_datos_report(doc, headers_oi, rows_oi)

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Data Analysis — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
#   FUNCIÓN 4: REPORTE RANGE
# ============================================================================
def _generar_reporte_range():
    """Genera reporte DOCX con información del Rango Esperado."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — RANGO ESPERADO", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    r = st.session_state.rango_resultado
    ticker_name = r.get("symbol", "N/A")

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run(f"Ticker: {ticker_name}")
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run_sub.font.name = "Calibri"
    run_sub.bold = True

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_fecha.font.name = "Calibri"

    doc.add_paragraph("")

    # Explicación
    _agregar_titulo_report(doc, "¿QUÉ ES EL RANGO ESPERADO?", level=1)
    p_exp = doc.add_paragraph()
    run_exp = p_exp.add_run(
        "El rango esperado es una estimación estadística del movimiento probable del precio del activo "
        "hasta la fecha de expiración, basado en la volatilidad implícita (IV) de las opciones. "
        "Se calcula con una desviación estándar (1σ), lo que significa que hay aproximadamente 68% de "
        "probabilidad de que el precio permanezca dentro del rango calculado."
    )
    run_exp.font.size = Pt(10)
    run_exp.font.name = "Calibri"

    doc.add_paragraph("")

    # Parámetros del cálculo
    dias = r.get('dias_restantes')
    _agregar_titulo_report(doc, "PARÁMETROS DEL CÁLCULO", level=1)
    _tabla_info_report(doc, {
        "Símbolo": ticker_name,
        "Precio Actual del Subyacente": f"${r['underlying_price']:,.2f}",
        "Fecha de Expiración": r["expiration"],
        "Días Restantes (DTE)": dias if dias else "N/A",
        "Delta Objetivo": f"±{r.get('target_delta', 'N/A')}",
    })

    # Rango calculado
    _agregar_titulo_report(doc, "RANGO DE PRECIOS ESPERADO (1σ)", level=1)
    _tabla_info_report(doc, {
        "Rango Inferior": f"${r['expected_range_low']:,.2f}",
        "Precio Actual": f"${r['underlying_price']:,.2f}",
        "Rango Superior": f"${r['expected_range_high']:,.2f}",
        "Bajada Esperada": f"-${r['downside_points']:,.2f} (-{r['downside_percent']:.2f}%)",
        "Subida Esperada": f"+${r['upside_points']:,.2f} (+{r['upside_percent']:.2f}%)",
        "Rango Total de Movimiento": f"${r['total_range_points']:,.2f} ({r['total_range_pct']:.2f}%)",
    })

    # Contratos utilizados
    _agregar_titulo_report(doc, "CONTRATOS UTILIZADOS EN EL CÁLCULO", level=1)
    p_cont = doc.add_paragraph()
    run_cont = p_cont.add_run(
        "El rango se calcula utilizando las opciones Call y Put con deltas más cercanos al objetivo configurado."
    )
    run_cont.font.size = Pt(10)
    run_cont.font.italic = True
    run_cont.font.name = "Calibri"

    _tabla_info_report(doc, {
        "Call Strike": f"${r['call_strike']}",
        "Call Delta": f"{r['call_delta']}",
        "Call IV": f"{r['call_iv']:.1f}%",
        "Put Strike": f"${r['put_strike']}",
        "Put Delta": f"{r['put_delta']}",
        "Put IV": f"{r['put_iv']:.1f}%",
    })

    # Interpretación
    _agregar_titulo_report(doc, "INTERPRETACIÓN", level=1)
    p_int = doc.add_paragraph()
    run_int = p_int.add_run(
        f"Basándose en la volatilidad implícita actual, se espera que {ticker_name} se mueva "
        f"entre ${r['expected_range_low']:,.2f} y ${r['expected_range_high']:,.2f} antes del "
        f"{r['expiration']}. Esto representa un rango de movimiento de ±{r['total_range_pct']:.1f}%.\n\n"
        f"Este rango puede utilizarse para:\n"
        f"• Planificar estrategias de trading direccionales (si esperas movimiento fuera del rango)\n"
        f"• Diseñar estrategias neutrales (si esperas que el precio permanezca dentro del rango)\n"
        f"• Identificar niveles de soporte y resistencia probables\n"
        f"• Evaluar el riesgo de posiciones existentes"
    )
    run_int.font.size = Pt(10)
    run_int.font.name = "Calibri"

    # Aviso
    doc.add_paragraph("")
    p_aviso = doc.add_paragraph()
    run_aviso = p_aviso.add_run(
        "⚠️ AVISO: Este cálculo es una estimación estadística basada en la volatilidad implícita "
        "y no garantiza que el precio permanecerá dentro del rango. Los movimientos del mercado "
        "pueden ser impredecibles, especialmente ante eventos inesperados o noticias significativas."
    )
    run_aviso.font.size = Pt(9)
    run_aviso.font.italic = True
    run_aviso.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_aviso.font.name = "Calibri"

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Rango Esperado — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
