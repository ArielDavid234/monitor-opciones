# -*- coding: utf-8 -*-
import io
import numpy as np
import pandas as pd
import streamlit as st
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from reports.docx_helpers import (
    _estilo_celda_report, _agregar_titulo_report,
    _tabla_info_report, _tabla_datos_report,
)

def _generar_reporte_range():
    """Genera reporte DOCX con información del Rango Esperado."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — RANGO ESPERADO", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    r = st.session_state.rango_resultado
    ticker_name = r.get("symbol", "N/A")

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run(f"Ticker: {ticker_name}")
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run_sub.font.name = "Calibri"
    run_sub.bold = True

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_fecha.font.name = "Calibri"

    doc.add_paragraph("")

    # Explicación
    _agregar_titulo_report(doc, "¿QUÉ ES EL RANGO ESPERADO?", level=1)
    p_exp = doc.add_paragraph()
    run_exp = p_exp.add_run(
        "El rango esperado es una estimación estadística del movimiento probable del precio del activo "
        "hasta la fecha de expiración, basado en la volatilidad implícita (IV) de las opciones. "
        "Se calcula con una desviación estándar (1σ), lo que significa que hay aproximadamente 68% de "
        "probabilidad de que el precio permanezca dentro del rango calculado."
    )
    run_exp.font.size = Pt(10)
    run_exp.font.name = "Calibri"

    doc.add_paragraph("")

    # Parámetros del cálculo
    dias = r.get('dias_restantes')
    _agregar_titulo_report(doc, "PARÁMETROS DEL CÁLCULO", level=1)
    _tabla_info_report(doc, {
        "Símbolo": ticker_name,
        "Precio Actual del Subyacente": f"${r['underlying_price']:,.2f}",
        "Fecha de Expiración": r["expiration"],
        "Días Restantes (DTE)": dias if dias else "N/A",
        "Delta Objetivo": f"±{r.get('target_delta', 'N/A')}",
    })

    # Rango calculado
    _agregar_titulo_report(doc, "RANGO DE PRECIOS ESPERADO (1σ)", level=1)
    _tabla_info_report(doc, {
        "Rango Inferior": f"${r['expected_range_low']:,.2f}",
        "Precio Actual": f"${r['underlying_price']:,.2f}",
        "Rango Superior": f"${r['expected_range_high']:,.2f}",
        "Bajada Esperada": f"-${r['downside_points']:,.2f} (-{r['downside_percent']:.2f}%)",
        "Subida Esperada": f"+${r['upside_points']:,.2f} (+{r['upside_percent']:.2f}%)",
        "Rango Total de Movimiento": f"${r['total_range_points']:,.2f} ({r['total_range_pct']:.2f}%)",
    })

    # Contratos utilizados
    _agregar_titulo_report(doc, "CONTRATOS UTILIZADOS EN EL CÁLCULO", level=1)
    p_cont = doc.add_paragraph()
    run_cont = p_cont.add_run(
        "El rango se calcula utilizando las opciones Call y Put con deltas más cercanos al objetivo configurado."
    )
    run_cont.font.size = Pt(10)
    run_cont.font.italic = True
    run_cont.font.name = "Calibri"

    _tabla_info_report(doc, {
        "Call Strike": f"${r['call_strike']}",
        "Call Delta": f"{r['call_delta']}",
        "Call IV": f"{r['call_iv']:.1f}%",
        "Put Strike": f"${r['put_strike']}",
        "Put Delta": f"{r['put_delta']}",
        "Put IV": f"{r['put_iv']:.1f}%",
    })

    # Interpretación
    _agregar_titulo_report(doc, "INTERPRETACIÓN", level=1)
    p_int = doc.add_paragraph()
    run_int = p_int.add_run(
        f"Basándose en la volatilidad implícita actual, se espera que {ticker_name} se mueva "
        f"entre ${r['expected_range_low']:,.2f} y ${r['expected_range_high']:,.2f} antes del "
        f"{r['expiration']}. Esto representa un rango de movimiento de ±{r['total_range_pct']:.1f}%.\n\n"
        f"Este rango puede utilizarse para:\n"
        f"• Planificar estrategias de trading direccionales (si esperas movimiento fuera del rango)\n"
        f"• Diseñar estrategias neutrales (si esperas que el precio permanezca dentro del rango)\n"
        f"• Identificar niveles de soporte y resistencia probables\n"
        f"• Evaluar el riesgo de posiciones existentes"
    )
    run_int.font.size = Pt(10)
    run_int.font.name = "Calibri"

    # Aviso
    doc.add_paragraph("")
    p_aviso = doc.add_paragraph()
    run_aviso = p_aviso.add_run(
        "⚠️ AVISO: Este cálculo es una estimación estadística basada en la volatilidad implícita "
        "y no garantiza que el precio permanecerá dentro del rango. Los movimientos del mercado "
        "pueden ser impredecibles, especialmente ante eventos inesperados o noticias significativas."
    )
    run_aviso.font.size = Pt(9)
    run_aviso.font.italic = True
    run_aviso.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_aviso.font.name = "Calibri"

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Rango Esperado — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
