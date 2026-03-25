# -*- coding: utf-8 -*-
"""
Generadores de reportes DOCX profesionales.
Extraídos de app_web.py — cero cambios de lógica.
"""
import io
import numpy as np
import pandas as pd
import streamlit as st
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from utils.formatters import (
    _fmt_precio, _fmt_entero, _fmt_monto, _fmt_iv, _fmt_lado,
    determinar_sentimiento,
)
from utils.helpers import _enriquecer_datos_opcion
from ui.components import format_market_cap


# ============================================================================
#                    HELPERS PARA GENERAR REPORTES DOCX
# ============================================================================
def _estilo_celda_report(cell, texto, negrita=False, color_fondo=None, color_texto=None, size=9, align="left"):
    """Aplica formato a una celda de tabla Word."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    run = p.add_run(str(texto))
    run.bold = negrita
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color_texto:
        run.font.color.rgb = color_texto
    if color_fondo:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(qn("w:shd"), {
            qn("w:fill"): color_fondo,
            qn("w:val"): "clear",
        })
        tcPr.append(shading)


def _agregar_titulo_report(doc, texto, level=2):
    """Agrega un título de sección con formato."""
    heading = doc.add_heading(texto, level=level)
    for run in heading.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)


def _tabla_info_report(doc, datos_dict, titulo=None):
    """Tabla de 2 columnas Campo/Valor para info resumida."""
    if titulo:
        p = doc.add_paragraph()
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for campo, valor in datos_dict.items():
        row = table.add_row()
        _estilo_celda_report(row.cells[0], campo, negrita=True, size=10)
        _estilo_celda_report(row.cells[1], str(valor), size=10)
    doc.add_paragraph("")


def _tabla_datos_report(doc, headers, rows_data):
    """Tabla con encabezados y filas de datos."""
    if not rows_data:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Encabezados
    for i, h in enumerate(headers):
        _estilo_celda_report(
            table.rows[0].cells[i], h,
            negrita=True, size=9,
            color_fondo="1E3A5F",
            color_texto=RGBColor(0xFF, 0xFF, 0xFF),
            align="center",
        )
    # Datos
    for row_idx, row_data in enumerate(rows_data):
        row = table.add_row()
        bg = "F0F4F8" if row_idx % 2 == 0 else None
        for i, val in enumerate(row_data):
            _estilo_celda_report(row.cells[i], str(val), size=9, color_fondo=bg)
    doc.add_paragraph("")


# ============================================================================
