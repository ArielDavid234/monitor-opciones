# -*- coding: utf-8 -*-
import io
import numpy as np
import pandas as pd
import streamlit as st
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from reports.docx_helpers import (
    _estilo_celda_report, _agregar_titulo_report,
    _tabla_info_report, _tabla_datos_report,
)

def _generar_reporte_data_analysis():
    """Genera reporte DOCX con análisis de sentimiento, soportes y resistencias del Live Scanning."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — DATA ANALYSIS", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run("Análisis de Sentimiento, Soportes y Resistencias")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.name = "Calibri"

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')

    # CONTENIDO
    if not st.session_state.datos_completos:
        p_sin = doc.add_paragraph()
        run_sin = p_sin.add_run("No hay datos de Live Scanning disponibles. Ejecuta el escaneo primero.")
        run_sin.font.size = Pt(11)
        run_sin.font.italic = True
        run_sin.font.name = "Calibri"
    else:
        df_analisis = pd.DataFrame(st.session_state.datos_completos)
        if "Prima_Volumen" in df_analisis.columns:
            df_analisis = df_analisis.rename(columns={"Prima_Volumen": "Prima_Vol"})

        ticker_symbol = st.session_state.get('ticker_actual') or st.session_state.get('ticker_anterior', 'N/A')
        precio_actual = st.session_state.get('precio_subyacente', None)

        # Header con ticker y precio
        _agregar_titulo_report(doc, f"TICKER: {ticker_symbol}", level=1)
        if precio_actual:
            p_precio = doc.add_paragraph()
            run_precio = p_precio.add_run(f"Precio Actual: ${precio_actual:,.2f}")
            run_precio.font.size = Pt(12)
            run_precio.font.bold = True
            run_precio.font.name = "Calibri"

        # ================================================================
        # SENTIMIENTO POR PRIMAS
        # ================================================================
        _agregar_titulo_report(doc, "💰 Desglose de Sentimiento por Primas", level=2)

        df_sent = df_analisis.copy()
        df_sent["_mid"] = (df_sent["Ask"] + df_sent["Bid"]) / 2

        mask_call = df_sent["Tipo"] == "CALL"
        mask_put = df_sent["Tipo"] == "PUT"
        mask_ask = df_sent["Ultimo"] >= df_sent["_mid"]
        mask_bid = df_sent["Ultimo"] < df_sent["_mid"]

        call_ask_val = df_sent.loc[mask_call & mask_ask, "Prima_Vol"].sum()
        call_bid_val = df_sent.loc[mask_call & mask_bid, "Prima_Vol"].sum()
        put_ask_val = df_sent.loc[mask_put & mask_ask, "Prima_Vol"].sum()
        put_bid_val = df_sent.loc[mask_put & mask_bid, "Prima_Vol"].sum()

        total_sent = call_ask_val + call_bid_val + put_ask_val + put_bid_val

        if total_sent > 0:
            bullish_total = call_ask_val + put_bid_val
            bearish_total = call_bid_val + put_ask_val
            net_pct = ((bullish_total - bearish_total) / total_sent) * 100

            _tabla_info_report(doc, {
                "📞 CALL Ask (Compra agresiva)": f"${call_ask_val:,.0f} (+{call_ask_val/total_sent*100:.1f}%)",
                "📞 CALL Bid (Venta agresiva)": f"${call_bid_val:,.0f} (-{call_bid_val/total_sent*100:.1f}%)",
                "📋 PUT Ask (Compra agresiva)": f"${put_ask_val:,.0f} (-{put_ask_val/total_sent*100:.1f}%)",
                "📋 PUT Bid (Venta agresiva)": f"${put_bid_val:,.0f} (+{put_bid_val/total_sent*100:.1f}%)",
                "Total Prima": f"${total_sent:,.0f}",
                "🟢 Alcista Total": f"${bullish_total:,.0f} ({bullish_total/total_sent*100:.1f}%)",
                "🔴 Bajista Total": f"${bearish_total:,.0f} ({bearish_total/total_sent*100:.1f}%)",
                "Sentimiento Neto": f"{'+' if net_pct >= 0 else ''}{net_pct:.1f}% ({'ALCISTA' if net_pct >= 0 else 'BAJISTA'})",
            })

        # ================================================================
        # SOPORTES Y RESISTENCIAS
        # ================================================================
        _agregar_titulo_report(doc, "🛡️ Soportes y Resistencias por Opciones", level=2)

        df_calls_sr = df_analisis[(df_analisis["Tipo"] == "CALL") & (df_analisis["Volumen"] > 0)].copy()
        df_puts_sr = df_analisis[(df_analisis["Tipo"] == "PUT") & (df_analisis["Volumen"] > 0)].copy()

        if not df_calls_sr.empty and not df_puts_sr.empty:
            # Top 5 CALL strikes → Soportes
            top_calls = df_calls_sr.groupby("Strike").agg(
                Vol_Total=("Volumen", "sum"),
                OI_Total=("OI", "sum"),
                Prima_Total=("Prima_Vol", "sum"),
            ).sort_values("Vol_Total", ascending=False).head(5).reset_index()

            # Top 5 PUT strikes → Resistencias
            top_puts = df_puts_sr.groupby("Strike").agg(
                Vol_Total=("Volumen", "sum"),
                OI_Total=("OI", "sum"),
                Prima_Total=("Prima_Vol", "sum"),
            ).sort_values("Vol_Total", ascending=False).head(5).reset_index()

            # Tabla de Soportes
            _agregar_titulo_report(doc, "🟢 Soportes (CALLs más tradeados)", level=3)
            headers_s = ["Nivel", "Strike", "Volumen", "OI", "Prima Total"]
            rows_s = []
            for idx, row in top_calls.iterrows():
                pct_str = ""
                if precio_actual and precio_actual > 0:
                    dist = ((row["Strike"] - precio_actual) / precio_actual) * 100
                    pct_str = f" ({'+' if dist >= 0 else ''}{dist:.1f}%)"
                rows_s.append([
                    f"S{idx+1}",
                    f"${row['Strike']:,.1f}{pct_str}",
                    f"{row['Vol_Total']:,.0f}",
                    f"{row['OI_Total']:,.0f}",
                    f"${row['Prima_Total']:,.0f}",
                ])
            _tabla_datos_report(doc, headers_s, rows_s)

            # Tabla de Resistencias
            _agregar_titulo_report(doc, "🔴 Resistencias (PUTs más tradeados)", level=3)
            headers_r = ["Nivel", "Strike", "Volumen", "OI", "Prima Total"]
            rows_r = []
            for idx, row in top_puts.iterrows():
                pct_str = ""
                if precio_actual and precio_actual > 0:
                    dist = ((row["Strike"] - precio_actual) / precio_actual) * 100
                    pct_str = f" ({'+' if dist >= 0 else ''}{dist:.1f}%)"
                rows_r.append([
                    f"R{idx+1}",
                    f"${row['Strike']:,.1f}{pct_str}",
                    f"{row['Vol_Total']:,.0f}",
                    f"{row['OI_Total']:,.0f}",
                    f"${row['Prima_Total']:,.0f}",
                ])
            _tabla_datos_report(doc, headers_r, rows_r)

        # ================================================================
        # DISTRIBUCIÓN CALL VS PUT
        # ================================================================
        _agregar_titulo_report(doc, "📊 Distribución CALL vs PUT", level=2)

        tipo_counts = df_analisis["Tipo"].value_counts()
        n_calls = tipo_counts.get("CALL", 0)
        n_puts = tipo_counts.get("PUT", 0)
        ratio_pc = n_puts / n_calls if n_calls > 0 else 0

        _tabla_info_report(doc, {
            "Total CALLs": f"{n_calls:,}",
            "Total PUTs": f"{n_puts:,}",
            "Put/Call Ratio": f"{ratio_pc:.3f}",
            "Interpretación": "Mayor actividad en CALLs (alcista)" if ratio_pc < 0.7 else "Ratio neutral",
        })

        # ================================================================
        # TOP 20 POR VOLUMEN
        # ================================================================
        _agregar_titulo_report(doc, "🎯 Top 20 Strikes por Volumen", level=2)

        vol_cols = ["Vencimiento", "Tipo", "Strike", "Volumen", "OI", "OI_Chg", "IV", "Ultimo", "Prima_Vol"]
        top_vol = df_analisis.nlargest(20, "Volumen")[[c for c in vol_cols if c in df_analisis.columns]].reset_index(drop=True)

        has_oi_chg = "OI_Chg" in top_vol.columns
        headers_vol = ["#", "Vencimiento", "Tipo", "Strike", "Volumen", "OI"]
        if has_oi_chg:
            headers_vol.append("OI Chg")
        headers_vol.extend(["IV", "Último", "Prima Total"])

        rows_vol = []
        for i, row in top_vol.iterrows():
            row_data = [
                i + 1,
                row.get("Vencimiento", "N/A"),
                row.get("Tipo", "N/A"),
                f"${row.get('Strike', 0):,.1f}",
                f"{row.get('Volumen', 0):,}",
                f"{row.get('OI', 0):,}",
            ]
            if has_oi_chg:
                oi_chg = row.get('OI_Chg', 0)
                row_data.append(f"+{int(oi_chg):,}" if oi_chg > 0 else f"{int(oi_chg):,}")
            row_data.extend([
                f"{row.get('IV', 0):.2f}%" if row.get('IV', 0) > 0 else "N/A",
                f"${row.get('Ultimo', 0):.2f}",
                f"${row.get('Prima_Vol', 0):,.0f}",
            ])
            rows_vol.append(row_data)
        _tabla_datos_report(doc, headers_vol, rows_vol)

        # ================================================================
        # TOP 20 POR OI
        # ================================================================
        _agregar_titulo_report(doc, "🏛️ Top 20 Strikes por Open Interest", level=2)

        oi_cols = ["Vencimiento", "Tipo", "Strike", "OI", "OI_Chg", "Volumen", "IV", "Ultimo", "Prima_Vol"]
        top_oi = df_analisis.nlargest(20, "OI")[[c for c in oi_cols if c in df_analisis.columns]].reset_index(drop=True)

        has_oi_chg_oi = "OI_Chg" in top_oi.columns
        headers_oi = ["#", "Vencimiento", "Tipo", "Strike", "OI"]
        if has_oi_chg_oi:
            headers_oi.append("OI Chg")
        headers_oi.extend(["Volumen", "IV", "Último", "Prima Total"])

        rows_oi = []
        for i, row in top_oi.iterrows():
            row_data = [
                i + 1,
                row.get("Vencimiento", "N/A"),
                row.get("Tipo", "N/A"),
                f"${row.get('Strike', 0):,.1f}",
                f"{row.get('OI', 0):,}",
            ]
            if has_oi_chg_oi:
                oi_chg = row.get('OI_Chg', 0)
                row_data.append(f"+{int(oi_chg):,}" if oi_chg > 0 else f"{int(oi_chg):,}")
            row_data.extend([
                f"{row.get('Volumen', 0):,}",
                f"{row.get('IV', 0):.2f}%" if row.get('IV', 0) > 0 else "N/A",
                f"${row.get('Ultimo', 0):.2f}",
                f"${row.get('Prima_Vol', 0):,.0f}",
            ])
            rows_oi.append(row_data)
        _tabla_datos_report(doc, headers_oi, rows_oi)

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Data Analysis — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
