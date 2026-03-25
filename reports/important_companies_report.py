# -*- coding: utf-8 -*-
import io
import numpy as np
import pandas as pd
import streamlit as st
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from reports.docx_helpers import (
    _estilo_celda_report, _agregar_titulo_report,
    _tabla_info_report, _tabla_datos_report,
)
from ui.components import format_market_cap

def _generar_reporte_important_companies():
    """Genera reporte DOCX con análisis detallado de Important Companies."""
    doc = Document()

    # Configurar página
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Portada
    doc.add_paragraph("")
    titulo = doc.add_heading("REPORTE — IMPORTANT COMPANIES", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run("Análisis de Empresas Consolidadas y Emergentes")
    run_sub.font.size = Pt(16)
    run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run_sub.font.name = "Calibri"
    run_sub.bold = True

    fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
    run_fecha.font.size = Pt(11)
    run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run_fecha.font.name = "Calibri"

    doc.add_paragraph("")

    # EMPRESAS CONSOLIDADAS
    if "proyecciones_resultados" in st.session_state and st.session_state.proyecciones_resultados:
        resultados = st.session_state.proyecciones_resultados

        _agregar_titulo_report(doc, f"EMPRESAS CONSOLIDADAS ({len(resultados)})", level=1)
        p_desc = doc.add_paragraph()
        run_d = p_desc.add_run(
            "Grandes corporaciones con historial probado y proyección de crecimiento sostenido. "
            "Análisis fundamental + técnico + sentimiento."
        )
        run_d.font.size = Pt(10)
        run_d.font.italic = True
        run_d.font.name = "Calibri"

        # Resumen métricas
        alta = sum(1 for r in resultados if r.get("veredicto", "").startswith("OPORTUNIDAD"))
        considerar = sum(1 for r in resultados if "CONSIDERAR" in r.get("veredicto", ""))
        mantener = sum(1 for r in resultados if "MANTENER" in r.get("veredicto", ""))
        precaucion = sum(1 for r in resultados if "PRECAUCIÓN" in r.get("veredicto", "") or "PRECAU" in r.get("veredicto", ""))

        _tabla_info_report(doc, {
            "Total Empresas": len(resultados),
            "Oportunidad de Compra": alta,
            "Considerar": considerar,
            "Mantener": mantener,
            "Precaución": precaucion,
        })

        # Tabla comparativa
        headers_comp = ["#", "Ticker", "Empresa", "Precio", "Score Fund.", "Score Téc.", "Score Comb.", "Veredicto",
                        "Crec. Ingresos", "Margen Op.", "P/E Fwd", "PEG", "Tendencia", "RSI", "Target", "Upside"]
        rows_comp = []
        for i, r in enumerate(resultados, 1):
            tecnico = r.get("tecnico", {})
            rows_comp.append([
                i,
                r["symbol"],
                r["nombre"][:30],
                f"${r['precio']:,.2f}",
                f"{r.get('score', 0)}/100",
                f"{r.get('score_tecnico', 0)}/100",
                f"{r.get('score_combinado', 0)}/100",
                r.get("veredicto", "N/A")[:25],
                f"{r['revenue_growth']*100:.1f}%",
                f"{r['operating_margins']*100:.1f}%",
                f"{r['forward_pe']:.1f}x" if r['forward_pe'] > 0 else "N/A",
                f"{r['peg_ratio']:.2f}" if r['peg_ratio'] > 0 else "N/A",
                tecnico.get("tendencia", "N/A"),
                f"{tecnico.get('rsi', 0):.0f}" if tecnico else "N/A",
                f"${r['target_mean']:,.0f}" if r.get('target_mean', 0) > 0 else "N/A",
                f"{r['upside_pct']:.1f}%" if r.get('upside_pct') else "N/A",
            ])
        _tabla_datos_report(doc, headers_comp, rows_comp)

        # Detalle por empresa
        for r in resultados:
            doc.add_page_break()
            _agregar_titulo_report(doc, f"{r['symbol']} — {r['nombre']}", level=2)

            tecnico = r.get("tecnico", {})

            # Info básica
            _tabla_info_report(doc, {
                "Precio Actual": f"${r['precio']:,.2f}",
                "Market Cap": format_market_cap(r.get("market_cap", 0)),
                "Sector": r.get("sector", "N/A"),
                "Industria": r.get("industria", "N/A"),
                "Score Combinado": f"{r.get('score_combinado', 0)}/100",
                "Veredicto": r.get("veredicto", "N/A"),
            })

            # Fundamental
            _agregar_titulo_report(doc, "📊 Análisis Fundamental", level=3)
            _tabla_info_report(doc, {
                "Ingresos Totales": f"${r.get('revenue', 0)/1e9:.1f}B" if r.get('revenue', 0) > 0 else "N/A",
                "Crecimiento Ingresos": f"{r['revenue_growth']*100:.1f}%",
                "Margen Bruto": f"{r['gross_margins']*100:.1f}%",
                "Margen Operativo": f"{r['operating_margins']*100:.1f}%",
                "Margen Neto": f"{r['profit_margins']*100:.1f}%",
                "P/E Forward": f"{r['forward_pe']:.1f}x" if r['forward_pe'] > 0 else "N/A",
                "P/E Trailing": f"{r['trailing_pe']:.1f}x" if r['trailing_pe'] > 0 else "N/A",
                "PEG Ratio": f"{r['peg_ratio']:.2f}" if r['peg_ratio'] > 0 else "N/A",
                "Free Cash Flow": f"${r.get('free_cashflow', 0)/1e9:.1f}B" if r.get('free_cashflow', 0) > 0 else "N/A",
                "Crecimiento Beneficios": f"{r['earnings_growth']*100:.1f}%",
            })

            # Técnico
            if tecnico:
                _agregar_titulo_report(doc, "📈 Análisis Técnico", level=3)
                _tabla_info_report(doc, {
                    "Tendencia": tecnico.get("tendencia", "N/A"),
                    "RSI (14)": f"{tecnico.get('rsi', 0):.0f}",
                    "ADX (14)": f"{tecnico.get('adx', 0):.0f}",
                    "SMA 20": f"${tecnico.get('sma_20', 0):,.2f}",
                    "SMA 50": f"${tecnico.get('sma_50', 0):,.2f}",
                    "SMA 200": f"${tecnico.get('sma_200', 0):,.2f}" if tecnico.get('sma_200', 0) > 0 else "N/A",
                    "Volumen Ratio": f"{tecnico.get('vol_ratio', 0):.2f}x",
                    "Soporte 20d": f"${tecnico.get('soporte_20d', 0):,.2f}",
                    "Resistencia 20d": f"${tecnico.get('resistencia_20d', 0):,.2f}",
                    "Rango 52sem": f"{tecnico.get('rango_52w_pct', 0):.0f}%",
                })

            # Sentimiento
            _agregar_titulo_report(doc, "🎯 Sentimiento", level=3)
            _tabla_info_report(doc, {
                "Recomendación": r.get("recommendation", "N/A").upper(),
                "Número Analistas": r.get("num_analysts", 0),
                "Target Medio": f"${r['target_mean']:,.2f}" if r.get('target_mean', 0) > 0 else "N/A",
                "Target Alto": f"${r['target_high']:,.2f}" if r.get('target_high', 0) > 0 else "N/A",
                "Target Bajo": f"${r['target_low']:,.2f}" if r.get('target_low', 0) > 0 else "N/A",
                "Upside Potencial": f"{r['upside_pct']:.1f}%" if r.get('upside_pct') else "N/A",
                "Beta": f"{r['beta']:.2f}" if r.get('beta', 0) > 0 else "N/A",
                "52 Week Low": f"${r['fifty_two_low']:,.2f}" if r.get('fifty_two_low', 0) > 0 else "N/A",
                "52 Week High": f"${r['fifty_two_high']:,.2f}" if r.get('fifty_two_high', 0) > 0 else "N/A",
            })

            # Razones fundamentales
            if r.get("razones"):
                p_raz = doc.add_paragraph()
                run_raz = p_raz.add_run("Factores del Score Fundamental:")
                run_raz.bold = True
                run_raz.font.size = Pt(10)
                run_raz.font.name = "Calibri"
                for razon in r["razones"]:
                    p_item = doc.add_paragraph(razon, style='List Bullet')
                    p_item.paragraph_format.left_indent = Pt(20)

            # Señales técnicas
            if r.get("señales_tecnicas"):
                p_sen = doc.add_paragraph()
                run_sen = p_sen.add_run("Señales Técnicas:")
                run_sen.bold = True
                run_sen.font.size = Pt(10)
                run_sen.font.name = "Calibri"
                for senal in r["señales_tecnicas"]:
                    p_item = doc.add_paragraph(senal, style='List Bullet')
                    p_item.paragraph_format.left_indent = Pt(20)

    # EMPRESAS EMERGENTES
    if "emergentes_resultados" in st.session_state and st.session_state.emergentes_resultados:
        doc.add_page_break()
        resultados_em = st.session_state.emergentes_resultados

        _agregar_titulo_report(doc, f"EMPRESAS EMERGENTES ({len(resultados_em)})", level=1)
        p_desc_em = doc.add_paragraph()
        run_d_em = p_desc_em.add_run(
            "Empresas innovadoras con alto potencial de crecimiento disruptivo a 10 años."
        )
        run_d_em.font.size = Pt(10)
        run_d_em.font.italic = True
        run_d_em.font.name = "Calibri"

        # Tabla comparativa emergentes
        headers_em = ["#", "Ticker", "Empresa", "Precio", "Score Comb.", "Veredicto",
                      "Crec. Ingresos", "P/E Fwd", "Target", "Upside"]
        rows_em = []
        for i, r in enumerate(resultados_em, 1):
            rows_em.append([
                i,
                r["symbol"],
                r["nombre"][:30],
                f"${r['precio']:,.2f}",
                f"{r.get('score_combinado', 0)}/100",
                r.get("veredicto", "N/A")[:25],
                f"{r['revenue_growth']*100:.1f}%",
                f"{r['forward_pe']:.1f}x" if r['forward_pe'] > 0 else "N/A",
                f"${r['target_mean']:,.0f}" if r.get('target_mean', 0) > 0 else "N/A",
                f"{r['upside_pct']:.1f}%" if r.get('upside_pct') else "N/A",
            ])
        _tabla_datos_report(doc, headers_em, rows_em)

    # Sin datos
    if ("proyecciones_resultados" not in st.session_state or not st.session_state.proyecciones_resultados) and \
       ("emergentes_resultados" not in st.session_state or not st.session_state.emergentes_resultados):
        p_sin = doc.add_paragraph()
        run_sin = p_sin.add_run("No hay datos de análisis disponibles. Ejecuta el análisis en Important Companies primero.")
        run_sin.font.size = Pt(11)
        run_sin.font.italic = True
        run_sin.font.name = "Calibri"

    # Pie de página
    doc.add_paragraph("")
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = pie.add_run(f"Monitor de Opciones — Reporte Important Companies — {fecha_legible}")
    run_pie.font.size = Pt(8)
    run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run_pie.font.name = "Calibri"

    # Retornar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
