# -*- coding: utf-8 -*-
"""
Monitor de Opciones — Punto de entrada para Streamlit Cloud.
Orquesta la UI importando lógica desde config/, core/ y ui/.
"""
import calendar
import io
import json
import os
import logging
import time
import threading
import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
import yfinance as yf
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# --- Logging ---
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

# --- Importar módulos del proyecto ---
from config.constants import (
    DEFAULT_MIN_VOLUME, DEFAULT_MIN_OI, DEFAULT_MIN_PRIMA, DEFAULT_QUICK_FILTER,
    DEFAULT_TARGET_DELTA, AUTO_REFRESH_INTERVAL,
)
from config.watchlists import WATCHLIST_EMPRESAS, WATCHLIST_EMERGENTES

from core.scanner import (
    BROWSER_PROFILES, crear_sesion_nueva, obtener_historial_contrato,
    ejecutar_escaneo, cargar_historial_csv,
    obtener_precio_actual, limpiar_cache_ticker,
)
from core.projections import analizar_proyeccion_empresa
from core.range_calc import calcular_rango_esperado
from core.clusters import detectar_compras_continuas
from core.news import obtener_noticias_financieras, filtrar_noticias
from core.oi_tracker import calcular_cambios_oi, resumen_oi, filtrar_contratos_oi
from core.barchart_oi import obtener_top_oi_changes, obtener_oi_simbolo
from core.economic_calendar import obtener_eventos_economicos

from ui.styles import CSS_STYLES
from ui.components import (
    format_market_cap, render_empresa_card, render_tabla_comparativa,
    analizar_watchlist, render_watchlist_preview, render_empresa_descriptions,
    render_analisis_completo,
    render_metric_card, render_metric_row, render_plotly_sparkline,
    render_pro_table, _sentiment_badge, _type_badge, _priority_badge, _badge_html,
)


# ============================================================================
#                    THREADING - ESCANEO EN SEGUNDO PLANO
# ============================================================================
def _ejecutar_escaneo_thread(results_dict, ticker_symbol, umbral_vol, umbral_oi, 
                             umbral_prima, umbral_filtro, csv_carpeta, guardar_csv):
    """Ejecuta el escaneo en un thread separado y almacena resultados."""
    try:
        # Usar modo paralelo para mejor performance (paralelo=True)
        alertas, datos, error, perfil, fechas = ejecutar_escaneo(
            ticker_symbol,
            umbral_vol,
            umbral_oi,
            umbral_prima,
            umbral_filtro,
            csv_carpeta,
            guardar_csv,
            paralelo=True,  # ← Escaneo paralelo activado
        )
        
        results_dict['alertas'] = alertas
        results_dict['datos'] = datos
        results_dict['error'] = error
        results_dict['perfil'] = perfil
        results_dict['fechas'] = fechas
        results_dict['completed'] = True
        results_dict['ticker'] = ticker_symbol
    except Exception as e:
        results_dict['error'] = str(e)
        results_dict['completed'] = True


# ============================================================================
#                    HELPERS DE FORMATEO REUTILIZABLES
# ============================================================================
def _fmt_dolar(x):
    """Formatea un valor como moneda: $1,234."""
    return f"${x:,.0f}" if x > 0 else "$0"


def _fmt_iv(x):
    """Formatea IV como porcentaje: 25.3%."""
    return f"{x:.1f}%" if x > 0 else "-"


def _fmt_precio(x):
    """Formatea un precio: $1.23."""
    return f"${x:.2f}" if x > 0 else "-"


def _fmt_entero(x):
    """Formatea un entero con separadores: 1,234."""
    return f"{int(x):,}"


def _fmt_monto(v):
    """Formatea un monto grande: $1.2M, $50K, $1,234."""
    if v >= 1_000_000:
        return f"${v / 1_000_000:.1f}M"
    elif v >= 1_000:
        return f"${v / 1_000:.0f}K"
    return f"${v:,.0f}"


def _fmt_oi_chg(x):
    """Formatea OI Change con signo: +1,234 o -567."""
    return f"+{int(x):,}" if x > 0 else f"{int(x):,}" if x < 0 else "0"


def _fmt_lado(lado):
    """Formatea el lado de ejecución con emoji indicador."""
    if lado == "Ask":
        return "🟢 Ask"   # Compra agresiva
    elif lado == "Bid":
        return "🔴 Bid"   # Venta agresiva
    elif lado == "Mid":
        return "⚪ Mid"
    return "➖ N/A"


def determinar_sentimiento(tipo_opcion, lado):
    """
    Determina el sentimiento de la operación según el tipo de opción y lado de ejecución.
    
    Alcista (Verde) - Apuesta a que el precio suba:
    - CALL + Ask (compra de CALL)
    - PUT + Bid (venta de PUT)
    
    Bajista (Rojo) - Apuesta a que el precio baje:
    - PUT + Ask (compra de PUT)
    - CALL + Bid (venta de CALL)
    
    Returns:
        tuple: (sentimiento_texto, emoji, color_hex)
    """
    if tipo_opcion == "CALL" and lado == "Ask":
        return "ALCISTA", "🟢", "#10b981"
    elif tipo_opcion == "PUT" and lado == "Bid":
        return "ALCISTA", "🟢", "#10b981"
    elif tipo_opcion == "PUT" and lado == "Ask":
        return "BAJISTA", "🔴", "#ef4444"
    elif tipo_opcion == "CALL" and lado == "Bid":
        return "BAJISTA", "🔴", "#ef4444"
    else:
        return "NEUTRAL", "⚪", "#94a3b8"


# ============================================================================
#                    SISTEMA DE FAVORITOS (persistencia JSON)
# ============================================================================
_FAVORITOS_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "favoritos.json")


def _cargar_favoritos():
    """Carga favoritos desde archivo JSON. Purga contratos expirados."""
    try:
        if os.path.exists(_FAVORITOS_PATH):
            with open(_FAVORITOS_PATH, "r", encoding="utf-8") as f:
                favoritos = json.load(f)
            # Purgar contratos expirados
            hoy = datetime.now().strftime("%Y-%m-%d")
            favoritos = [fav for fav in favoritos if fav.get("Vencimiento", "9999-12-31") >= hoy]
            _guardar_favoritos(favoritos)  # persistir la limpieza
            return favoritos
    except Exception as e:
        logger.warning("Error cargando favoritos: %s", e)
    return []


def _guardar_favoritos(favoritos):
    """Guarda la lista de favoritos en archivo JSON."""
    try:
        os.makedirs(os.path.dirname(_FAVORITOS_PATH), exist_ok=True)
        with open(_FAVORITOS_PATH, "w", encoding="utf-8") as f:
            json.dump(favoritos, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error("Error guardando favoritos: %s", e)


def _agregar_favorito(contrato_data):
    """Agrega un contrato a favoritos si no existe ya."""
    favoritos = st.session_state.get("favoritos", [])
    contrato_id = contrato_data.get("Contrato", "")
    if not contrato_id:
        return False
    # Verificar que no exista ya
    if any(f.get("Contrato") == contrato_id for f in favoritos):
        return False
    # Agregar timestamp de cuando se marcó
    contrato_data["Guardado_En"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    favoritos.append(contrato_data)
    st.session_state.favoritos = favoritos
    _guardar_favoritos(favoritos)
    return True


def _eliminar_favorito(contrato_id):
    """Elimina un contrato de favoritos por su símbolo."""
    favoritos = st.session_state.get("favoritos", [])
    favoritos = [f for f in favoritos if f.get("Contrato") != contrato_id]
    st.session_state.favoritos = favoritos
    _guardar_favoritos(favoritos)


def _es_favorito(contrato_id):
    """Verifica si un contrato ya estí en favoritos."""
    favoritos = st.session_state.get("favoritos", [])
    return any(f.get("Contrato") == contrato_id for f in favoritos)


def _fetch_barchart_oi(simbolo, progress_bar=None):
    """Obtiene datos de OI de Barchart para un símbolo y actualiza session_state.
    
    Args:
        simbolo: Ticker del símbolo
        progress_bar: (opcional) Objeto st.progress() para actualizar durante el fetch
    """
    try:
        if progress_bar:
            progress_bar.progress(0.25, text="Cargando datos...")
        
        df_calls, err_c = obtener_oi_simbolo(simbolo, tipo="call")
        
        if progress_bar:
            progress_bar.progress(0.60, text="Cargando datos...")
        
        df_puts, err_p = obtener_oi_simbolo(simbolo, tipo="put")
        
        if progress_bar:
            progress_bar.progress(0.90, text="Cargando datos...")

        frames = []
        if df_calls is not None and not df_calls.empty:
            df_calls["Tipo"] = "CALL"
            frames.append(df_calls)
        if df_puts is not None and not df_puts.empty:
            df_puts["Tipo"] = "PUT"
            frames.append(df_puts)

        if frames:
            combined = pd.concat(frames, ignore_index=True)
            combined = combined.sort_values("OI_Chg", ascending=False).reset_index(drop=True)
            st.session_state.barchart_data = combined
            st.session_state.barchart_error = None
        else:
            err_msg = err_c or err_p or "Sin datos de Barchart"
            st.session_state.barchart_data = None
            st.session_state.barchart_error = err_msg
        
        if progress_bar:
            progress_bar.progress(1.0, text="Cargando datos...")
    except Exception as e:
        st.session_state.barchart_data = None
        st.session_state.barchart_error = f"Error Barchart: {e}"


def _inyectar_oi_chg_barchart():
    """Inyecta OI_Chg real de Barchart en datos_completos, alertas_actuales y clusters."""
    bc = st.session_state.get("barchart_data")
    if bc is None or bc.empty:
        return

    # Crear mapa (Vencimiento, Tipo, Strike) → OI_Chg de Barchart
    bc_map = {}
    for _, row in bc.iterrows():
        tipo = row.get("Tipo", "")
        strike = row.get("Strike", 0)
        venc = row.get("Vencimiento", "")
        oi_chg = int(row.get("OI_Chg", 0) or 0)
        key = (str(venc), tipo, float(strike))
        # Si hay duplicados, sumar o quedarnos con el de mayor magnitud
        if key not in bc_map or abs(oi_chg) > abs(bc_map[key]):
            bc_map[key] = oi_chg

    if not bc_map:
        return

    # Inyectar en datos_completos
    for d in st.session_state.datos_completos:
        key = (str(d.get("Vencimiento", "")), d.get("Tipo", ""), float(d.get("Strike", 0)))
        if key in bc_map:
            d["OI_Chg"] = bc_map[key]

    # Inyectar en alertas_actuales
    for a in st.session_state.alertas_actuales:
        key = (str(a.get("Vencimiento", "")), a.get("Tipo_Opcion", ""), float(a.get("Strike", 0)))
        if key in bc_map:
            a["OI_Chg"] = bc_map[key]

    # Inyectar en clusters (detalle)
    for c in st.session_state.clusters_detectados:
        total_chg = 0
        for det in c.get("Detalle", []):
            key = (str(det.get("Vencimiento", "")), det.get("Tipo_Opcion", c.get("Tipo_Opcion", "")), float(det.get("Strike", 0)))
            if key in bc_map:
                det["OI_Chg"] = bc_map[key]
                total_chg += bc_map[key]
        if total_chg != 0:
            c["OI_Chg_Total"] = total_chg


def _enriquecer_datos_opcion(datos, precio_subyacente=None):
    """Enriquece datos de opciones con mîtricas derivadas calculadas."""
    if not isinstance(datos, (list, pd.DataFrame)):
        return datos
    
    # Si es DataFrame, convertir a lista de dicts
    if isinstance(datos, pd.DataFrame):
        datos_lista = datos.to_dict('records')
    else:
        datos_lista = datos.copy()
    
    for item in datos_lista:
        try:
            # Bísicos
            ask = float(item.get('Ask', 0) or 0)
            bid = float(item.get('Bid', 0) or 0) 
            strike = float(item.get('Strike', 0) or 0)
            volumen = int(item.get('Volumen', 0) or 0)
            oi = int(item.get('OI', 0) or 0)
            iv = float(item.get('IV', 0) or 0)
            
            # Bid/Ask Spread
            if ask > 0 and bid > 0:
                spread = ask - bid
                spread_pct = (spread / ask) * 100 if ask > 0 else 0
                item['Spread'] = spread
                item['Spread_Pct'] = spread_pct
                item['Mid_Price'] = (ask + bid) / 2
            else:
                item['Spread'] = 0
                item['Spread_Pct'] = 0
                item['Mid_Price'] = ask or bid or 0
            
            # Volume/OI Ratio (liquidez relativa)
            item['Vol_OI_Ratio'] = volumen / oi if oi > 0 else 0
            
            # Liquidity Score (0-100)
            # Basado en: volumen, OI, spread estrechamente
            vol_score = min(volumen / 100, 1) * 40  # max 40 pts por volumen
            oi_score = min(oi / 500, 1) * 30        # max 30 pts por OI
            spread_score = max(0, 1 - item.get('Spread_Pct', 100)/10) * 30  # max 30 pts por spread estrecho
            item['Liquidity_Score'] = vol_score + oi_score + spread_score
            
            # Moneyness (si tenemos precio subyacente)
            if precio_subyacente and strike > 0:
                if item.get('Tipo_Opcion', '') == 'CALL' or item.get('Tipo', '') == 'CALL':
                    moneyness = strike / precio_subyacente
                    if moneyness < 0.95:
                        item['Moneyness'] = 'ITM'
                    elif moneyness > 1.05:
                        item['Moneyness'] = 'OTM'
                    else:
                        item['Moneyness'] = 'ATM'
                else:  # PUT
                    moneyness = precio_subyacente / strike
                    if moneyness < 0.95:
                        item['Moneyness'] = 'ITM' 
                    elif moneyness > 1.05:
                        item['Moneyness'] = 'OTM'
                    else:
                        item['Moneyness'] = 'ATM'
                
                # Distancia porcentual del precio actual
                item['Distance_Pct'] = abs(strike - precio_subyacente) / precio_subyacente * 100
            else:
                item['Moneyness'] = 'N/A'
                item['Distance_Pct'] = 0
            
            # Premium/Underlying Ratio
            mid_price = item.get('Mid_Price', 0)
            if precio_subyacente and mid_price > 0:
                item['Premium_Ratio'] = (mid_price / precio_subyacente) * 100
            else:
                item['Premium_Ratio'] = 0
            
            # Time Value (si no tenemos valor intrínseco exacto, aproximamos)
            if precio_subyacente and strike > 0 and mid_price > 0:
                tipo = item.get('Tipo_Opcion', item.get('Tipo', ''))
                if tipo == 'CALL':
                    intrinsic = max(precio_subyacente - strike, 0)
                else:  # PUT
                    intrinsic = max(strike - precio_subyacente, 0)
                item['Time_Value'] = max(mid_price - intrinsic, 0)
                item['Time_Value_Pct'] = (item['Time_Value'] / mid_price * 100) if mid_price > 0 else 0
            else:
                item['Time_Value'] = 0
                item['Time_Value_Pct'] = 0
                
        except (ValueError, TypeError, KeyError) as e:
            # Si hay error en algún cílculo, continuar con valores por defecto
            continue
    
    return datos_lista if not isinstance(datos, pd.DataFrame) else pd.DataFrame(datos_lista)


# ============================================================================
#                    CONFIGURACIÓN DE PÁGINA
# ============================================================================
st.set_page_config(
    page_title="OPTIONSKING Analytics",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ============================================================================
#                    ESTILOS CSS
# ============================================================================
# CSS complementario para reforzar el tema oscuro profesional
_CUSTOM_CSS = """
<style>
    /* Fondo general ultra oscuro */
    .stApp { background-color: #0f172a; color: white; }

    /* Sidebar custom */
    section[data-testid="stSidebar"] {
        background-color: #1e293b;
        border-right: 1px solid #334155;
    }

    /* Cards y containers */
    .metric-card, .stAlert, div.block-container {
        background-color: #1e293b !important;
        border-radius: 12px;
        padding: 1.5rem;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.5);
        border: 1px solid #334155;
    }

    /* Mîtricas superiores */
    .stMetric { background-color: #1e293b; border-radius: 12px; padding: 1rem; }
    .stMetric > label { color: #94a3b8; }
    .stMetric > div { color: white; font-size: 1.8rem; }

    /* Tablas pro */
    table { background-color: #1e293b; }
    thead tr { background-color: #0f172a !important; }
    tbody tr:hover { background-color: #334155 !important; }

    /* Verde neón y rojo */
    .positive { color: #00ff88; }
    .negative { color: #ef4444; }
    .badge-green { background-color: #10b981; color: white; padding: 4px 10px; border-radius: 8px; }
    .badge-red { background-color: #ef4444; color: white; padding: 4px 10px; border-radius: 8px; }

    /* Gauge y charts */
    .js-plotly-plot { background-color: #1e293b !important; }
</style>
"""

# Combina CUSTOM_CSS (base) + CSS_STYLES (avanzado — tiene !important, gana donde corresponda)
st.markdown(_CUSTOM_CSS + CSS_STYLES, unsafe_allow_html=True)

# Forzar dark mode en <html> para navegadores y componentes internos
st.markdown(
    '<script>document.documentElement.setAttribute("data-theme","dark");'
    'document.documentElement.style.colorScheme="dark";</script>'
    '<meta name="color-scheme" content="dark">',
    unsafe_allow_html=True,
)

# ============================================================================
#                    INICIALIZAR SESSION STATE
# ============================================================================
_DEFAULTS = {
    "alertas_actuales": [],
    "datos_completos": [],
    "scan_count": 0,
    "last_scan_time": None,
    "last_perfil": None,
    "scan_error": None,
    "datos_anteriores": [],
    "oi_cambios": None,
    "fechas_escaneadas": [],
    "auto_scan": False,
    "clusters_detectados": [],
    "ticker_anterior": "SPY",
    "trigger_scan": False,
    "todas_las_fechas": [],
    "rango_resultado": None,
    "rango_error": None,
    "scanning_active": False,
    "scan_thread": None,
    "scan_thread_results": None,
    "noticias_data": [],
    "noticias_last_refresh": None,
    "barchart_data": None,
    "barchart_error": None,
    "noticias_auto_refresh": False,
    "noticias_filtro": "Todas",
    "favoritos": [],
    "eventos_economicos": [],
    "eventos_last_refresh": None,
}
for _key, _val in _DEFAULTS.items():
    if _key not in st.session_state:
        st.session_state[_key] = _val

# Cargar favoritos desde disco al inicio
if not st.session_state.favoritos:
    st.session_state.favoritos = _cargar_favoritos()

# ============================================================================
#                    SIDEBAR - CONFIGURACIÓN
# ============================================================================
with st.sidebar:
    # -- OPTIONSKING Analytics Logo --
    st.markdown("""
        <div style="text-align: center; padding: 2rem 0 1rem;">
            <div style="display:inline-block;width:56px;height:56px;margin-bottom:10px;">
                <svg viewBox="0 0 64 64" fill="none" xmlns="http://www.w3.org/2000/svg">
                    <defs><linearGradient id="cg" x1="0" y1="0" x2="1" y2="1">
                        <stop offset="0%" stop-color="#00ff88"/><stop offset="100%" stop-color="#10b981"/>
                    </linearGradient></defs>
                    <path d="M8 48h48l-6-28-10 12-10-20-10 20-10-12z" fill="url(#cg)" stroke="#00ff88" stroke-width="1.5"/>
                    <rect x="8" y="48" width="48" height="6" rx="2" fill="url(#cg)"/>
                    <circle cx="32" cy="12" r="3" fill="#00ff88"/>
                    <circle cx="12" cy="22" r="2.5" fill="#10b981"/>
                    <circle cx="52" cy="22" r="2.5" fill="#10b981"/>
                </svg>
            </div>
            <h1 style="color: #00ff88; font-size: 36px; margin:0; font-weight:800; letter-spacing:-0.02em;">OPTIONSKING</h1>
            <p style="color: white; font-size: 22px; margin:4px 0 0 0; font-weight:500;">Analytics</p>
        </div>
        <hr style="border-color: #334155; margin: 0.5rem 0 1rem 0;">
    """, unsafe_allow_html=True)

# -- Menú de navegación con emojis (siempre disponible) --
    pagina = st.radio(
        "Navegación",
        ["🔍 Live Scanning", "📊 Open Interest", "📈 Data Analysis",
         "📐 Range", "⭐ Favorites", "🏢 Important Companies", "📰 News & Calendar", "📋 Reports"],
        index=0,
        label_visibility="collapsed",
    )

    st.markdown("---")

    # -- Avatar / User Section --
    st.markdown(
        '<div style="text-align:center; margin-top:2rem; padding:1rem 0;">'  
        '<div style="width:48px;height:48px;border-radius:50%;background:linear-gradient(135deg,#00ff88,#10b981);'
        'display:inline-flex;align-items:center;justify-content:center;font-size:20px;font-weight:700;color:#0f172a;'
        'margin-bottom:8px;box-shadow:0 0 16px rgba(0,255,136,0.2);">AD</div>'
        '<div style="color:white;font-weight:600;font-size:0.9rem;">Ariel David</div>'
        '<div style="color:#64748b;font-size:0.75rem;">● Pro Plan</div>'
        '</div>',
        unsafe_allow_html=True,
    )

# Valores por defecto de umbrales (se configuran en Live Scanning)
if "umbral_vol" not in st.session_state:
    st.session_state.umbral_vol = DEFAULT_MIN_VOLUME
if "umbral_oi" not in st.session_state:
    st.session_state.umbral_oi = DEFAULT_MIN_OI
if "umbral_prima" not in st.session_state:
    st.session_state.umbral_prima = DEFAULT_MIN_PRIMA
if "umbral_filtro" not in st.session_state:
    st.session_state.umbral_filtro = DEFAULT_QUICK_FILTER

# Guardado automítico siempre activo
csv_carpeta = "alertas"
guardar_csv = True

# rango_delta se configura en la pígina Range
if "rango_delta" not in st.session_state:
    st.session_state.rango_delta = DEFAULT_TARGET_DELTA
rango_delta = st.session_state.rango_delta

# Obtener pígina actual (navegación siempre disponible)
if "current_page" not in st.session_state:
    st.session_state.current_page = "🔍 Live Scanning"
st.session_state.current_page = pagina
pagina = st.session_state.current_page

# ============================================================================
#                    ENCABEZADO PRINCIPAL
# ============================================================================
# Header superior: ticker input
ticker_symbol = st.text_input(
    "🔍 Símbolo del Ticker", value="SPY", max_chars=10,
    help="Ingresa el símbolo de la acción (ej: SPY, AAPL, TSLA, QQQ)",
    placeholder="Escribe un ticker... (SPY, AAPL, TSLA, QQQ)",
    label_visibility="collapsed",
).strip().upper()

# Detectar cambio de ticker → auto-escanear
if ticker_symbol and ticker_symbol != st.session_state.ticker_anterior:
    st.session_state.ticker_anterior = ticker_symbol
    st.session_state.alertas_actuales = []
    st.session_state.datos_completos = []
    st.session_state.datos_anteriores = []
    st.session_state.oi_cambios = None
    st.session_state.barchart_data = None
    st.session_state.barchart_error = None
    st.session_state.clusters_detectados = []
    st.session_state.rango_resultado = None
    st.session_state.rango_error = None
    st.session_state.scan_error = None
    st.session_state.fechas_escaneadas = []
    limpiar_cache_ticker(ticker_symbol)  # Forzar datos frescos del nuevo ticker
    st.session_state.trigger_scan = True
    st.rerun()

st.markdown(
    f"""
    <div class="scanner-header">
        <h1>👑 OPTIONS<span style="color: #00ff88;">KING</span> Analytics</h1>
        <p class="subtitle">
            Escíner institucional de actividad inusual en opciones — <b style="color: #00ff88;">{ticker_symbol}</b>
        </p>
        <span class="badge">● LIVE • Anílisis Avanzado</span>
    </div>
    """,
    unsafe_allow_html=True,
)

# ============================================================================
#                    NAVEGACIÓN POR RADIO (SIDEBAR)
# ============================================================================
# Variables de umbrales (disponibles en todas las píginas)
umbral_vol = st.session_state.umbral_vol
umbral_oi = st.session_state.umbral_oi
umbral_prima = st.session_state.umbral_prima
umbral_filtro = st.session_state.umbral_filtro

# ============================================================================
#   🔍 LIVE SCANNING
# ============================================================================
if pagina == "🔍 Live Scanning":

    # --- Umbrales de filtrado ---
    with st.expander("📊 Umbrales de Filtrado", expanded=False):
        _umb_c1, _umb_c2, _umb_c3, _umb_c4 = st.columns(4)
        with _umb_c1:
            umbral_vol = st.number_input("Volumen mínimo", value=st.session_state.umbral_vol, step=1_000, format="%d",
                                          help="Solo muestra contratos con volumen ≥ este valor", key="inp_umbral_vol")
        with _umb_c2:
            umbral_oi = st.number_input("Open Interest mínimo", value=st.session_state.umbral_oi, step=1_000, format="%d",
                                         help="Solo muestra contratos con OI ≥ este valor", key="inp_umbral_oi")
        with _umb_c3:
            umbral_prima = st.number_input("Prima Total mínima ($)", value=st.session_state.umbral_prima, step=500_000, format="%d",
                                            help="Prima Total = Volumen × Precio × 100", key="inp_umbral_prima")
        with _umb_c4:
            umbral_filtro = st.number_input("Filtro rípido (vol/oi mín.)", value=st.session_state.umbral_filtro, step=100, format="%d",
                                             help="Ignora opciones con vol Y oi debajo de este umbral", key="inp_umbral_filtro")
        # Guardar en session_state para persistir entre píginas
        st.session_state.umbral_vol = umbral_vol
        st.session_state.umbral_oi = umbral_oi
        st.session_state.umbral_prima = umbral_prima
        st.session_state.umbral_filtro = umbral_filtro

    col_btn1, col_btn2 = st.columns([1, 1])

    with col_btn1:
        scan_btn = st.button("🚀 Escanear Ahora", type="primary", use_container_width=True)
    with col_btn2:
        auto_scan = st.checkbox("🔄 Auto-escaneo (5 min)")

    if st.session_state.last_scan_time:
        st.markdown(
            f"""
            <div class="status-bar">
                <div class="status-dot"></div>
                <span>Último escaneo: <b>{st.session_state.last_scan_time}</b></span>
                <span>Perfil TLS: <b>{st.session_state.last_perfil}</b></span>
                <span>Ciclos: <b>{st.session_state.scan_count}</b></span>
                <span>Fechas: <b>{len(st.session_state.fechas_escaneadas)}</b></span>
            </div>
            """,
            unsafe_allow_html=True,
        )

    # --- Ejecutar escaneo ---
    auto_trigger = st.session_state.trigger_scan
    if auto_trigger:
        st.session_state.trigger_scan = False

    # Verificar si hay un thread de escaneo activo
    if st.session_state.scan_thread is not None:
        thread = st.session_state.scan_thread
        results = st.session_state.scan_thread_results
        
        # Verificar si el thread ha terminado
        if not thread.is_alive() and results.get('completed', False):
            # Procesar resultados del thread
            alertas = results.get('alertas', [])
            datos = results.get('datos', [])
            error = results.get('error', None)
            perfil = results.get('perfil', None)
            fechas = results.get('fechas', [])
            
            if error:
                st.error(f"❌ Error en escaneo: {error}")
                st.session_state.scan_error = error
            else:
                st.session_state.alertas_actuales = alertas
                st.session_state.datos_completos = datos
                st.session_state.scan_count += 1
                st.session_state.last_scan_time = datetime.now().strftime("%H:%M:%S")
                
                # Capturar precio subyacente usando caché TTL (evita rate-limiting)
                precio, _err_precio = obtener_precio_actual(ticker_symbol)
                if precio is not None:
                    st.session_state.precio_subyacente = precio
                st.session_state.last_perfil = perfil
                st.session_state.scan_error = None
                st.session_state.fechas_escaneadas = fechas
                n_alertas = len(alertas)
                n_opciones = len(datos)

                # Calcular cambios en OI entre escaneos (para oi_tracker)
                if st.session_state.datos_anteriores:
                    st.session_state.oi_cambios = calcular_cambios_oi(
                        datos, st.session_state.datos_anteriores
                    )

                # Inicializar OI_Chg en 0 (será sobrescrito por Barchart)
                for d in st.session_state.datos_completos:
                    d["OI_Chg"] = 0
                for a in st.session_state.alertas_actuales:
                    a["OI_Chg"] = 0

                # Auto-fetch Barchart OI Changes (fuente real de OI_Chg)
                progress_bar = st.progress(0, text="Cargando datos...")
                _fetch_barchart_oi(ticker_symbol, progress_bar=progress_bar)
                progress_bar.empty()  # Limpiar barra al terminar

                # Inyectar OI_Chg real de Barchart en datos_completos y alertas
                _inyectar_oi_chg_barchart()

                # Detectar clusters DESPUÉS de inyectar OI_Chg
                clusters = detectar_compras_continuas(alertas, umbral_prima)
                st.session_state.clusters_detectados = clusters

                st.success(f"✅ Escaneo completado — {n_alertas} alertas en {n_opciones:,} opciones")
            
            # Limpiar thread y resultados
            st.session_state.scan_thread = None
            st.session_state.scan_thread_results = None
            st.session_state.scanning_active = False
            st.rerun()
        else:
            # Thread todavía está corriendo - mostrar indicador
            st.info(f"🔄 Escaneo en progreso para **{ticker_symbol}**... Puedes navegar libremente mientras tanto.")
            # Forzar rerun para actualizar estado
            time.sleep(0.5)
            st.rerun()

    # Iniciar nuevo escaneo si se presionó el botón
    if scan_btn or auto_trigger or (auto_scan and st.session_state.auto_scan):
        # Solo iniciar si no hay un thread activo
        if st.session_state.scan_thread is None or not st.session_state.scan_thread.is_alive():
            st.session_state.scanning_active = True
            
            # Guardar datos anteriores para comparar OI
            if st.session_state.datos_completos:
                st.session_state.datos_anteriores = st.session_state.datos_completos.copy()

            # Crear diccionario compartido para resultados
            results_dict = {}
            st.session_state.scan_thread_results = results_dict
            
            # Iniciar thread de escaneo
            thread = threading.Thread(
                target=_ejecutar_escaneo_thread,
                args=(results_dict, ticker_symbol, umbral_vol, umbral_oi, 
                      umbral_prima, umbral_filtro, csv_carpeta, guardar_csv),
                daemon=True
            )
            thread.start()
            st.session_state.scan_thread = thread
            
            st.info(f"🚀 Escaneo iniciado para **{ticker_symbol}**... Puedes continuar usando la aplicación.")
            time.sleep(0.5)
            st.rerun()

    st.session_state.auto_scan = auto_scan

    # --- Mîtricas rípidas ---
    if st.session_state.datos_completos:
        st.markdown("### 📊 Mîtricas del Escaneo")
        datos_df = pd.DataFrame(st.session_state.datos_completos)
        _n_calls = len(datos_df[datos_df["Tipo"] == "CALL"])
        _n_puts = len(datos_df[datos_df["Tipo"] == "PUT"])
        _n_alertas = len(st.session_state.alertas_actuales)
        _n_clusters = len(st.session_state.clusters_detectados)
        _total = len(datos_df)
        _call_pct = (_n_calls / _total * 100) if _total else 0
        _put_pct = (_n_puts / _total * 100) if _total else 0
        _pc_ratio = _n_puts / _n_calls if _n_calls > 0 else 0
        _total_vol = int(datos_df["Volumen"].sum()) if "Volumen" in datos_df.columns else 0
        _total_oi = int(datos_df["OI"].sum()) if "OI" in datos_df.columns else 0
        _total_prima = datos_df["Prima_Volumen"].sum() if "Prima_Volumen" in datos_df.columns else 0
        _flow_pct = _call_pct - _put_pct  # positive = bullish flow
        _spk = sorted(datos_df["Volumen"].dropna().tail(12).tolist()) if "Volumen" in datos_df.columns else None
        _spk_oi = sorted(datos_df["OI"].dropna().tail(12).tolist()) if "OI" in datos_df.columns else None

        st.markdown(render_metric_row([
            render_metric_card("Flow Sentiment", f"{_flow_pct:+.1f}%", delta=_flow_pct, sparkline_data=_spk),
            render_metric_card("Total Volume", f"{_total_vol:,}", delta=_call_pct, delta_suffix="% calls"),
            render_metric_card("Gamma Exposure", _fmt_monto(_total_prima), sparkline_data=_spk_oi),
            render_metric_card("Put/Call Ratio", f"{_pc_ratio:.2f}", delta=-(_pc_ratio - 1) * 100 if _pc_ratio != 0 else 0, color_override="#ef4444" if _pc_ratio > 1 else "#00ff88"),
            render_metric_card("Unusual Alerts", f"{_n_alertas}", delta=float(_n_clusters), delta_suffix=" clusters"),
        ]), unsafe_allow_html=True)

    # --- Mostrar alertas ---
    if st.session_state.alertas_actuales:
        st.markdown("### 🚨 Alertas Detectadas")

        st.markdown(
            """
            <div class="leyenda-colores">
                <div style="font-weight: 600; color: #f1f5f9; margin-bottom: 8px; font-size: 0.9rem;">🎨 Guía de Prioridades</div>
                <span class="leyenda-item"><span class="dot-green">●</span> <b>VERDE</b> — Mayor prima detectada. Míxima atención: contrato con mís dinero en juego.</span>
                <span class="leyenda-item"><span class="dot-red">●</span> <b>ROJO</b> — Actividad institucional. Vol <u>y</u> OI superan umbrales + prima alta.</span>
                <span class="leyenda-item"><span class="dot-orange">●</span> <b>NARANJA</b> — Actividad notable. Vol y OI superan umbrales.</span>
                <span class="leyenda-item"><span class="dot-purple">●</span> <b>MORADO</b> — Compra continua. Múltiples contratos similares cerca del umbral = posible mismo comprador institucional.</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

        alertas_sorted = sorted(
            st.session_state.alertas_actuales,
            key=lambda a: a["Prima_Volumen"],
            reverse=True,
        )
        max_prima = max(a["Prima_Volumen"] for a in alertas_sorted)

        for i, alerta in enumerate(alertas_sorted):
            tipo = alerta["Tipo_Alerta"]
            prima_mayor = alerta["Prima_Volumen"]

            es_top = (prima_mayor == max_prima) and (i == 0)
            if es_top:
                css_class = "alerta-top"
                emoji = "🟢"
                etiqueta = "MAYOR PRIMA"
            elif tipo == "PRINCIPAL":
                css_class = "alerta-principal"
                emoji = "🔴"
                etiqueta = "ACTIVIDAD INSTITUCIONAL"
            else:
                css_class = "alerta-prima"
                emoji = "🟠"
                etiqueta = "PRIMA ALTA"

            # Determinar sentimiento para colorear
            sentimiento_txt, sentimiento_emoji, sentimiento_color = determinar_sentimiento(
                alerta["Tipo_Opcion"], alerta.get("Lado", "N/A")
            )

            razones = []
            if alerta["Volumen"] >= umbral_vol:
                razones.append(f"Vol {alerta['Volumen']:,} ≥ {umbral_vol:,}")
            if alerta["OI"] >= umbral_oi:
                razones.append(f"OI {alerta['OI']:,} ≥ {umbral_oi:,}")
            if alerta["Prima_Volumen"] >= umbral_prima:
                razones.append(f"Prima Total ${alerta['Prima_Volumen']:,.0f} ≥ ${umbral_prima:,.0f}")
            if es_top:
                razones.insert(0, f"💰 Mayor prima del escaneo: ${prima_mayor:,.0f}")
            razon_html = " | ".join(razones)

            prima_vol_fmt = f"${alerta['Prima_Volumen']:,.0f}"
            contract_sym_card = alerta.get("Contrato", "")

            expander_label = (
                f"{emoji} {etiqueta} — {alerta['Tipo_Opcion']} Strike ${alerta['Strike']} | "
                f"Venc: {alerta['Vencimiento']} | Vol: {alerta['Volumen']:,} | "
                f"Prima: ${prima_mayor:,.0f}"
            )

            with st.expander(expander_label, expanded=False):
                # ⭐ Botón de favorito rípido (arriba del detalle)
                if contract_sym_card:
                    ya_fav_top = _es_favorito(contract_sym_card)
                    star_icon = "⭐" if ya_fav_top else "☆"
                    star_label = f"{star_icon} Favorito" if ya_fav_top else f"{star_icon} Marcar Favorito"
                    col_star, _ = st.columns([1, 4])
                    with col_star:
                        if st.button(star_label, key=f"star_top_{i}_{contract_sym_card}", disabled=ya_fav_top, use_container_width=True):
                            fav_data_top = {
                                "Contrato": contract_sym_card,
                                "Ticker": alerta.get("Ticker", ticker_symbol),
                                "Tipo_Opcion": alerta["Tipo_Opcion"],
                                "Strike": alerta["Strike"],
                                "Vencimiento": alerta["Vencimiento"],
                                "Volumen": alerta["Volumen"],
                                "OI": alerta["OI"],
                                "OI_Chg": alerta.get("OI_Chg", 0),
                                "Ask": alerta["Ask"],
                                "Bid": alerta["Bid"],
                                "Ultimo": alerta["Ultimo"],
                                "Lado": alerta.get("Lado", "N/A"),
                                "IV": alerta.get("IV", 0),
                                "Prima_Volumen": alerta["Prima_Volumen"],
                                "Tipo_Alerta": alerta["Tipo_Alerta"],
                            }
                            if _agregar_favorito(fav_data_top):
                                st.rerun()

                st.markdown(
                    f"""
                    <div class="{css_class}" style="margin-bottom: 0; border-left: 5px solid {sentimiento_color} !important;">
                        <div style="display: flex; justify-content: space-between; align-items: center;">
                            <div>
                                <strong>{emoji} {etiqueta}</strong> — 
                                <b>{alerta['Tipo_Opcion']}</b> | 
                                Strike: <b>${alerta['Strike']}</b> | 
                                Venc: <b>{alerta['Vencimiento']}</b>
                            </div>
                            <div style="padding: 4px 12px; border-radius: 8px; background: {sentimiento_color}20; border: 1px solid {sentimiento_color}; font-size: 0.75rem; font-weight: 700;">
                                {sentimiento_emoji} {sentimiento_txt}
                            </div>
                        </div>
                        Vol: <b>{alerta['Volumen']:,}</b> | 
                        Prima Total: <b>{prima_vol_fmt}</b> |
                        Ask: ${alerta['Ask']} | Bid: ${alerta['Bid']} | Último: ${alerta['Ultimo']} |
                        <b>Lado: {_fmt_lado(alerta.get('Lado', 'N/A'))}</b><br>
                        <span class="razon-alerta">📌 {razon_html}</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                # --- Detalles + Grífica del contrato ---
                if contract_sym_card:
                    col_chart, col_details = st.columns([3, 1])

                    with col_details:
                        st.markdown("**Detalles del contrato:**")
                        st.markdown(f"- **Símbolo:** `{contract_sym_card}`")
                        st.markdown(f"- **Tipo:** {alerta['Tipo_Opcion']}")
                        st.markdown(f"- **Strike:** ${alerta['Strike']}")
                        st.markdown(f"- **Vencimiento:** {alerta['Vencimiento']}")
                        st.markdown(f"- **Volumen:** {alerta['Volumen']:,}")
                        st.markdown(f"- **Ask:** ${alerta['Ask']}")
                        st.markdown(f"- **Bid:** ${alerta['Bid']}")
                        st.markdown(f"- **Último:** ${alerta['Ultimo']}")
                        st.markdown(f"- **Lado:** {_fmt_lado(alerta.get('Lado', 'N/A'))}")
                        st.markdown(f"- **Prima Total:** ${prima_mayor:,.0f}")

                        # Botón de favorito
                        ya_fav = _es_favorito(contract_sym_card)
                        btn_label = "⭐ Ya en Favoritos" if ya_fav else "☆ Guardar en Favoritos"
                        if st.button(btn_label, key=f"fav_btn_{i}_{contract_sym_card}", disabled=ya_fav, use_container_width=True):
                            fav_data = {
                                "Contrato": contract_sym_card,
                                "Ticker": alerta.get("Ticker", ticker_symbol),
                                "Tipo_Opcion": alerta["Tipo_Opcion"],
                                "Strike": alerta["Strike"],
                                "Vencimiento": alerta["Vencimiento"],
                                "Volumen": alerta["Volumen"],
                                "OI": alerta["OI"],
                                "OI_Chg": alerta.get("OI_Chg", 0),
                                "Ask": alerta["Ask"],
                                "Bid": alerta["Bid"],
                                "Ultimo": alerta["Ultimo"],
                                "Lado": alerta.get("Lado", "N/A"),
                                "IV": alerta.get("IV", 0),
                                "Prima_Volumen": alerta["Prima_Volumen"],
                                "Tipo_Alerta": alerta["Tipo_Alerta"],
                            }
                            if _agregar_favorito(fav_data):
                                st.success(f"⭐ {contract_sym_card} guardado en Favoritos")
                                st.rerun()

                    with col_chart:
                        with st.spinner("Cargando grífica..."):
                            hist_df_card, hist_err_card = obtener_historial_contrato(contract_sym_card)

                        if hist_err_card:
                            st.warning(f"⚠️ Error al cargar historial: {hist_err_card}")
                        elif hist_df_card.empty:
                            st.info("ℹ️ No hay datos históricos disponibles para este contrato.")
                        else:
                            st.markdown(f"**Precio del contrato** — `{contract_sym_card}`")
                            chart_price = hist_df_card[["Close"]].copy()
                            chart_price.columns = ["Precio"]
                            st.line_chart(chart_price, height=300)

                            if "Volume" in hist_df_card.columns:
                                chart_vol = hist_df_card[["Volume"]].copy()
                                chart_vol.columns = ["Volumen"]
                                st.bar_chart(chart_vol, height=180)

                            with st.expander("🗓️ Datos históricos completos"):
                                display_hist = hist_df_card.copy()
                                display_hist.index = display_hist.index.strftime("%Y-%m-%d %H:%M")
                                for col in ["Open", "High", "Low", "Close"]:
                                    if col in display_hist.columns:
                                        display_hist[col] = display_hist[col].apply(
                                            lambda x: f"${x:.2f}" if pd.notna(x) else "-"
                                        )
                                st.dataframe(display_hist, width="stretch", hide_index=False)
                else:
                    st.info("ℹ️ No se encontró el símbolo del contrato.")

        # ── Two-column dashboard layout ──────────────────────────────
        alertas_df = pd.DataFrame(alertas_sorted)

        def asignar_prioridad(row):
            prima_m = row["Prima_Volumen"]
            if prima_m == max_prima:
                return "TOP"
            elif row["Tipo_Alerta"] == "PRINCIPAL":
                return "INSTITUCIONAL"
            else:
                return "PRIMA ALTA"

        def _sent_badge_row(row):
            return _sentiment_badge(row["Tipo_Opcion"], row.get("Lado", "N/A"))

        alertas_df.insert(0, "Prioridad", alertas_df.apply(asignar_prioridad, axis=1))
        alertas_df.insert(1, "Sentimiento", alertas_df.apply(_sent_badge_row, axis=1))
        if "Tipo_Opcion" in alertas_df.columns:
            alertas_df["Tipo_Opcion"] = alertas_df["Tipo_Opcion"].apply(_type_badge)
        if "Lado" in alertas_df.columns:
            alertas_df["Lado"] = alertas_df["Lado"].apply(_fmt_lado)
        alertas_df = alertas_df.rename(columns={"Prima_Volumen": "Prima Total"})
        alertas_df["Prima Total"] = alertas_df["Prima Total"].apply(_fmt_dolar)
        cols_ocultar = [c for c in ["OI", "OI_Chg"] if c in alertas_df.columns]
        _tbl_df = alertas_df.drop(columns=cols_ocultar, errors="ignore")

        _col_left, _col_right = st.columns([1, 1], gap="medium")

        # ── LEFT COLUMN: Unusual Activity + Net Flow + Clusters ──
        with _col_left:
            st.markdown(
                render_pro_table(
                    _tbl_df,
                    title="📋 Unusual Activity — Alertas",
                    badge_count=f"{len(_tbl_df)} alertas",
                    footer_text=f"Ordenadas por prima · {len(_tbl_df)} resultados",
                    special_format={"Prioridad": _priority_badge},
                ),
                unsafe_allow_html=True,
            )

            # --- Net Flow bar chart (Calls vs Puts) ---
            _calls_prima = sum(
                d.get("Prima_Volumen", 0) for d in alertas_sorted
                if d.get("Tipo_Opcion") == "CALL"
            )
            _puts_prima = sum(
                d.get("Prima_Volumen", 0) for d in alertas_sorted
                if d.get("Tipo_Opcion") == "PUT"
            )
            if _calls_prima > 0 or _puts_prima > 0:
                _net_fig = go.Figure()
                _net_fig.add_trace(go.Bar(
                    x=["CALLS"], y=[_calls_prima],
                    marker_color="#00ff88", name="Calls",
                    text=[f"${_calls_prima:,.0f}"], textposition="auto",
                    textfont=dict(color="#ffffff", size=12),
                ))
                _net_fig.add_trace(go.Bar(
                    x=["PUTS"], y=[_puts_prima],
                    marker_color="#ef4444", name="Puts",
                    text=[f"${_puts_prima:,.0f}"], textposition="auto",
                    textfont=dict(color="#ffffff", size=12),
                ))
                _net_fig.update_layout(
                    title=dict(text="Net Premium Flow", font=dict(color="#e2e8f0", size=14)),
                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                    font=dict(color="#94a3b8", family="Inter"),
                    height=260, margin=dict(l=10, r=10, t=40, b=10),
                    showlegend=False,
                    yaxis=dict(gridcolor="rgba(51,65,85,0.4)", tickformat="$,.0f"),
                    xaxis=dict(showgrid=False),
                    bargap=0.35,
                )
                st.plotly_chart(_net_fig, use_container_width=True, config={"displayModeBar": False})

            # --- CLUSTERS ---
            if st.session_state.clusters_detectados:
                st.markdown("#### 🔗 Compras Continuas")
                st.markdown(
                    '<div style="background:rgba(139,92,246,0.06);border:1px solid rgba(139,92,246,0.15);'
                    'border-radius:12px;padding:10px 14px;margin-bottom:12px;font-size:0.78rem;color:#c4b5fd;">'
                    '⚠️ <b>Actividad institucional fragmentada</b> — Múltiples contratos similares con strikes '
                    'cercanos y primas cerca del umbral.</div>',
                    unsafe_allow_html=True,
                )

                for idx_c, cluster in enumerate(st.session_state.clusters_detectados):
                    rango_str = (
                        f"${cluster['Strike_Min']} - ${cluster['Strike_Max']}"
                        if cluster["Strike_Min"] != cluster["Strike_Max"]
                        else f"${cluster['Strike_Min']}"
                    )
                    st.markdown(
                        f'<div class="alerta-cluster">'
                        f'<strong>🟣 COMPRA CONTINUA</strong> '
                        f'<span class="cluster-badge">{cluster["Contratos"]} contratos</span><br>'
                        f'<b>{cluster["Tipo_Opcion"]}</b> | Venc: <b>{cluster["Vencimiento"]}</b> | '
                        f'Rango: <b>{rango_str}</b><br>'
                        f'Prima: <b>${cluster["Prima_Total"]:,.0f}</b> | '
                        f'Vol: <b>{cluster["Vol_Total"]:,}</b></div>',
                        unsafe_allow_html=True,
                    )

                if len(st.session_state.clusters_detectados) > 0:
                    clusters_table = []
                    for c in st.session_state.clusters_detectados:
                        clusters_table.append({
                            "Tipo": c["Tipo_Opcion"],
                            "Vencimiento": c["Vencimiento"],
                            "Contratos": c["Contratos"],
                            "Rango Strikes": f"${c['Strike_Min']} - ${c['Strike_Max']}",
                            "Prima Total": f"${c['Prima_Total']:,.0f}",
                            "Vol Total": f"{c['Vol_Total']:,}",
                        })
                    st.markdown(
                        render_pro_table(pd.DataFrame(clusters_table),
                                         title="🔗 Clusters Detectados",
                                         badge_count=f"{len(clusters_table)}"),
                        unsafe_allow_html=True,
                    )

        # ── RIGHT COLUMN: Options Flow Screener ──
        with _col_right:
            st.markdown(
                '<div style="font-size:1.05rem;font-weight:700;color:#e2e8f0;margin-bottom:8px;">'
                '🔍 Options Flow Screener</div>',
                unsafe_allow_html=True,
            )
            if st.session_state.datos_completos:
                datos_df = pd.DataFrame(st.session_state.datos_completos)

                _rf1, _rf2 = st.columns(2)
                with _rf1:
                    filtro_tipo = st.selectbox(
                        "Tipo", ["Todos", "CALL", "PUT"], key="filtro_tipo_scanner"
                    )
                with _rf2:
                    filtro_fecha = st.selectbox(
                        "Vencimiento",
                        ["Todos"] + sorted(datos_df["Vencimiento"].unique().tolist()),
                        key="filtro_fecha_scanner",
                    )
                min_vol_filtro = st.number_input(
                    "Volumen mínimo", value=0, step=100, key="min_vol_scanner"
                )

                df_filtered = datos_df.copy()
                if filtro_tipo != "Todos":
                    df_filtered = df_filtered[df_filtered["Tipo"] == filtro_tipo]
                if filtro_fecha != "Todos":
                    df_filtered = df_filtered[df_filtered["Vencimiento"] == filtro_fecha]
                if min_vol_filtro > 0:
                    df_filtered = df_filtered[df_filtered["Volumen"] >= min_vol_filtro]

                display_df = df_filtered.copy()
                if "Prima_Vol" in display_df.columns:
                    display_df = display_df.rename(columns={"Prima_Vol": "Prima Total"})
                    display_df["Prima Total"] = display_df["Prima Total"].apply(_fmt_dolar)
                display_df["IV"] = display_df["IV"].apply(_fmt_iv)

                cols_ocultar_df = [c for c in ["OI", "OI_Chg"] if c in display_df.columns]
                st.dataframe(
                    display_df.drop(columns=cols_ocultar_df, errors="ignore").sort_values("Volumen", ascending=False),
                    use_container_width=True,
                    hide_index=True,
                    height=600,
                )
                st.caption(f"Mostrando {len(df_filtered):,} de {len(datos_df):,} opciones")
            else:
                st.info("Ejecuta un escaneo para ver el flujo de opciones.")

    elif st.session_state.scan_count > 0 and not st.session_state.scan_error:
        st.success("✅ Sin alertas relevantes en este ciclo.")

    # --- Options Flow Screener (when no alerts but data exists) ---
    if not st.session_state.alertas_actuales and st.session_state.datos_completos:
        st.markdown(
            '<div style="font-size:1.05rem;font-weight:700;color:#e2e8f0;margin-bottom:8px;">'
            '🔍 Options Flow Screener</div>',
            unsafe_allow_html=True,
        )
        datos_df = pd.DataFrame(st.session_state.datos_completos)

        col_f1, col_f2, col_f3 = st.columns(3)
        with col_f1:
            filtro_tipo = st.selectbox(
                "Tipo", ["Todos", "CALL", "PUT"], key="filtro_tipo_scanner_noalert"
            )
        with col_f2:
            filtro_fecha = st.selectbox(
                "Vencimiento",
                ["Todos"] + sorted(datos_df["Vencimiento"].unique().tolist()),
                key="filtro_fecha_scanner_noalert",
            )
        with col_f3:
            min_vol_filtro = st.number_input(
                "Volumen mínimo", value=0, step=100, key="min_vol_scanner_noalert"
            )

        df_filtered = datos_df.copy()
        if filtro_tipo != "Todos":
            df_filtered = df_filtered[df_filtered["Tipo"] == filtro_tipo]
        if filtro_fecha != "Todos":
            df_filtered = df_filtered[df_filtered["Vencimiento"] == filtro_fecha]
        if min_vol_filtro > 0:
            df_filtered = df_filtered[df_filtered["Volumen"] >= min_vol_filtro]

        display_df = df_filtered.copy()
        if "Prima_Vol" in display_df.columns:
            display_df = display_df.rename(columns={"Prima_Vol": "Prima Total"})
            display_df["Prima Total"] = display_df["Prima Total"].apply(_fmt_dolar)
        display_df["IV"] = display_df["IV"].apply(_fmt_iv)

        cols_ocultar_df = [c for c in ["OI", "OI_Chg"] if c in display_df.columns]
        st.dataframe(
            display_df.drop(columns=cols_ocultar_df, errors="ignore").sort_values("Volumen", ascending=False),
            use_container_width=True,
            hide_index=True,
            height=500,
        )
        st.caption(f"Mostrando {len(df_filtered):,} de {len(datos_df):,} opciones")

    # --- Datos del Último Escaneo ---
    if st.session_state.datos_completos:
        st.markdown("---")
        st.markdown("#### 📊 Datos del Último Escaneo")
        datos_df_esc = pd.DataFrame(st.session_state.datos_completos)
        _a_calls = len(datos_df_esc[datos_df_esc["Tipo"] == "CALL"])
        _a_puts = len(datos_df_esc[datos_df_esc["Tipo"] == "PUT"])
        _a_total = len(datos_df_esc)
        _a_alertas = len(st.session_state.alertas_actuales)
        _a_clusters = len(st.session_state.clusters_detectados)
        _a_cpct = (_a_calls / _a_total * 100) if _a_total else 0
        _a_ppct = (_a_puts / _a_total * 100) if _a_total else 0
        _a_spk = sorted(datos_df_esc["Volumen"].dropna().tail(12).tolist()) if "Volumen" in datos_df_esc.columns else None
        st.markdown(render_metric_row([
            render_metric_card("Opciones", f"{_a_total:,}", sparkline_data=_a_spk),
            render_metric_card("Calls", f"{_a_calls:,}", delta=_a_cpct),
            render_metric_card("Puts", f"{_a_puts:,}", delta=_a_ppct, color_override="#ef4444"),
            render_metric_card("Alertas", f"{_a_alertas}"),
            render_metric_card("Clusters", f"{_a_clusters}"),
        ]), unsafe_allow_html=True)

        with st.expander("🔍 Ver todas las opciones escaneadas", expanded=False):
            datos_enriquecidos = _enriquecer_datos_opcion(
                st.session_state.datos_completos,
                precio_subyacente=st.session_state.get('precio_subyacente')
            )
            display_scan = pd.DataFrame(datos_enriquecidos)

            if 'Prima_Vol' in display_scan.columns:
                display_scan["Prima Total"] = display_scan["Prima_Vol"].apply(_fmt_monto)
            if 'IV' in display_scan.columns:
                display_scan["IV_F"] = display_scan["IV"].apply(_fmt_iv)
            if 'Spread_Pct' in display_scan.columns:
                display_scan["Spread_%"] = display_scan["Spread_Pct"].apply(lambda x: f"{x:.1f}%" if pd.notna(x) else "-")
            if 'Liquidity_Score' in display_scan.columns:
                display_scan["Liquidez"] = display_scan["Liquidity_Score"].apply(lambda x: f"{x:.0f}" if pd.notna(x) else "-")
            if 'Lado' in display_scan.columns:
                display_scan["Lado_F"] = display_scan["Lado"].apply(_fmt_lado)

            if 'Tipo' in display_scan.columns and 'Lado' in display_scan.columns:
                display_scan["Sentimiento"] = display_scan.apply(
                    lambda row: f"{determinar_sentimiento(row['Tipo'], row.get('Lado', 'N/A'))[1]} {determinar_sentimiento(row['Tipo'], row.get('Lado', 'N/A'))[0]}",
                    axis=1
                )

            cols_mostrar = ['Sentimiento', 'Tipo', 'Strike', 'Vencimiento', 'Volumen', 'Ask', 'Bid', 'Spread_%',
                           'Ultimo', 'Lado_F', 'IV_F', 'Moneyness', 'Prima Total', 'Liquidez']
            cols_disponibles = [c for c in cols_mostrar if c in display_scan.columns]
            cols_ocultar_h = [c for c in ["OI", "OI_Chg"] if c in display_scan.columns]
            display_final = display_scan.drop(columns=cols_ocultar_h, errors="ignore")

            st.dataframe(
                display_final[cols_disponibles] if cols_disponibles else display_final,
                use_container_width=True, hide_index=True, height=400,
            )

            csv_enriquecido = pd.DataFrame(datos_enriquecidos).to_csv(index=False).encode("utf-8")
            st.download_button(
                "📈 Descargar Datos Enriquecidos (CSV)",
                csv_enriquecido,
                f"opciones_enriquecidas_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                "text/csv",
                key="dl_datos_enriquecidos_escaneo",
                help="Incluye mîtricas adicionales: spread, moneyness, liquidez, ratios, etc."
            )

        # --- Clusters de Compra Continua ---
        if st.session_state.clusters_detectados:
            st.markdown("##### 🔗 Clusters de Compra Continua")
            clusters_table_esc = []
            for c in st.session_state.clusters_detectados:
                clusters_table_esc.append({
                    "Tipo": c["Tipo_Opcion"],
                    "Vencimiento": c["Vencimiento"],
                    "Contratos": c["Contratos"],
                    "Rango Strikes": f"${c['Strike_Min']} - ${c['Strike_Max']}",
                    "Prima Total": _fmt_monto(c['Prima_Total']),
                    "Prima Prom.": _fmt_monto(c['Prima_Promedio']),
                    "Vol Total": _fmt_entero(c['Vol_Total']),
                    "OI Total": _fmt_entero(c['OI_Total']),
                    "OI Chg": _fmt_oi_chg(c.get('OI_Chg_Total', 0)),
                })
            st.markdown(
                render_pro_table(pd.DataFrame(clusters_table_esc),
                                 title="🔗 Clusters de Compra Continua",
                                 badge_count=f"{len(clusters_table_esc)}"),
                unsafe_allow_html=True,
            )

    # Auto-refresh con countdown visual
    if auto_scan and st.session_state.scan_count > 0:
        countdown = AUTO_REFRESH_INTERVAL  # Configurable desde constants.py
        placeholder = st.empty()
        progress_bar = st.progress(1.0)
        for remaining in range(countdown, 0, -1):
            mins, secs = divmod(remaining, 60)
            pct = remaining / countdown
            placeholder.markdown(
                f'<div style="background:#1e293b;border:1px solid #334155;border-radius:10px;'
                f'padding:10px 18px;display:flex;align-items:center;gap:12px;font-size:0.85rem;">'
                f'<span style="color:#00ff88;font-size:1.1rem;">🔄</span>'
                f'<span style="color:#94a3b8;">Próximo escaneo en</span>'
                f'<span style="color:#ffffff;font-weight:700;font-family:JetBrains Mono,monospace;">'
                f'{mins}:{secs:02d}</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
            progress_bar.progress(pct)
            time.sleep(1)
        placeholder.empty()
        progress_bar.empty()
        st.rerun()


# ============================================================================
#   📊 OPEN INTEREST
# ============================================================================
elif pagina == "📊 Open Interest":
    st.markdown("### 📊 Open Interest")

    # ================================================================
    #  TOP OI CHANGES (Barchart) — Auto-cargado al escanear
    # ================================================================
    st.markdown("#### 🔥 Top Cambios en OI — Barchart")
    st.caption("Se actualiza automíticamente con cada escaneo • Fuente: Barchart.com")

    # Filtro tipo + OI Chg mínimo
    col_f1, col_f2 = st.columns([1, 1])
    with col_f1:
        bc_tipo_filtro = st.radio(
            "Filtrar por tipo", ["Todos", "📞 CALL", "📋 PUT"],
            horizontal=True, key="bc_tipo_filtro", index=0,
        )
    with col_f2:
        bc_min_chg = st.number_input(
            "OI Chg mínimo", value=0, step=5, min_value=0, key="bc_min_chg",
        )

    # Botón para recarga manual (sin necesidad de re-escanear)
    col_btn1, col_btn2 = st.columns([1, 3])
    with col_btn1:
        bc_refresh = st.button("🔄 Actualizar OI", key="bc_refresh")

    if bc_refresh:
        sim_bc = st.session_state.get("ticker_anterior", "SPY")
        progress_bar = st.progress(0, text="Cargando datos...")
        _fetch_barchart_oi(sim_bc, progress_bar=progress_bar)
        _inyectar_oi_chg_barchart()
        progress_bar.empty()

    # Mostrar error
    if st.session_state.barchart_error:
        st.warning(f"⚠️ {st.session_state.barchart_error}")

    # Mostrar datos
    if st.session_state.barchart_data is not None and not st.session_state.barchart_data.empty:
        df_bc_all = st.session_state.barchart_data.copy()

        # Aplicar filtro tipo
        if bc_tipo_filtro == "📞 CALL":
            df_bc_all = df_bc_all[df_bc_all["Tipo"] == "CALL"]
        elif bc_tipo_filtro == "📋 PUT":
            df_bc_all = df_bc_all[df_bc_all["Tipo"] == "PUT"]

        # Aplicar filtro OI Chg mínimo (valor absoluto)
        if bc_min_chg > 0:
            df_bc_all = df_bc_all[df_bc_all["OI_Chg"].abs() >= bc_min_chg]

        n_total = len(df_bc_all)

        if n_total == 0:
            st.info("Sin contratos que cumplan los filtros seleccionados.")
        else:
            # Separar positivos y negativos
            df_positivos = df_bc_all[df_bc_all["OI_Chg"] > 0].sort_values("OI_Chg", ascending=False).reset_index(drop=True)
            df_negativos = df_bc_all[df_bc_all["OI_Chg"] < 0].sort_values("OI_Chg", ascending=True).reset_index(drop=True)

            n_pos = len(df_positivos)
            n_neg = len(df_negativos)
            n_calls = len(df_bc_all[df_bc_all["Tipo"] == "CALL"]) if "Tipo" in df_bc_all.columns else 0
            n_puts = len(df_bc_all[df_bc_all["Tipo"] == "PUT"]) if "Tipo" in df_bc_all.columns else 0

            # Calcular contratos cerrados (OI_Chg negativo)
            contratos_cerrados_total = int(df_negativos["OI_Chg"].sum()) if n_neg > 0 else 0
            calls_cerrados = int(df_negativos[df_negativos["Tipo"] == "CALL"]["OI_Chg"].sum()) if n_neg > 0 and "Tipo" in df_negativos.columns else 0
            puts_cerrados = int(df_negativos[df_negativos["Tipo"] == "PUT"]["OI_Chg"].sum()) if n_neg > 0 and "Tipo" in df_negativos.columns else 0

            # Calcular contratos abiertos (OI_Chg positivo)
            contratos_abiertos_total = int(df_positivos["OI_Chg"].sum()) if n_pos > 0 else 0
            calls_abiertos = int(df_positivos[df_positivos["Tipo"] == "CALL"]["OI_Chg"].sum()) if n_pos > 0 and "Tipo" in df_positivos.columns else 0
            puts_abiertos = int(df_positivos[df_positivos["Tipo"] == "PUT"]["OI_Chg"].sum()) if n_pos > 0 and "Tipo" in df_positivos.columns else 0

            # Mîtricas rípidas
            _pos_pct = (n_pos / n_total * 100) if n_total else 0
            _neg_pct = (n_neg / n_total * 100) if n_total else 0
            st.markdown(render_metric_row([
                render_metric_card("Total Contratos", f"{n_total:,}"),
                render_metric_card("CALLs", f"{n_calls:,}", delta=(n_calls / n_total * 100) if n_total else 0),
                render_metric_card("PUTs", f"{n_puts:,}", delta=(n_puts / n_total * 100) if n_total else 0, color_override="#ef4444"),
                render_metric_card("Señales Positivas", f"{n_pos:,}", delta=_pos_pct),
                render_metric_card("Señales Negativas", f"{n_neg:,}", delta=_neg_pct, color_override="#ef4444"),
            ]), unsafe_allow_html=True)

            # Segunda fila de mîtricas: Contratos abiertos vs cerrados
            st.markdown("---")
            st.markdown("##### 📈 Flujo de Contratos")
            _net_flow = contratos_abiertos_total + contratos_cerrados_total
            _open_spk = [max(0, v) for v in df_positivos["OI_Chg"].head(10).tolist()] if n_pos > 1 else None
            _close_spk = [abs(v) for v in df_negativos["OI_Chg"].head(10).tolist()] if n_neg > 1 else None
            st.markdown(render_metric_row([
                render_metric_card("Contratos Abiertos", f"{contratos_abiertos_total:,}", delta="Nuevas posiciones", sparkline_data=_open_spk),
                render_metric_card("CALLs Abiertos", f"{calls_abiertos:,}"),
                render_metric_card("PUTs Abiertos", f"{puts_abiertos:,}"),
                render_metric_card("Contratos Cerrados", f"{contratos_cerrados_total:,}", delta="Posiciones cerradas", sparkline_data=_close_spk, color_override="#ef4444"),
                render_metric_card("CALLs Cerrados", f"{calls_cerrados:,}"),
                render_metric_card("PUTs Cerrados", f"{puts_cerrados:,}"),
            ]), unsafe_allow_html=True)

            st.markdown("---")

            # --- Columnas de tabla ---
            display_cols = ["Tipo", "Ticker", "Strike", "Vencimiento", "DTE",
                            "Volumen", "OI", "OI_Chg", "IV", "Delta", "Último"]

            def _formatear_tabla_oi(df_raw):
                """Formatea un DataFrame de OI para mostrar."""
                cols = [c for c in display_cols if c in df_raw.columns]
                df_fmt = df_raw[cols].copy()
                df_fmt["OI_Chg"] = df_fmt["OI_Chg"].apply(
                    lambda x: f"+{int(x):,}" if x > 0 else f"{int(x):,}" if x < 0 else "0"
                )
                df_fmt["Volumen"] = df_fmt["Volumen"].apply(lambda x: f"{int(x):,}")
                df_fmt["OI"] = df_fmt["OI"].apply(lambda x: f"{int(x):,}")
                df_fmt["IV"] = df_fmt["IV"].apply(lambda x: f"{x:.1f}%" if x > 0 else "-")
                df_fmt["Delta"] = df_fmt["Delta"].apply(lambda x: f"{x:.3f}" if x != 0 else "-")
                df_fmt["Último"] = df_fmt["Último"].apply(lambda x: f"${x:.2f}" if x > 0 else "-")
                df_fmt["Strike"] = df_fmt["Strike"].apply(lambda x: f"${x:,.1f}")
                return df_fmt

            def _mostrar_tabla_paginada(df_raw, df_fmt, key_prefix, emoji_func):
                """Muestra tabla con paginación y emojis."""
                n = len(df_fmt)
                if n == 0:
                    st.info("Sin contratos en esta categoría.")
                    return

                # Indicador visual
                df_show = df_fmt.copy()
                df_show.insert(0, "", df_raw["OI_Chg"].apply(emoji_func))

                contratos_por_grupo = 20
                if n > contratos_por_grupo:
                    rangos = []
                    for i in range(0, n, contratos_por_grupo):
                        inicio_r = i + 1
                        fin_r = min(i + contratos_por_grupo, n)
                        rangos.append(f"{inicio_r}-{fin_r}")
                    rango_sel = st.selectbox(
                        f"Rango de contratos (Total: {n:,})",
                        rangos, key=f"{key_prefix}_rango",
                    )
                    inicio_idx, fin_idx = map(int, rango_sel.split("-"))
                else:
                    inicio_idx, fin_idx = 1, n

                df_pagina = df_show.iloc[inicio_idx - 1 : fin_idx]
                st.dataframe(
                    df_pagina,
                    use_container_width=True,
                    hide_index=True,
                    height=min(500, 35 * len(df_pagina) + 38),
                )
                st.caption(f"Mostrando {inicio_idx}-{fin_idx} de {n:,} contratos")

            # ========================================
            # TABLA 1: OI Chg POSITIVO (Abriendo posiciones)
            # ========================================
            st.markdown("#### 🟢 OI Chg Positivo — Abriendo Posiciones")
            st.caption("Contratos donde el Open Interest aumentó → nuevas posiciones abiertas")

            if n_pos > 0:
                df_pos_fmt = _formatear_tabla_oi(df_positivos)
                _mostrar_tabla_paginada(
                    df_positivos, df_pos_fmt, "oi_pos",
                    lambda x: "🔥" if x >= 50 else ("🟢" if x >= 20 else "")
                )
            else:
                st.info("Sin contratos con OI Chg positivo.")

            st.markdown("---")

            # ========================================
            # TABLA 2: OI Chg NEGATIVO (Cerrando posiciones)
            # ========================================
            st.markdown("#### 🔴 OI Chg Negativo — Cerrando Posiciones")
            st.caption("Contratos donde el Open Interest disminuyó → posiciones cerradas o ejercidas")

            if n_neg > 0:
                df_neg_fmt = _formatear_tabla_oi(df_negativos)
                _mostrar_tabla_paginada(
                    df_negativos, df_neg_fmt, "oi_neg",
                    lambda x: "🔥" if x <= -50 else ("🔴" if x <= -20 else "")
                )
            else:
                st.info("Sin contratos con OI Chg negativo.")
    elif st.session_state.scan_count == 0:
        st.info("⏳ **Ejecutí un escaneo** en 🔍 Live Scanning para cargar los datos de Open Interest automíticamente.")


# ============================================================================
#   📋 REPORTS
# ============================================================================
elif pagina == "📋 Reports":
    st.markdown("### 📋 Reports")
    st.markdown(
        """
        <div class="watchlist-info">
            💾 <b>Centro de Reportes</b> — Descarga reportes detallados en formato DOCX.
            Los reportes se generan con los datos cargados en cada sección.
        </div>
        """,
        unsafe_allow_html=True,
    )

    # =============================================
    # HELPERS PARA GENERAR REPORTES DOCX
    # =============================================
    def _estilo_celda_report(cell, texto, negrita=False, color_fondo=None, color_texto=None, size=9, align="left"):
        """Aplica formato a una celda de tabla Word."""
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        p.space_before = Pt(1)
        p.space_after = Pt(1)
        run = p.add_run(str(texto))
        run.bold = negrita
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        if color_texto:
            run.font.color.rgb = color_texto
        if color_fondo:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = tcPr.makeelement(qn("w:shd"), {
                qn("w:fill"): color_fondo,
                qn("w:val"): "clear",
            })
            tcPr.append(shading)

    def _agregar_titulo_report(doc, texto, level=2):
        """Agrega un título de sección con formato."""
        heading = doc.add_heading(texto, level=level)
        for run in heading.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    def _tabla_info_report(doc, datos_dict, titulo=None):
        """Tabla de 2 columnas Campo/Valor para info resumida."""
        if titulo:
            p = doc.add_paragraph()
            run = p.add_run(titulo)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light List Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for campo, valor in datos_dict.items():
            row = table.add_row()
            _estilo_celda_report(row.cells[0], campo, negrita=True, size=10)
            _estilo_celda_report(row.cells[1], str(valor), size=10)
        doc.add_paragraph("")

    def _tabla_datos_report(doc, headers, rows_data):
        """Tabla con encabezados y filas de datos."""
        if not rows_data:
            return
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light List Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Encabezados
        for i, h in enumerate(headers):
            _estilo_celda_report(
                table.rows[0].cells[i], h,
                negrita=True, size=9,
                color_fondo="1E3A5F",
                color_texto=RGBColor(0xFF, 0xFF, 0xFF),
                align="center",
            )
        # Datos
        for row_idx, row_data in enumerate(rows_data):
            row = table.add_row()
            bg = "F0F4F8" if row_idx % 2 == 0 else None
            for i, val in enumerate(row_data):
                _estilo_celda_report(row.cells[i], str(val), size=9, color_fondo=bg)
        doc.add_paragraph("")

    # =============================================
    # FUNCIÓN 1: REPORTE LIVE SCANNING
    # =============================================
    def _generar_reporte_live_scanning():
        """Genera reporte DOCX con todos los datos del Live Scanning."""
        doc = Document()
        
        # Configurar pígina
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # Portada
        doc.add_paragraph("")
        titulo = doc.add_heading("REPORTE — LIVE SCANNING", level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        ticker_name = st.session_state.get("ticker_anterior", "N/A")
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = subtitulo.add_run(f"Ticker: {ticker_name}")
        run_sub.font.size = Pt(18)
        run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
        run_sub.font.name = "Calibri"
        run_sub.bold = True

        fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        fecha_p = doc.add_paragraph()
        fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
        run_fecha.font.size = Pt(11)
        run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run_fecha.font.name = "Calibri"

        doc.add_paragraph("")

        # Resumen Ejecutivo
        n_opciones = len(st.session_state.datos_completos)
        n_alertas = len(st.session_state.alertas_actuales)
        n_clusters = len(st.session_state.clusters_detectados)
        precio_subyacente = st.session_state.get('precio_subyacente', 0)
        
        datos_calls = [d for d in st.session_state.datos_completos if d.get('Tipo', '') == 'CALL']
        datos_puts = [d for d in st.session_state.datos_completos if d.get('Tipo', '') == 'PUT']
        
        vol_promedio = np.mean([d.get('Volumen', 0) for d in st.session_state.datos_completos]) if st.session_state.datos_completos else 0
        oi_promedio = np.mean([d.get('OI', 0) for d in st.session_state.datos_completos]) if st.session_state.datos_completos else 0
        iv_promedio = np.mean([d.get('IV', 0) for d in st.session_state.datos_completos if d.get('IV', 0) > 0]) if st.session_state.datos_completos else 0
        
        principales = [a for a in st.session_state.alertas_actuales if a.get("Tipo_Alerta") == "PRINCIPAL"]
        prima_alta = [a for a in st.session_state.alertas_actuales if a.get("Tipo_Alerta") == "PRIMA_ALTA"]

        _agregar_titulo_report(doc, "RESUMEN EJECUTIVO", level=1)
        _tabla_info_report(doc, {
            "Ticker Analizado": ticker_name,
            "Precio Subyacente": f"${precio_subyacente:,.2f}" if precio_subyacente > 0 else "N/D",
            "Fecha del Reporte": fecha_legible,
            "Total Opciones Escaneadas": f"{n_opciones:,}",
            "Calls vs Puts": f"{len(datos_calls):,} calls / {len(datos_puts):,} puts",
            "Alertas Detectadas": f"{n_alertas} ({len(principales)} institucionales, {len(prima_alta)} prima alta)",
            "Clusters de Compra": f"{n_clusters}",
            "Volumen Promedio": f"{vol_promedio:,.0f}",
            "OI Promedio": f"{oi_promedio:,.0f}",
            "IV Promedio": f"{iv_promedio:.1f}%" if iv_promedio > 0 else "N/D",
        })

        # Alertas Institucionales
        if principales:
            _agregar_titulo_report(doc, f"ALERTAS INSTITUCIONALES ({len(principales)})", level=1)
            p_desc = doc.add_paragraph()
            run_d = p_desc.add_run(
                "Operaciones con prima significativa que sugieren actividad institucional."
            )
            run_d.font.size = Pt(10)
            run_d.font.italic = True
            run_d.font.name = "Calibri"

            headers = ["#", "Tipo", "Strike", "Vencimiento", "Volumen", "OI", 
                       "Ask", "Bid", "Último", "IV", "Sentimiento", "Lado", "Prima Total", "Contrato"]
            rows = []
            principales_enriq = _enriquecer_datos_opcion(principales, precio_subyacente)
            
            for i, a in enumerate(principales_enriq, 1):
                sent_txt, sent_emoji, _ = determinar_sentimiento(a["Tipo_Opcion"], a.get("Lado", "N/A"))
                rows.append([
                    i, a["Tipo_Opcion"], _fmt_precio(a['Strike']), a["Vencimiento"],
                    _fmt_entero(a['Volumen']), _fmt_entero(a['OI']),
                    _fmt_precio(a['Ask']), _fmt_precio(a['Bid']), _fmt_precio(a['Ultimo']),
                    _fmt_iv(a.get('IV', 0)),
                    f"{sent_emoji} {sent_txt}", _fmt_lado(a.get('Lado', 'N/A')),
                    _fmt_monto(a['Prima_Volumen']),
                    a.get("Contrato", "N/A"),
                ])
            _tabla_datos_report(doc, headers, rows)

        # Alertas Prima Alta
        if prima_alta:
            _agregar_titulo_report(doc, f"ALERTAS PRIMA ALTA ({len(prima_alta)})", level=1)
            p_desc = doc.add_paragraph()
            run_d = p_desc.add_run(
                "Opciones con volumen y open interest por encima de los umbrales configurados."
            )
            run_d.font.size = Pt(10)
            run_d.font.italic = True
            run_d.font.name = "Calibri"

            headers = ["#", "Tipo", "Strike", "Vencimiento", "Volumen", "OI",
                       "Ask", "Bid", "Último", "IV", "Sentimiento", "Lado", "Prima Total"]
            rows = []
            prima_alta_enriq = _enriquecer_datos_opcion(prima_alta, precio_subyacente)
            
            for i, a in enumerate(prima_alta_enriq, 1):
                sent_txt, sent_emoji, _ = determinar_sentimiento(a["Tipo_Opcion"], a.get("Lado", "N/A"))
                rows.append([
                    i, a["Tipo_Opcion"], _fmt_precio(a['Strike']), a["Vencimiento"],
                    _fmt_entero(a['Volumen']), _fmt_entero(a['OI']),
                    _fmt_precio(a['Ask']), _fmt_precio(a['Bid']), _fmt_precio(a['Ultimo']),
                    _fmt_iv(a.get('IV', 0)),
                    f"{sent_emoji} {sent_txt}", _fmt_lado(a.get('Lado', 'N/A')),
                    _fmt_monto(a['Prima_Volumen']),
                ])
            _tabla_datos_report(doc, headers, rows)

        # Clusters
        if st.session_state.clusters_detectados:
            _agregar_titulo_report(doc, f"CLUSTERS DE COMPRA CONTINUA ({n_clusters})", level=1)
            p_desc = doc.add_paragraph()
            run_d = p_desc.add_run(
                "Grupos de contratos con strikes cercanos y primas similares en la misma expiración."
            )
            run_d.font.size = Pt(10)
            run_d.font.italic = True
            run_d.font.name = "Calibri"

            headers_cl = ["#", "Tipo", "Vencimiento", "Contratos", "Rango Strikes",
                          "Prima Total", "Prima Prom.", "Vol Total", "OI Total"]
            rows_cl = []
            for i, c in enumerate(st.session_state.clusters_detectados, 1):
                rows_cl.append([
                    i, c["Tipo_Opcion"], c["Vencimiento"], c["Contratos"],
                    f"${c['Strike_Min']} — ${c['Strike_Max']}",
                    _fmt_monto(c['Prima_Total']), _fmt_monto(c['Prima_Promedio']),
                    _fmt_entero(c['Vol_Total']), _fmt_entero(c['OI_Total']),
                ])
            _tabla_datos_report(doc, headers_cl, rows_cl)

            # Detalle de cada cluster
            for i, c in enumerate(st.session_state.clusters_detectados, 1):
                if c.get("Detalle"):
                    p_cl = doc.add_paragraph()
                    run_cl = p_cl.add_run(f"Detalle Cluster #{i} — {c['Tipo_Opcion']} Venc. {c['Vencimiento']}")
                    run_cl.bold = True
                    run_cl.font.size = Pt(10)
                    run_cl.font.name = "Calibri"

                    headers_det = ["#", "Strike", "Volumen", "OI", "Prima Total"]
                    rows_det = []
                    for j, d in enumerate(c["Detalle"], 1):
                        rows_det.append([
                            j, _fmt_precio(d['Strike']),
                            _fmt_entero(d['Volumen']), _fmt_entero(d['OI']),
                            _fmt_monto(d['Prima_Volumen']),
                        ])
                    _tabla_datos_report(doc, headers_det, rows_det)

        # Todas las opciones escaneadas
        if st.session_state.datos_completos:
            _agregar_titulo_report(doc, "TODAS LAS OPCIONES ESCANEADAS", level=1)

            datos_sorted = sorted(
                st.session_state.datos_completos,
                key=lambda x: x.get("Prima_Volumen", 0), reverse=True,
            )

            p_info = doc.add_paragraph()
            run_info = p_info.add_run(f"Total de opciones: {len(datos_sorted):,}")
            run_info.font.size = Pt(10)
            run_info.font.italic = True
            run_info.font.name = "Calibri"

            headers_opt = ["Tipo", "Vencimiento", "Strike", "Volumen", "OI",
                           "Ask", "Bid", "Último", "Lado", "IV", "Prima Total"]
            rows_opt = []
            datos_enriquecidos = _enriquecer_datos_opcion(datos_sorted, precio_subyacente)
            
            for d in datos_enriquecidos:
                rows_opt.append([
                    d["Tipo"], d["Vencimiento"], _fmt_precio(d['Strike']),
                    _fmt_entero(d['Volumen']), _fmt_entero(d['OI']),
                    _fmt_precio(d['Ask']), _fmt_precio(d['Bid']), _fmt_precio(d['Ultimo']),
                    _fmt_lado(d.get('Lado', 'N/A')), _fmt_iv(d.get('IV', 0)),
                    _fmt_monto(d.get('Prima_Volumen', 0)),
                ])
            _tabla_datos_report(doc, headers_opt, rows_opt)

        # Pie de pígina
        doc.add_paragraph("")
        pie = doc.add_paragraph()
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_pie = pie.add_run(f"Monitor de Opciones — Reporte Live Scanning — {fecha_legible}")
        run_pie.font.size = Pt(8)
        run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        run_pie.font.name = "Calibri"

        # Retornar bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # =============================================
    # FUNCIÓN 2: REPORTE OPEN INTEREST
    # =============================================
    def _generar_reporte_open_interest():
        """Genera reporte DOCX con anílisis de Open Interest."""
        doc = Document()
        
        # Configurar pígina
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # Portada
        doc.add_paragraph("")
        titulo = doc.add_heading("REPORTE — OPEN INTEREST", level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        ticker_name = st.session_state.get("ticker_anterior", "N/A")
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = subtitulo.add_run(f"Ticker: {ticker_name}")
        run_sub.font.size = Pt(18)
        run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
        run_sub.font.name = "Calibri"
        run_sub.bold = True

        fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        fecha_p = doc.add_paragraph()
        fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
        run_fecha.font.size = Pt(11)
        run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run_fecha.font.name = "Calibri"

        doc.add_paragraph("")

        # Datos de Barchart
        if st.session_state.barchart_data is not None and not st.session_state.barchart_data.empty:
            df_bc = st.session_state.barchart_data.copy()
            
            df_positivos = df_bc[df_bc["OI_Chg"] > 0].sort_values("OI_Chg", ascending=False)
            df_negativos = df_bc[df_bc["OI_Chg"] < 0].sort_values("OI_Chg", ascending=True)
            
            n_total = len(df_bc)
            n_pos = len(df_positivos)
            n_neg = len(df_negativos)
            n_calls = len(df_bc[df_bc["Tipo"] == "CALL"]) if "Tipo" in df_bc.columns else 0
            n_puts = len(df_bc[df_bc["Tipo"] == "PUT"]) if "Tipo" in df_bc.columns else 0
            
            contratos_abiertos = int(df_positivos["OI_Chg"].sum()) if n_pos > 0 else 0
            contratos_cerrados = int(df_negativos["OI_Chg"].sum()) if n_neg > 0 else 0

            # Resumen
            _agregar_titulo_report(doc, "RESUMEN DE CAMBIOS EN OPEN INTEREST", level=1)
            _tabla_info_report(doc, {
                "Ticker": ticker_name,
                "Fecha": fecha_legible,
                "Total Contratos Analizados": f"{n_total:,}",
                "Calls": f"{n_calls:,}",
                "Puts": f"{n_puts:,}",
                "Contratos Abiertos (OI Positivo)": f"{contratos_abiertos:,}",
                "Contratos Cerrados (OI Negativo)": f"{contratos_cerrados:,}",
                "Señales Positivas": f"{n_pos:,}",
                "Señales Negativas": f"{n_neg:,}",
            })

            # Tabla OI Positivo
            if n_pos > 0:
                _agregar_titulo_report(doc, f"OI POSITIVO — ABRIENDO POSICIONES ({n_pos})", level=1)
                p_desc = doc.add_paragraph()
                run_d = p_desc.add_run(
                    "Contratos donde el Open Interest aumentó, indicando nuevas posiciones abiertas."
                )
                run_d.font.size = Pt(10)
                run_d.font.italic = True
                run_d.font.name = "Calibri"

                headers_pos = ["#", "Tipo", "Strike", "Vencimiento", "DTE", "Volumen", "OI", "OI Chg", "IV", "Delta", "Último"]
                rows_pos = []
                for i, row in enumerate(df_positivos.head(100).itertuples(), 1):
                    rows_pos.append([
                        i,
                        row.Tipo if hasattr(row, 'Tipo') else "N/A",
                        f"${row.Strike:,.1f}" if hasattr(row, 'Strike') else "N/A",
                        row.Vencimiento if hasattr(row, 'Vencimiento') else "N/A",
                        f"{row.DTE}d" if hasattr(row, 'DTE') else "N/A",
                        f"{int(row.Volumen):,}" if hasattr(row, 'Volumen') else "N/A",
                        f"{int(row.OI):,}" if hasattr(row, 'OI') else "N/A",
                        f"+{int(row.OI_Chg):,}" if hasattr(row, 'OI_Chg') else "N/A",
                        f"{row.IV:.1f}%" if hasattr(row, 'IV') and row.IV > 0 else "N/A",
                        f"{row.Delta:.3f}" if hasattr(row, 'Delta') and row.Delta != 0 else "N/A",
                        f"${row.Último:.2f}" if hasattr(row, 'Último') and row.Último > 0 else "N/A",
                    ])
                _tabla_datos_report(doc, headers_pos, rows_pos)

            # Tabla OI Negativo
            if n_neg > 0:
                _agregar_titulo_report(doc, f"OI NEGATIVO — CERRANDO POSICIONES ({n_neg})", level=1)
                p_desc = doc.add_paragraph()
                run_d = p_desc.add_run(
                    "Contratos donde el Open Interest disminuyó, indicando posiciones cerradas o ejercidas."
                )
                run_d.font.size = Pt(10)
                run_d.font.italic = True
                run_d.font.name = "Calibri"

                headers_neg = ["#", "Tipo", "Strike", "Vencimiento", "DTE", "Volumen", "OI", "OI Chg", "IV", "Delta", "Último"]
                rows_neg = []
                for i, row in enumerate(df_negativos.head(100).itertuples(), 1):
                    rows_neg.append([
                        i,
                        row.Tipo if hasattr(row, 'Tipo') else "N/A",
                        f"${row.Strike:,.1f}" if hasattr(row, 'Strike') else "N/A",
                        row.Vencimiento if hasattr(row, 'Vencimiento') else "N/A",
                        f"{row.DTE}d" if hasattr(row, 'DTE') else "N/A",
                        f"{int(row.Volumen):,}" if hasattr(row, 'Volumen') else "N/A",
                        f"{int(row.OI):,}" if hasattr(row, 'OI') else "N/A",
                        f"{int(row.OI_Chg):,}" if hasattr(row, 'OI_Chg') else "N/A",
                        f"{row.IV:.1f}%" if hasattr(row, 'IV') and row.IV > 0 else "N/A",
                        f"{row.Delta:.3f}" if hasattr(row, 'Delta') and row.Delta != 0 else "N/A",
                        f"${row.Último:.2f}" if hasattr(row, 'Último') and row.Último > 0 else "N/A",
                    ])
                _tabla_datos_report(doc, headers_neg, rows_neg)

        else:
            # Sin datos
            p_sin = doc.add_paragraph()
            run_sin = p_sin.add_run("No hay datos de Open Interest disponibles. Ejecuta un escaneo primero.")
            run_sin.font.size = Pt(11)
            run_sin.font.italic = True
            run_sin.font.name = "Calibri"

        # Pie de pígina
        doc.add_paragraph("")
        pie = doc.add_paragraph()
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_pie = pie.add_run(f"Monitor de Opciones — Reporte Open Interest — {fecha_legible}")
        run_pie.font.size = Pt(8)
        run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        run_pie.font.name = "Calibri"

        # Retornar bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # =============================================
    # FUNCIÓN 3: REPORTE IMPORTANT COMPANIES
    # =============================================
    def _generar_reporte_important_companies():
        """Genera reporte DOCX con anílisis detallado de Important Companies."""
        doc = Document()
        
        # Configurar pígina
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # Portada
        doc.add_paragraph("")
        titulo = doc.add_heading("REPORTE — IMPORTANT COMPANIES", level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = subtitulo.add_run("Anílisis de Empresas Consolidadas y Emergentes")
        run_sub.font.size = Pt(16)
        run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
        run_sub.font.name = "Calibri"
        run_sub.bold = True

        fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        fecha_p = doc.add_paragraph()
        fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
        run_fecha.font.size = Pt(11)
        run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run_fecha.font.name = "Calibri"

        doc.add_paragraph("")

        # EMPRESAS CONSOLIDADAS
        if "proyecciones_resultados" in st.session_state and st.session_state.proyecciones_resultados:
            resultados = st.session_state.proyecciones_resultados
            
            _agregar_titulo_report(doc, f"EMPRESAS CONSOLIDADAS ({len(resultados)})", level=1)
            p_desc = doc.add_paragraph()
            run_d = p_desc.add_run(
                "Grandes corporaci ones con historial probado y proyección de crecimiento sostenido. "
                "Anílisis fundamental + tîcnico + sentimiento."
            )
            run_d.font.size = Pt(10)
            run_d.font.italic = True
            run_d.font.name = "Calibri"

            # Resumen mîtricas
            alta = sum(1 for r in resultados if r.get("veredicto", "").startswith("OPORTUNIDAD"))
            considerar = sum(1 for r in resultados if "CONSIDERAR" in r.get("veredicto", ""))
            mantener = sum(1 for r in resultados if "MANTENER" in r.get("veredicto", ""))
            precaucion = sum(1 for r in resultados if "PRECAUCIÓN" in r.get("veredicto", "") or "PRECAU" in r.get("veredicto", ""))

            _tabla_info_report(doc, {
                "Total Empresas": len(resultados),
                "Oportunidad de Compra": alta,
                "Considerar": considerar,
                "Mantener": mantener,
                "Precaución": precaucion,
            })

            # Tabla comparativa
            headers_comp = ["#", "Ticker", "Empresa", "Precio", "Score Fund.", "Score Tîc.", "Score Comb.", "Veredicto",
                            "Crec. Ingresos", "Margen Op.", "P/E Fwd", "PEG", "Tendencia", "RSI", "Target", "Upside"]
            rows_comp = []
            for i, r in enumerate(resultados, 1):
                tecnico = r.get("tecnico", {})
                rows_comp.append([
                    i,
                    r["symbol"],
                    r["nombre"][:30],
                    f"${r['precio']:,.2f}",
                    f"{r.get('score', 0)}/100",
                    f"{r.get('score_tecnico', 0)}/100",
                    f"{r.get('score_combinado', 0)}/100",
                    r.get("veredicto", "N/A")[:25],
                    f"{r['revenue_growth']*100:.1f}%",
                    f"{r['operating_margins']*100:.1f}%",
                    f"{r['forward_pe']:.1f}x" if r['forward_pe'] > 0 else "N/A",
                    f"{r['peg_ratio']:.2f}" if r['peg_ratio'] > 0 else "N/A",
                    tecnico.get("tendencia", "N/A"),
                    f"{tecnico.get('rsi', 0):.0f}" if tecnico else "N/A",
                    f"${r['target_mean']:,.0f}" if r.get('target_mean', 0) > 0 else "N/A",
                    f"{r['upside_pct']:.1f}%" if r.get('upside_pct') else "N/A",
                ])
            _tabla_datos_report(doc, headers_comp, rows_comp)

            # Detalle por empresa
            for r in resultados:
                doc.add_page_break()
                _agregar_titulo_report(doc, f"{r['symbol']} — {r['nombre']}", level=2)
                
                tecnico = r.get("tecnico", {})
                
                # Info bísica
                _tabla_info_report(doc, {
                    "Precio Actual": f"${r['precio']:,.2f}",
                    "Market Cap": format_market_cap(r.get("market_cap", 0)),
                    "Sector": r.get("sector", "N/A"),
                    "Industria": r.get("industria", "N/A"),
                    "Score Combinado": f"{r.get('score_combinado', 0)}/100",
                    "Veredicto": r.get("veredicto", "N/A"),
                })

                # Fundamental
                _agregar_titulo_report(doc, "📊 Anílisis Fundamental", level=3)
                _tabla_info_report(doc, {
                    "Ingresos Totales": f"${r.get('revenue', 0)/1e9:.1f}B" if r.get('revenue', 0) > 0 else "N/A",
                    "Crecimiento Ingresos": f"{r['revenue_growth']*100:.1f}%",
                    "Margen Bruto": f"{r['gross_margins']*100:.1f}%",
                    "Margen Operativo": f"{r['operating_margins']*100:.1f}%",
                    "Margen Neto": f"{r['profit_margins']*100:.1f}%",
                    "P/E Forward": f"{r['forward_pe']:.1f}x" if r['forward_pe'] > 0 else "N/A",
                    "P/E Trailing": f"{r['trailing_pe']:.1f}x" if r['trailing_pe'] > 0 else "N/A",
                    "PEG Ratio": f"{r['peg_ratio']:.2f}" if r['peg_ratio'] > 0 else "N/A",
                    "Free Cash Flow": f"${r.get('free_cashflow', 0)/1e9:.1f}B" if r.get('free_cashflow', 0) > 0 else "N/A",
                    "Crecimiento Beneficios": f"{r['earnings_growth']*100:.1f}%",
                })

                # Tîcnico
                if tecnico:
                    _agregar_titulo_report(doc, "📈 Anílisis Tîcnico", level=3)
                    _tabla_info_report(doc, {
                        "Tendencia": tecnico.get("tendencia", "N/A"),
                        "RSI (14)": f"{tecnico.get('rsi', 0):.0f}",
                        "ADX (14)": f"{tecnico.get('adx', 0):.0f}",
                        "SMA 20": f"${tecnico.get('sma_20', 0):,.2f}",
                        "SMA 50": f"${tecnico.get('sma_50', 0):,.2f}",
                        "SMA 200": f"${tecnico.get('sma_200', 0):,.2f}" if tecnico.get('sma_200', 0) > 0 else "N/A",
                        "Volumen Ratio": f"{tecnico.get('vol_ratio', 0):.2f}x",
                        "Soporte 20d": f"${tecnico.get('soporte_20d', 0):,.2f}",
                        "Resistencia 20d": f"${tecnico.get('resistencia_20d', 0):,.2f}",
                        "Rango 52sem": f"{tecnico.get('rango_52w_pct', 0):.0f}%",
                    })

                # Sentimiento
                _agregar_titulo_report(doc, "🎯 Sentimiento", level=3)
                _tabla_info_report(doc, {
                    "Recomendación": r.get("recommendation", "N/A").upper(),
                    "Número Analistas": r.get("num_analysts", 0),
                    "Target Medio": f"${r['target_mean']:,.2f}" if r.get('target_mean', 0) > 0 else "N/A",
                    "Target Alto": f"${r['target_high']:,.2f}" if r.get('target_high', 0) > 0 else "N/A",
                    "Target Bajo": f"${r['target_low']:,.2f}" if r.get('target_low', 0) > 0 else "N/A",
                    "Upside Potencial": f"{r['upside_pct']:.1f}%" if r.get('upside_pct') else "N/A",
                    "Beta": f"{r['beta']:.2f}" if r.get('beta', 0) > 0 else "N/A",
                    "52 Week Low": f"${r['fifty_two_low']:,.2f}" if r.get('fifty_two_low', 0) > 0 else "N/A",
                    "52 Week High": f"${r['fifty_two_high']:,.2f}" if r.get('fifty_two_high', 0) > 0 else "N/A",
                })

                # Razones fundamentales
                if r.get("razones"):
                    p_raz = doc.add_paragraph()
                    run_raz = p_raz.add_run("Factores del Score Fundamental:")
                    run_raz.bold = True
                    run_raz.font.size = Pt(10)
                    run_raz.font.name = "Calibri"
                    for razon in r["razones"]:
                        p_item = doc.add_paragraph(razon, style='List Bullet')
                        p_item.paragraph_format.left_indent = Pt(20)

                # Señales tîcnicas
                if r.get("señales_tecnicas"):
                    p_sen = doc.add_paragraph()
                    run_sen = p_sen.add_run("Señales Tîcnicas:")
                    run_sen.bold = True
                    run_sen.font.size = Pt(10)
                    run_sen.font.name = "Calibri"
                    for senal in r["señales_tecnicas"]:
                        p_item = doc.add_paragraph(senal, style='List Bullet')
                        p_item.paragraph_format.left_indent = Pt(20)

        # EMPRESAS EMERGENTES
        if "emergentes_resultados" in st.session_state and st.session_state.emergentes_resultados:
            doc.add_page_break()
            resultados_em = st.session_state.emergentes_resultados
            
            _agregar_titulo_report(doc, f"EMPRESAS EMERGENTES ({len(resultados_em)})", level=1)
            p_desc_em = doc.add_paragraph()
            run_d_em = p_desc_em.add_run(
                "Empresas innovadoras con alto potencial de crecimiento disruptivo a 10 años."
            )
            run_d_em.font.size = Pt(10)
            run_d_em.font.italic = True
            run_d_em.font.name = "Calibri"

            # Tabla comparativa emergentes (igual formato que consolidadas)
            headers_em = ["#", "Ticker", "Empresa", "Precio", "Score Comb.", "Veredicto",
                          "Crec. Ingresos", "P/E Fwd", "Target", "Upside"]
            rows_em = []
            for i, r in enumerate(resultados_em, 1):
                rows_em.append([
                    i,
                    r["symbol"],
                    r["nombre"][:30],
                    f"${r['precio']:,.2f}",
                    f"{r.get('score_combinado', 0)}/100",
                    r.get("veredicto", "N/A")[:25],
                    f"{r['revenue_growth']*100:.1f}%",
                    f"{r['forward_pe']:.1f}x" if r['forward_pe'] > 0 else "N/A",
                    f"${r['target_mean']:,.0f}" if r.get('target_mean', 0) > 0 else "N/A",
                    f"{r['upside_pct']:.1f}%" if r.get('upside_pct') else "N/A",
                ])
            _tabla_datos_report(doc, headers_em, rows_em)

        # Sin datos
        if ("proyecciones_resultados" not in st.session_state or not st.session_state.proyecciones_resultados) and \
           ("emergentes_resultados" not in st.session_state or not st.session_state.emergentes_resultados):
            p_sin = doc.add_paragraph()
            run_sin = p_sin.add_run("No hay datos de anílisis disponibles. Ejecuta el anílisis en Important Companies primero.")
            run_sin.font.size = Pt(11)
            run_sin.font.italic = True
            run_sin.font.name = "Calibri"

        # Pie de pígina
        doc.add_paragraph("")
        pie = doc.add_paragraph()
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_pie = pie.add_run(f"Monitor de Opciones — Reporte Important Companies — {fecha_legible}")
        run_pie.font.size = Pt(8)
        run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        run_pie.font.name = "Calibri"

        # Retornar bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # =============================================
    # FUNCIÓN 3.5: REPORTE DATA ANALYSIS (REAL)
    # =============================================
    def _generar_reporte_data_analysis():
        """Genera reporte DOCX con anílisis de sentimiento, soportes y resistencias del Live Scanning."""
        doc = Document()
        
        # Configurar pígina
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # Portada
        doc.add_paragraph("")
        titulo = doc.add_heading("REPORTE — DATA ANALYSIS", level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = subtitulo.add_run("Anílisis de Sentimiento, Soportes y Resistencias")
        run_sub.font.size = Pt(11)
        run_sub.font.italic = True
        run_sub.font.name = "Calibri"

        fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')

        # CONTENIDO
        if not st.session_state.datos_completos:
            p_sin = doc.add_paragraph()
            run_sin = p_sin.add_run("No hay datos de Live Scanning disponibles. Ejecuta el escaneo primero.")
            run_sin.font.size = Pt(11)
            run_sin.font.italic = True
            run_sin.font.name = "Calibri"
        else:
            df_analisis = pd.DataFrame(st.session_state.datos_completos)
            if "Prima_Volumen" in df_analisis.columns:
                df_analisis = df_analisis.rename(columns={"Prima_Volumen": "Prima_Vol"})

            ticker_symbol = st.session_state.get('ticker_actual') or st.session_state.get('ticker_anterior', 'N/A')
            precio_actual = st.session_state.get('precio_subyacente', None)

            # Header con ticker y precio
            _agregar_titulo_report(doc, f"TICKER: {ticker_symbol}", level=1)
            if precio_actual:
                p_precio = doc.add_paragraph()
                run_precio = p_precio.add_run(f"Precio Actual: ${precio_actual:,.2f}")
                run_precio.font.size = Pt(12)
                run_precio.font.bold = True
                run_precio.font.name = "Calibri"

            # ================================================================
            # SENTIMIENTO POR PRIMAS
            # ================================================================
            _agregar_titulo_report(doc, "💰 Desglose de Sentimiento por Primas", level=2)

            df_sent = df_analisis.copy()
            df_sent["_mid"] = (df_sent["Ask"] + df_sent["Bid"]) / 2

            mask_call = df_sent["Tipo"] == "CALL"
            mask_put = df_sent["Tipo"] == "PUT"
            mask_ask = df_sent["Ultimo"] >= df_sent["_mid"]
            mask_bid = df_sent["Ultimo"] < df_sent["_mid"]

            call_ask_val = df_sent.loc[mask_call & mask_ask, "Prima_Vol"].sum()
            call_bid_val = df_sent.loc[mask_call & mask_bid, "Prima_Vol"].sum()
            put_ask_val = df_sent.loc[mask_put & mask_ask, "Prima_Vol"].sum()
            put_bid_val = df_sent.loc[mask_put & mask_bid, "Prima_Vol"].sum()

            total_sent = call_ask_val + call_bid_val + put_ask_val + put_bid_val

            if total_sent > 0:
                bullish_total = call_ask_val + put_bid_val
                bearish_total = call_bid_val + put_ask_val
                net_pct = ((bullish_total - bearish_total) / total_sent) * 100

                _tabla_info_report(doc, {
                    "📞 CALL Ask (Compra agresiva)": f"${call_ask_val:,.0f} (+{call_ask_val/total_sent*100:.1f}%)",
                    "📞 CALL Bid (Venta agresiva)": f"${call_bid_val:,.0f} (-{call_bid_val/total_sent*100:.1f}%)",
                    "📋 PUT Ask (Compra agresiva)": f"${put_ask_val:,.0f} (-{put_ask_val/total_sent*100:.1f}%)",
                    "📋 PUT Bid (Venta agresiva)": f"${put_bid_val:,.0f} (+{put_bid_val/total_sent*100:.1f}%)",
                    "Total Prima": f"${total_sent:,.0f}",
                    "🟢 Alcista Total": f"${bullish_total:,.0f} ({bullish_total/total_sent*100:.1f}%)",
                    "🔴 Bajista Total": f"${bearish_total:,.0f} ({bearish_total/total_sent*100:.1f}%)",
                    "Sentimiento Neto": f"{'+' if net_pct >= 0 else ''}{net_pct:.1f}% ({'ALCISTA' if net_pct >= 0 else 'BAJISTA'})",
                })

            # ================================================================
            # SOPORTES Y RESISTENCIAS
            # ================================================================
            _agregar_titulo_report(doc, "🛡️ Soportes y Resistencias por Opciones", level=2)

            df_calls_sr = df_analisis[(df_analisis["Tipo"] == "CALL") & (df_analisis["Volumen"] > 0)].copy()
            df_puts_sr = df_analisis[(df_analisis["Tipo"] == "PUT") & (df_analisis["Volumen"] > 0)].copy()

            if not df_calls_sr.empty and not df_puts_sr.empty:
                # Top 5 CALL strikes → Soportes
                top_calls = df_calls_sr.groupby("Strike").agg(
                    Vol_Total=("Volumen", "sum"),
                    OI_Total=("OI", "sum"),
                    Prima_Total=("Prima_Vol", "sum"),
                ).sort_values("Vol_Total", ascending=False).head(5).reset_index()

                # Top 5 PUT strikes → Resistencias
                top_puts = df_puts_sr.groupby("Strike").agg(
                    Vol_Total=("Volumen", "sum"),
                    OI_Total=("OI", "sum"),
                    Prima_Total=("Prima_Vol", "sum"),
                ).sort_values("Vol_Total", ascending=False).head(5).reset_index()

                # Tabla de Soportes
                _agregar_titulo_report(doc, "🟢 Soportes (CALLs mís tradeados)", level=3)
                headers_s = ["Nivel", "Strike", "Volumen", "OI", "Prima Total"]
                rows_s = []
                for idx, row in top_calls.iterrows():
                    pct_str = ""
                    if precio_actual and precio_actual > 0:
                        dist = ((row["Strike"] - precio_actual) / precio_actual) * 100
                        pct_str = f" ({'+' if dist >= 0 else ''}{dist:.1f}%)"
                    rows_s.append([
                        f"S{idx+1}",
                        f"${row['Strike']:,.1f}{pct_str}",
                        f"{row['Vol_Total']:,.0f}",
                        f"{row['OI_Total']:,.0f}",
                        f"${row['Prima_Total']:,.0f}",
                    ])
                _tabla_datos_report(doc, headers_s, rows_s)

                # Tabla de Resistencias
                _agregar_titulo_report(doc, "🔴 Resistencias (PUTs mís tradeados)", level=3)
                headers_r = ["Nivel", "Strike", "Volumen", "OI", "Prima Total"]
                rows_r = []
                for idx, row in top_puts.iterrows():
                    pct_str = ""
                    if precio_actual and precio_actual > 0:
                        dist = ((row["Strike"] - precio_actual) / precio_actual) * 100
                        pct_str = f" ({'+' if dist >= 0 else ''}{dist:.1f}%)"
                    rows_r.append([
                        f"R{idx+1}",
                        f"${row['Strike']:,.1f}{pct_str}",
                        f"{row['Vol_Total']:,.0f}",
                        f"{row['OI_Total']:,.0f}",
                        f"${row['Prima_Total']:,.0f}",
                    ])
                _tabla_datos_report(doc, headers_r, rows_r)

            # ================================================================
            # DISTRIBUCIÓN CALL VS PUT
            # ================================================================
            _agregar_titulo_report(doc, "📊 Distribución CALL vs PUT", level=2)
            
            tipo_counts = df_analisis["Tipo"].value_counts()
            n_calls = tipo_counts.get("CALL", 0)
            n_puts = tipo_counts.get("PUT", 0)
            ratio_pc = n_puts / n_calls if n_calls > 0 else 0

            _tabla_info_report(doc, {
                "Total CALLs": f"{n_calls:,}",
                "Total PUTs": f"{n_puts:,}",
                "Put/Call Ratio": f"{ratio_pc:.3f}",
                "Interpretación": "Mayor actividad en CALLs (alcista)" if ratio_pc < 0.7 else "Ratio neutral",
            })

            # ================================================================
            # TOP 20 POR VOLUMEN
            # ================================================================
            _agregar_titulo_report(doc, "🎯 Top 20 Strikes por Volumen", level=2)
            
            vol_cols = ["Vencimiento", "Tipo", "Strike", "Volumen", "IV", "Ultimo", "Prima_Vol"]
            top_vol = df_analisis.nlargest(20, "Volumen")[[c for c in vol_cols if c in df_analisis.columns]].reset_index(drop=True)

            headers_vol = ["#", "Vencimiento", "Tipo", "Strike", "Volumen", "IV", "Último", "Prima Total"]
            rows_vol = []
            for i, row in top_vol.iterrows():
                rows_vol.append([
                    i + 1,
                    row.get("Vencimiento", "N/A"),
                    row.get("Tipo", "N/A"),
                    f"${row.get('Strike', 0):,.1f}",
                    f"{row.get('Volumen', 0):,}",
                    f"{row.get('IV', 0):.2f}%" if row.get('IV', 0) > 0 else "N/A",
                    f"${row.get('Ultimo', 0):.2f}",
                    f"${row.get('Prima_Vol', 0):,.0f}",
                ])
            _tabla_datos_report(doc, headers_vol, rows_vol)

            # ================================================================
            # TOP 20 POR OI
            # ================================================================
            _agregar_titulo_report(doc, "🏛️ Top 20 Strikes por Open Interest", level=2)
            
            oi_cols = ["Vencimiento", "Tipo", "Strike", "OI", "Volumen", "IV", "Ultimo", "Prima_Vol"]
            top_oi = df_analisis.nlargest(20, "OI")[[c for c in oi_cols if c in df_analisis.columns]].reset_index(drop=True)

            headers_oi = ["#", "Vencimiento", "Tipo", "Strike", "OI", "Volumen", "IV", "Último", "Prima Total"]
            rows_oi = []
            for i, row in top_oi.iterrows():
                rows_oi.append([
                    i + 1,
                    row.get("Vencimiento", "N/A"),
                    row.get("Tipo", "N/A"),
                    f"${row.get('Strike', 0):,.1f}",
                    f"{row.get('OI', 0):,}",
                    f"{row.get('Volumen', 0):,}",
                    f"{row.get('IV', 0):.2f}%" if row.get('IV', 0) > 0 else "N/A",
                    f"${row.get('Ultimo', 0):.2f}",
                    f"${row.get('Prima_Vol', 0):,.0f}",
                ])
            _tabla_datos_report(doc, headers_oi, rows_oi)

        # Pie de pígina
        doc.add_paragraph("")
        pie = doc.add_paragraph()
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_pie = pie.add_run(f"Monitor de Opciones — Reporte Data Analysis — {fecha_legible}")
        run_pie.font.size = Pt(8)
        run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        run_pie.font.name = "Calibri"

        # Retornar bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # =============================================
    # FUNCIÓN 4: REPORTE RANGE
    # =============================================
    def _generar_reporte_range():
        """Genera reporte DOCX con información del Rango Esperado."""
        doc = Document()
        
        # Configurar pígina
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        # Portada
        doc.add_paragraph("")
        titulo = doc.add_heading("REPORTE — RANGO ESPERADO", level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        r = st.session_state.rango_resultado
        ticker_name = r.get("symbol", "N/A")
        
        subtitulo = doc.add_paragraph()
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = subtitulo.add_run(f"Ticker: {ticker_name}")
        run_sub.font.size = Pt(18)
        run_sub.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
        run_sub.font.name = "Calibri"
        run_sub.bold = True

        fecha_legible = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
        fecha_p = doc.add_paragraph()
        fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fecha = fecha_p.add_run(f"Generado: {fecha_legible}")
        run_fecha.font.size = Pt(11)
        run_fecha.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run_fecha.font.name = "Calibri"

        doc.add_paragraph("")

        # Explicación
        _agregar_titulo_report(doc, "¿QUÉ ES EL RANGO ESPERADO?", level=1)
        p_exp = doc.add_paragraph()
        run_exp = p_exp.add_run(
            "El rango esperado es una estimación estadística del movimiento probable del precio del activo "
            "hasta la fecha de expiración, basado en la volatilidad implícita (IV) de las opciones. "
            "Se calcula con una desviación estíndar (1σ), lo que significa que hay aproximadamente 68% de "
            "probabilidad de que el precio permanezca dentro del rango calculado."
        )
        run_exp.font.size = Pt(10)
        run_exp.font.name = "Calibri"

        doc.add_paragraph("")

        # Parímetros del cílculo
        dias = r.get('dias_restantes')
        _agregar_titulo_report(doc, "PARÁMETROS DEL CÁLCULO", level=1)
        _tabla_info_report(doc, {
            "Símbolo": ticker_name,
            "Precio Actual del Subyacente": f"${r['underlying_price']:,.2f}",
            "Fecha de Expiración": r["expiration"],
            "Días Restantes (DTE)": dias if dias else "N/A",
            "Delta Objetivo": f"±{r.get('target_delta', 'N/A')}",
        })

        # Rango calculado
        _agregar_titulo_report(doc, "RANGO DE PRECIOS ESPERADO (1σ)", level=1)
        _tabla_info_report(doc, {
            "Rango Inferior": f"${r['expected_range_low']:,.2f}",
            "Precio Actual": f"${r['underlying_price']:,.2f}",
            "Rango Superior": f"${r['expected_range_high']:,.2f}",
            "Bajada Esperada": f"-${r['downside_points']:,.2f} (-{r['downside_percent']:.2f}%)",
            "Subida Esperada": f"+${r['upside_points']:,.2f} (+{r['upside_percent']:.2f}%)",
            "Rango Total de Movimiento": f"${r['total_range_points']:,.2f} ({r['total_range_pct']:.2f}%)",
        })

        # Contratos utilizados
        _agregar_titulo_report(doc, "CONTRATOS UTILIZADOS EN EL CÁLCULO", level=1)
        p_cont = doc.add_paragraph()
        run_cont = p_cont.add_run(
            "El rango se calcula utilizando las opciones Call y Put con deltas mís cercanos al objetivo configurado."
        )
        run_cont.font.size = Pt(10)
        run_cont.font.italic = True
        run_cont.font.name = "Calibri"

        _tabla_info_report(doc, {
            "Call Strike": f"${r['call_strike']}",
            "Call Delta": f"{r['call_delta']}",
            "Call IV": f"{r['call_iv']:.1f}%",
            "Put Strike": f"${r['put_strike']}",
            "Put Delta": f"{r['put_delta']}",
            "Put IV": f"{r['put_iv']:.1f}%",
        })

        # Interpretación
        _agregar_titulo_report(doc, "INTERPRETACIÓN", level=1)
        p_int = doc.add_paragraph()
        run_int = p_int.add_run(
            f"Basíndose en la volatilidad implícita actual, se espera que {ticker_name} se mueva "
            f"entre ${r['expected_range_low']:,.2f} y ${r['expected_range_high']:,.2f} antes del "
            f"{r['expiration']}. Esto representa un rango de movimiento de ±{r['total_range_pct']:.1f}%.\n\n"
            f"Este rango puede utilizarse para:\n"
            f"• Planificar estrategias de trading direccionales (si esperas movimiento fuera del rango)\n"
            f"• Diseñar estrategias neutrales (si esperas que el precio permanezca dentro del rango)\n"
            f"• Identificar niveles de soporte y resistencia probables\n"
            f"• Evaluar el riesgo de posiciones existentes"
        )
        run_int.font.size = Pt(10)
        run_int.font.name = "Calibri"

        # Aviso
        doc.add_paragraph("")
        p_aviso = doc.add_paragraph()
        run_aviso = p_aviso.add_run(
            "⚠️ AVISO: Este cílculo es una estimación estadística basada en la volatilidad implícita "
            "y no garantiza que el precio permanecerí dentro del rango. Los movimientos del mercado "
            "pueden ser impredecibles, especialmente ante eventos inesperados o noticias significativas."
        )
        run_aviso.font.size = Pt(9)
        run_aviso.font.italic = True
        run_aviso.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run_aviso.font.name = "Calibri"

        # Pie de pígina
        doc.add_paragraph("")
        pie = doc.add_paragraph()
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_pie = pie.add_run(f"Monitor de Opciones — Reporte Rango Esperado — {fecha_legible}")
        run_pie.font.size = Pt(8)
        run_pie.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        run_pie.font.name = "Calibri"

        # Retornar bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # =============================================
    # BOTONES DE DESCARGA
    # =============================================
    st.markdown("---")
    st.markdown("#### 📥 Descargar Reportes")
    st.caption("Genera reportes detallados en formato DOCX con los datos cargados en cada sección.")

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    # Verificar disponibilidad de datos
    tiene_scanning = st.session_state.scan_count > 0 and st.session_state.datos_completos
    tiene_oi = tiene_scanning and st.session_state.barchart_data is not None and not st.session_state.barchart_data.empty
    tiene_analysis = ("proyecciones_resultados" in st.session_state and st.session_state.proyecciones_resultados) or \
                     ("emergentes_resultados" in st.session_state and st.session_state.emergentes_resultados)
    tiene_range = st.session_state.rango_resultado is not None

    # Botón 1: Live Scanning
    if tiene_scanning:
        ticker_name = st.session_state.get("ticker_anterior", "SCAN")
        with st.spinner("📊 Generando reporte de Live Scanning..."):
            try:
                docx_scanning = _generar_reporte_live_scanning()
                st.download_button(
                    "📊 Descargar Reporte Live Scanning (DOCX)",
                    docx_scanning,
                    f"reporte_live_scanning_{ticker_name}_{timestamp}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_scanning",
                    help="Descarga todos los datos escaneados: alertas, clusters, y todas las opciones analizadas.",
                )
            except Exception as e:
                st.error(f"⚠️ Error al generar reporte de Live Scanning: {e}")
    else:
        st.info("📊 **Reporte Live Scanning** — Ejecuta un escaneo primero en 🔍 Live Scanning")

    # Botón 2: Open Interest
    if tiene_oi:
        ticker_name = st.session_state.get("ticker_anterior", "SCAN")
        with st.spinner("📊 Generando reporte de Open Interest..."):
            try:
                docx_oi = _generar_reporte_open_interest()
                st.download_button(
                    "📊 Descargar Reporte Open Interest (DOCX)",
                    docx_oi,
                    f"reporte_open_interest_{ticker_name}_{timestamp}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_oi",
                    help="Descarga el anílisis completo de cambios en Open Interest (OI positivo y negativo).",
                )
            except Exception as e:
                st.error(f"⚠️ Error al generar reporte de Open Interest: {e}")
    else:
        st.info("📊 **Reporte Open Interest** — Ejecuta un escaneo primero en 🔍 Live Scanning")

    # Botón 3: Important Companies
    if tiene_analysis:
        with st.spinner("📊 Generando reporte de Important Companies..."):
            try:
                docx_important = _generar_reporte_important_companies()
                st.download_button(
                    "🏢 Descargar Reporte Important Companies (DOCX)",
                    docx_important,
                    f"reporte_important_companies_{timestamp}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_important",
                    help="Descarga el anílisis completo de Important Companies: fundamental, tîcnico, sentimiento y veredicto.",
                )
            except Exception as e:
                st.error(f"⚠️ Error al generar reporte de Important Companies: {e}")
    else:
        st.info("🏢 **Reporte Important Companies** — Ejecuta el anílisis en 🏢 Important Companies primero")

    # Botón 4: Data Analysis
    if tiene_scanning:
        ticker_name = st.session_state.get("ticker_anterior", "ANALYSIS")
        with st.spinner("📊 Generando reporte de Data Analysis..."):
            try:
                docx_analysis = _generar_reporte_data_analysis()
                st.download_button(
                    "📈 Descargar Reporte Data Analysis (DOCX)",
                    docx_analysis,
                    f"reporte_data_analysis_{ticker_name}_{timestamp}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_analysis",
                    help="Descarga el anílisis de sentimiento, soportes y resistencias basado en el Live Scanning.",
                )
            except Exception as e:
                st.error(f"⚠️ Error al generar reporte de Data Analysis: {e}")
    else:
        st.info("📈 **Reporte Data Analysis** — Ejecuta un escaneo primero en 🔍 Live Scanning")

    # Botón 5: Range
    if tiene_range:
        ticker_name = st.session_state.rango_resultado.get("symbol", "RANGE")
        with st.spinner("📊 Generando reporte de Rango Esperado..."):
            try:
                docx_range = _generar_reporte_range()
                st.download_button(
                    "📐 Descargar Reporte Rango Esperado (DOCX)",
                    docx_range,
                    f"reporte_rango_{ticker_name}_{timestamp}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_range",
                    help="Descarga el cílculo detallado del rango esperado con explicación e interpretación.",
                )
            except Exception as e:
                st.error(f"⚠️ Error al generar reporte de Rango: {e}")
    else:
        st.info("📐 **Reporte Rango Esperado** — Calcula el rango en 📐 Range primero")

    st.markdown("---")
    st.success("✅ Selecciona los reportes que deseas descargar. Los archivos .docx son editables y tienen estructura profesional.")


# ============================================================================
#   📈 DATA ANALYSIS
# ============================================================================
elif pagina == "📈 Data Analysis":
    st.markdown("### 📈 Data Analysis")

    if not st.session_state.datos_completos:
        st.info("Ejecuta un escaneo primero para ver los anílisis.")
    else:
        df_analisis = pd.DataFrame(st.session_state.datos_completos)
        # Renombrar columna para consistencia en esta sección
        if "Prima_Volumen" in df_analisis.columns:
            df_analisis = df_analisis.rename(columns={"Prima_Volumen": "Prima_Vol"})
        
        titulo_datos = f"Datos del último escaneo — {ticker_symbol}"
        
        st.caption(f"*{titulo_datos}* — {len(df_analisis):,} registros")

        # ================================================================
        # DESGLOSE DE SENTIMIENTO POR PRIMAS
        # ================================================================
        st.markdown("### 💰 Desglose de Sentimiento por Primas")
        st.markdown("---")

        # Clasificar opciones por lado de ejecución (Bid vs Ask)
        df_sent = df_analisis.copy()
        df_sent["_mid"] = (df_sent["Ask"] + df_sent["Bid"]) / 2

        mask_call = df_sent["Tipo"] == "CALL"
        mask_put = df_sent["Tipo"] == "PUT"
        mask_ask = df_sent["Ultimo"] >= df_sent["_mid"]
        mask_bid = df_sent["Ultimo"] < df_sent["_mid"]

        # CALL Ask = compra agresiva de calls → ALCISTA
        # CALL Bid = venta agresiva de calls → BAJISTA
        # PUT Ask = compra agresiva de puts → BAJISTA
        # PUT Bid = venta agresiva de puts → ALCISTA
        call_ask_val = df_sent.loc[mask_call & mask_ask, "Prima_Vol"].sum()
        call_bid_val = df_sent.loc[mask_call & mask_bid, "Prima_Vol"].sum()
        put_ask_val = df_sent.loc[mask_put & mask_ask, "Prima_Vol"].sum()
        put_bid_val = df_sent.loc[mask_put & mask_bid, "Prima_Vol"].sum()

        total_sent = call_ask_val + call_bid_val + put_ask_val + put_bid_val

        if total_sent > 0:
            # Porcentajes con signo: + alcista, - bajista
            rows_data = [
                ("📞 CALL Ask", "Compra agresiva", call_ask_val, +(call_ask_val / total_sent * 100), True),
                ("📞 CALL Bid", "Venta agresiva", call_bid_val, -(call_bid_val / total_sent * 100), False),
                ("📋 PUT Ask", "Compra agresiva", put_ask_val, -(put_ask_val / total_sent * 100), False),
                ("📋 PUT Bid", "Venta agresiva", put_bid_val, +(put_bid_val / total_sent * 100), True),
            ]

            bullish_total = call_ask_val + put_bid_val
            bearish_total = call_bid_val + put_ask_val
            net_pct = ((bullish_total - bearish_total) / total_sent) * 100

            max_abs = max(abs(r[3]) for r in rows_data)
            if max_abs == 0:
                max_abs = 1

            # Generar HTML compacto con clases CSS para evitar truncamiento
            rows_html = ""
            for label, desc, amount, pct, is_bull in rows_data:
                cc = "g" if is_bull else "r"
                pct_str = f"+{pct:.1f}%" if pct >= 0 else f"{pct:.1f}%"
                bar_w = abs(pct) / max_abs * 44

                if is_bull:
                    fill_s = f"left:50%;width:{bar_w:.1f}%;background:linear-gradient(90deg,rgba(16,185,129,.6),rgba(5,150,105,.2));border-radius:0 6px 6px 0"
                else:
                    fill_s = f"right:50%;width:{bar_w:.1f}%;background:linear-gradient(270deg,rgba(239,68,68,.6),rgba(185,28,28,.2));border-radius:6px 0 0 6px"

                rows_html += (
                    f'<div class="sr"><div class="sl"><div class="slt">{label}</div>'
                    f'<div class="sld">{desc}</div></div>'
                    f'<div class="sa {cc}">{_fmt_monto(amount)}</div>'
                    f'<div class="sb"><div class="sm"></div>'
                    f'<div class="sf" style="{fill_s}"></div></div>'
                    f'<div class="sp {cc}">{pct_str}</div></div>'
                )

            # Barra de sentimiento neto
            net_color = "#10b981" if net_pct >= 0 else "#ef4444"
            net_label = "ALCISTA" if net_pct >= 0 else "BAJISTA"
            net_emoji = "🟢" if net_pct >= 0 else "🔴"
            net_pct_str = f"+{net_pct:.1f}%" if net_pct >= 0 else f"{net_pct:.1f}%"
            bull_pct = bullish_total / total_sent * 100
            bear_pct = bearish_total / total_sent * 100
            net_bar_w = max(abs(bull_pct - bear_pct) / 100 * 44, 8)
            nc = "g" if net_pct >= 0 else "r"

            if net_pct >= 0:
                net_fill = f"left:50%;width:{net_bar_w:.1f}%;background:linear-gradient(90deg,rgba(16,185,129,.8),rgba(5,150,105,.3));border-radius:0 6px 6px 0"
            else:
                net_fill = f"right:50%;width:{net_bar_w:.1f}%;background:linear-gradient(270deg,rgba(239,68,68,.8),rgba(185,28,28,.3));border-radius:6px 0 0 6px"

            # --- OKA Sentiment Gauge (Plotly) ---
            gauge_score = max(0, min(100, 50 + net_pct / 2))  # Normalizar a 0-100
            if net_pct >= 10:
                gauge_lbl = "ALCISTA"
            elif net_pct <= -10:
                gauge_lbl = "BAJISTA"
            else:
                gauge_lbl = "NEUTRAL"

            fig_gauge = go.Figure(go.Indicator(
                mode="gauge+number+delta",
                value=gauge_score,
                domain={"x": [0, 1], "y": [0, 1]},
                title={"text": f"OKA Sentiment Index — {gauge_lbl}", "font": {"size": 16, "color": "white"}},
                number={"font": {"size": 42, "color": "white"}, "suffix": "/100"},
                delta={"reference": 50, "increasing": {"color": "#00ff88"}, "decreasing": {"color": "#ef4444"}},
                gauge={
                    "axis": {"range": [0, 100], "tickwidth": 1, "tickcolor": "#475569", "tickfont": {"color": "#94a3b8", "size": 11}},
                    "bar": {"color": "#00ff88", "thickness": 0.3},
                    "bgcolor": "#0f172a",
                    "borderwidth": 0,
                    "steps": [
                        {"range": [0, 30], "color": "rgba(239, 68, 68, 0.25)"},
                        {"range": [30, 50], "color": "rgba(245, 158, 11, 0.15)"},
                        {"range": [50, 70], "color": "rgba(16, 185, 129, 0.15)"},
                        {"range": [70, 100], "color": "rgba(0, 255, 136, 0.2)"},
                    ],
                    "threshold": {
                        "line": {"color": "white", "width": 3},
                        "thickness": 0.8,
                        "value": gauge_score,
                    },
                },
            ))
            fig_gauge.update_layout(
                paper_bgcolor="#1e293b",
                plot_bgcolor="#1e293b",
                font={"color": "white", "family": "Inter, sans-serif"},
                height=400,
                margin=dict(l=30, r=30, t=60, b=10),
            )
            st.plotly_chart(fig_gauge, use_container_width=True)

            # Bullish / Bearish / Neutral label debajo del gauge
            _gauge_color = "#00ff88" if gauge_lbl == "ALCISTA" else "#ef4444" if gauge_lbl == "BAJISTA" else "#f59e0b"
            st.markdown(
                f'<h3 style="text-align:center;color:{_gauge_color};margin:-10px 0 8px;font-weight:800;">{gauge_lbl}</h3>',
                unsafe_allow_html=True,
            )

            # Footer stats below gauge
            st.markdown(
                f'<div style="display:flex;justify-content:space-around;padding:8px 0 12px;'
                f'background:#1e293b;border-radius:0 0 12px 12px;margin-top:-10px;">'
                f'<div style="text-align:center"><div style="color:#94a3b8;font-size:.75rem">Bullish</div>'
                f'<div style="color:#10b981;font-weight:700">{bull_pct:.1f}%</div></div>'
                f'<div style="text-align:center"><div style="color:#94a3b8;font-size:.75rem">Score</div>'
                f'<div style="color:white;font-weight:700">{gauge_score:.0f}/100</div></div>'
                f'<div style="text-align:center"><div style="color:#94a3b8;font-size:.75rem">Bearish</div>'
                f'<div style="color:#ef4444;font-weight:700">{bear_pct:.1f}%</div></div>'
                f'</div>',
                unsafe_allow_html=True,
            )

            st.markdown(
                f'<div class="sp0">'
                f'<div class="tt">💰 Desglose de Sentimiento por Primas</div>'
                f'<div class="ts">Prima ejecutada por lado del order book — Compras vs Ventas agresivas</div>'
                f'{rows_html}'
                f'<div class="sn"><div class="snr">'
                f'<div class="snl"><div class="snt">{net_emoji} NETO</div><div class="snd {nc}">{net_label}</div></div>'
                f'<div class="sa {nc}">{_fmt_monto(abs(bullish_total - bearish_total))}</div>'
                f'<div class="sb"><div class="sm"></div><div class="sf" style="{net_fill}"></div></div>'
                f'<div class="sp {nc}">{net_pct_str}</div>'
                f'</div></div>'
                f'<div class="ssum">'
                f'<div class="ssi"><div class="ssh">🟢 Alcista</div><div class="ssv g">{_fmt_monto(bullish_total)}</div><div class="ssp g">{bull_pct:.1f}%</div></div>'
                f'<div class="ssi"><div class="ssh">📊 Total</div><div class="ssv w">{_fmt_monto(total_sent)}</div><div class="ssp gy">100%</div></div>'
                f'<div class="ssi"><div class="ssh">🔴 Bajista</div><div class="ssv r">{_fmt_monto(bearish_total)}</div><div class="ssp r">{bear_pct:.1f}%</div></div>'
                f'</div></div>',
                unsafe_allow_html=True,
            )
        else:
            st.info("Sin datos suficientes para calcular el sentimiento por primas.")

        st.markdown("---")

        # ================================================================
        # SOPORTES Y RESISTENCIAS POR VOLUMEN DE OPCIONES
        # ================================================================
        st.markdown("### 🛡️ Soportes y Resistencias por Opciones")

        # Obtener precio actual
        precio_actual = st.session_state.get('precio_subyacente', None)

        # Separar CALLs y PUTs con volumen > 0
        df_calls_sr = df_analisis[(df_analisis["Tipo"] == "CALL") & (df_analisis["Volumen"] > 0)].copy()
        df_puts_sr = df_analisis[(df_analisis["Tipo"] == "PUT") & (df_analisis["Volumen"] > 0)].copy()

        if not df_calls_sr.empty and not df_puts_sr.empty:
            # Top 5 strikes con mís volumen en CALLs → Soportes
            top_calls = df_calls_sr.groupby("Strike").agg(
                Vol_Total=("Volumen", "sum"),
                OI_Total=("OI", "sum"),
                Prima_Total=("Prima_Vol", "sum"),
                Contratos=("Volumen", "count"),
            ).sort_values("Vol_Total", ascending=False).head(5).reset_index()

            # Top 5 strikes con mís volumen en PUTs → Resistencias
            top_puts = df_puts_sr.groupby("Strike").agg(
                Vol_Total=("Volumen", "sum"),
                OI_Total=("OI", "sum"),
                Prima_Total=("Prima_Vol", "sum"),
                Contratos=("Volumen", "count"),
            ).sort_values("Vol_Total", ascending=False).head(5).reset_index()

            col_sr1, col_sr2 = st.columns(2)

            with col_sr1:
                st.markdown("#### 🟢 Soportes (Más tradeados)")
                for idx_s, row_s in top_calls.iterrows():
                    pct_dist = ""
                    if precio_actual and precio_actual > 0:
                        dist = ((row_s["Strike"] - precio_actual) / precio_actual) * 100
                        pct_dist = f" ({'+' if dist >= 0 else ''}{dist:.1f}%)"
                    st.markdown(
                        f"""
                        <div style="background: rgba(16, 185, 129, 0.08); border: 1px solid rgba(16, 185, 129, 0.2); 
                             border-radius: 10px; padding: 10px 14px; margin-bottom: 8px;">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <div>
                                    <span style="font-size: 1.1rem; font-weight: 700; color: #10b981;">
                                        S{idx_s + 1}: ${row_s['Strike']:,.1f}
                                    </span>
                                    <span style="font-size: 0.8rem; color: #94a3b8;">{pct_dist}</span>
                                </div>
                                <div style="text-align: right;">
                                    <span style="font-size: 0.82rem; color: #f1f5f9;">
                                        Vol: <b>{row_s['Vol_Total']:,.0f}</b>
                                    </span>
                                </div>
                            </div>
                            <div style="font-size: 0.75rem; color: #94a3b8; margin-top: 4px;">
                                OI: {row_s['OI_Total']:,.0f} | Prima: {_fmt_monto(row_s['Prima_Total'])} | {int(row_s['Contratos'])} contratos
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

            with col_sr2:
                st.markdown("#### � Resistencias (Más tradeados)")
                for idx_r, row_r in top_puts.iterrows():
                    pct_dist = ""
                    if precio_actual and precio_actual > 0:
                        dist = ((row_r["Strike"] - precio_actual) / precio_actual) * 100
                        pct_dist = f" ({'+' if dist >= 0 else ''}{dist:.1f}%)"
                    st.markdown(
                        f"""
                        <div style="background: rgba(239, 68, 68, 0.08); border: 1px solid rgba(239, 68, 68, 0.2); 
                             border-radius: 10px; padding: 10px 14px; margin-bottom: 8px;">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <div>
                                    <span style="font-size: 1.1rem; font-weight: 700; color: #ef4444;">
                                        R{idx_r + 1}: ${row_r['Strike']:,.1f}
                                    </span>
                                    <span style="font-size: 0.8rem; color: #94a3b8;">{pct_dist}</span>
                                </div>
                                <div style="text-align: right;">
                                    <span style="font-size: 0.82rem; color: #f1f5f9;">
                                        Vol: <b>{row_r['Vol_Total']:,.0f}</b>
                                    </span>
                                </div>
                            </div>
                            <div style="font-size: 0.75rem; color: #94a3b8; margin-top: 4px;">
                                OI: {row_r['OI_Total']:,.0f} | Prima: {_fmt_monto(row_r['Prima_Total'])} | {int(row_r['Contratos'])} contratos
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

            # Barra visual de niveles
            if precio_actual and precio_actual > 0:
                st.markdown("---")
                st.markdown("#### 📍 Mapa de Niveles vs Precio Actual")

                # Combinar todos los niveles
                niveles_s = [(s, "S", v) for s, v in zip(top_calls["Strike"], top_calls["Vol_Total"])]
                niveles_r = [(s, "R", v) for s, v in zip(top_puts["Strike"], top_puts["Vol_Total"])]
                todos_niveles = sorted(niveles_r + niveles_s, key=lambda x: x[0])

                max_vol_nivel = max(n[2] for n in todos_niveles) if todos_niveles else 1
                all_strikes = [n[0] for n in todos_niveles] + [precio_actual]
                rango_min = min(all_strikes) * 0.998
                rango_max = max(all_strikes) * 1.002
                rango_total = rango_max - rango_min if rango_max > rango_min else 1

                mapa_html = '<div style="position:relative; height:60px; background:#0f172a; border-radius:10px; margin:10px 0 20px 0; border:1px solid #1e293b;">'

                # Líneas de niveles
                for strike_n, tipo_n, vol_n in todos_niveles:
                    pos_pct = ((strike_n - rango_min) / rango_total) * 100
                    pos_pct = max(2, min(98, pos_pct))
                    color = "#10b981" if tipo_n == "S" else "#ef4444"
                    opacity = 0.4 + 0.6 * (vol_n / max_vol_nivel)
                    label = tipo_n
                    mapa_html += (
                        f'<div style="position:absolute; left:{pos_pct:.1f}%; top:0; bottom:0; '
                        f'width:2px; background:{color}; opacity:{opacity:.2f};"></div>'
                        f'<div style="position:absolute; left:{pos_pct:.1f}%; top:2px; transform:translateX(-50%); '
                        f'font-size:0.65rem; font-weight:700; color:{color};">{label} ${strike_n:,.0f}</div>'
                        f'<div style="position:absolute; left:{pos_pct:.1f}%; bottom:2px; transform:translateX(-50%); '
                        f'font-size:0.6rem; color:#64748b;">{vol_n:,.0f}</div>'
                    )

                # Línea del precio actual
                pos_precio = ((precio_actual - rango_min) / rango_total) * 100
                pos_precio = max(2, min(98, pos_precio))
                mapa_html += (
                    f'<div style="position:absolute; left:{pos_precio:.1f}%; top:0; bottom:0; '
                    f'width:3px; background:#f59e0b; z-index:5;"></div>'
                    f'<div style="position:absolute; left:{pos_precio:.1f}%; top:50%; transform:translate(-50%,-50%); '
                    f'background:#f59e0b; color:#000; font-size:0.7rem; font-weight:800; padding:2px 6px; '
                    f'border-radius:4px; z-index:6; white-space:nowrap;">📍 ${precio_actual:,.2f}</div>'
                )

                mapa_html += '</div>'
                st.markdown(mapa_html, unsafe_allow_html=True)

                # Resumen de niveles cercanos
                soportes_abajo = sorted([n for n in niveles_s if n[0] < precio_actual], key=lambda x: x[0], reverse=True)
                resistencias_arriba = sorted([n for n in niveles_r if n[0] > precio_actual], key=lambda x: x[0])

                col_near1, col_near2 = st.columns(2)
                with col_near1:
                    if soportes_abajo:
                        s_cercano = soportes_abajo[0]
                        dist_s = ((s_cercano[0] - precio_actual) / precio_actual) * 100
                        st.metric("🟢 Soporte mís cercano", f"${s_cercano[0]:,.1f}", 
                                 delta=f"{dist_s:.2f}% abajo", delta_color="normal")
                    else:
                        st.info("Sin soportes por debajo del precio actual")
                with col_near2:
                    if resistencias_arriba:
                        r_cercana = resistencias_arriba[0]
                        dist_r = ((r_cercana[0] - precio_actual) / precio_actual) * 100
                        st.metric("🔴 Resistencia mís cercana", f"${r_cercana[0]:,.1f}", 
                                 delta=f"+{dist_r:.2f}% arriba", delta_color="inverse")
                    else:
                        st.info("Sin resistencias por encima del precio actual")
        else:
            st.info("No hay suficientes datos de CALLs y PUTs para calcular soportes y resistencias.")

        st.markdown("---")

        col_a1, col_a2 = st.columns(2)

        with col_a1:
            st.markdown("#### 📊 Distribución CALL vs PUT")
            tipo_counts = df_analisis["Tipo"].value_counts()
            st.bar_chart(tipo_counts)

            n_calls = tipo_counts.get("CALL", 0)
            n_puts = tipo_counts.get("PUT", 0)
            ratio_pc = n_puts / n_calls if n_calls > 0 else 0
            st.metric("Put/Call Ratio", f"{ratio_pc:.3f}")
            if ratio_pc < 0.7:
                st.success("📈 Ratio < 0.7: Mayor actividad en CALLs (sentimiento alcista)")
            else:
                st.info("↔️ Ratio neutral")

        with col_a2:
            st.markdown("#### 📅 Volumen por Vencimiento")
            vol_by_date = (
                df_analisis.groupby("Vencimiento")["Volumen"]
                .sum()
                .sort_index()
            )
            st.bar_chart(vol_by_date)

        st.markdown("#### 🎯 Top 20 Strikes por Volumen")
        vol_cols = ["Vencimiento", "Tipo", "Strike", "Volumen", "IV", "Ultimo", "Prima_Vol", "Lado"]
        top_vol = (
            df_analisis.nlargest(20, "Volumen")[[c for c in vol_cols if c in df_analisis.columns]]
            .reset_index(drop=True)
        )
        top_vol_display = top_vol.copy()
        top_vol_display = top_vol_display.rename(columns={"Prima_Vol": "Prima Total"})
        if "Tipo" in top_vol_display.columns and "Lado" in top_vol_display.columns:
            top_vol_display.insert(0, "Sentimiento", top_vol_display.apply(
                lambda row: _sentiment_badge(row["Tipo"], row.get("Lado", "N/A")), axis=1
            ))
        if "Tipo" in top_vol_display.columns:
            top_vol_display["Tipo"] = top_vol_display["Tipo"].apply(_type_badge)
        top_vol_display["Prima Total"] = top_vol_display["Prima Total"].apply(_fmt_dolar)
        if "Lado" in top_vol_display.columns:
            top_vol_display["Lado"] = top_vol_display["Lado"].apply(_fmt_lado)
        st.markdown(
            render_pro_table(top_vol_display, title="🎯 Top 20 por Volumen", badge_count="20"),
            unsafe_allow_html=True,
        )

        st.markdown("#### 🏛️ Top 20 Strikes por Open Interest")
        oi_cols = ["Vencimiento", "Tipo", "Strike", "OI", "Volumen", "IV", "Ultimo", "Prima_Vol", "Lado"]
        top_oi = (
            df_analisis.nlargest(20, "OI")[[c for c in oi_cols if c in df_analisis.columns]]
            .reset_index(drop=True)
        )
        top_oi_display = top_oi.copy()
        top_oi_display = top_oi_display.rename(columns={"Prima_Vol": "Prima Total"})
        if "Tipo" in top_oi_display.columns and "Lado" in top_oi_display.columns:
            top_oi_display.insert(0, "Sentimiento", top_oi_display.apply(
                lambda row: _sentiment_badge(row["Tipo"], row.get("Lado", "N/A")), axis=1
            ))
        if "Tipo" in top_oi_display.columns:
            top_oi_display["Tipo"] = top_oi_display["Tipo"].apply(_type_badge)
        top_oi_display["Prima Total"] = top_oi_display["Prima Total"].apply(_fmt_dolar)
        if "Lado" in top_oi_display.columns:
            top_oi_display["Lado"] = top_oi_display["Lado"].apply(_fmt_lado)
        st.markdown(
            render_pro_table(top_oi_display, title="🏛️ Top 20 por Open Interest", badge_count="20"),
            unsafe_allow_html=True,
        )

        col_iv1, col_iv2 = st.columns(2)
        with col_iv1:
            st.markdown("#### 📉 Volatilidad Implícita por Strike (CALLs)")
            calls_iv = df_analisis[
                (df_analisis["Tipo"] == "CALL") & (df_analisis["IV"] > 0)
            ].sort_values("Strike")
            if not calls_iv.empty:
                chart_data_calls = calls_iv[["Strike", "IV"]].set_index("Strike")
                st.line_chart(chart_data_calls)
        with col_iv2:
            st.markdown("#### 📉 Volatilidad Implícita por Strike (PUTs)")
            puts_iv = df_analisis[
                (df_analisis["Tipo"] == "PUT") & (df_analisis["IV"] > 0)
            ].sort_values("Strike")
            if not puts_iv.empty:
                chart_data_puts = puts_iv[["Strike", "IV"]].set_index("Strike")
                st.line_chart(chart_data_puts)

        # Desglose por vencimiento
        df_calls_s = df_analisis[df_analisis["Tipo"] == "CALL"]
        df_puts_s = df_analisis[df_analisis["Tipo"] == "PUT"]
        col_pv1, col_pv2 = st.columns(2)

        with col_pv1:
            st.markdown("#### 📞 Prima Total en CALLs por Vencimiento")
            if not df_calls_s.empty:
                prima_calls_venc = df_calls_s.groupby("Vencimiento").agg(
                    Prima_Total=("Prima_Vol", "sum"),
                    Contratos=("Volumen", "count"),
                    Volumen_Total=("Volumen", "sum"),
                ).sort_values("Prima_Total", ascending=False).reset_index()

                display_pc = prima_calls_venc.copy()
                display_pc["Prima_Total"] = display_pc["Prima_Total"].apply(_fmt_dolar)
                display_pc["Volumen_Total"] = display_pc["Volumen_Total"].apply(_fmt_entero)
                st.markdown(
                    render_pro_table(display_pc, title="📞 CALLs por Vencimiento"),
                    unsafe_allow_html=True,
                )
            else:
                st.info("Sin datos de CALLs.")

        with col_pv2:
            st.markdown("#### 📋 Prima Total en PUTs por Vencimiento")
            if not df_puts_s.empty:
                prima_puts_venc = df_puts_s.groupby("Vencimiento").agg(
                    Prima_Total=("Prima_Vol", "sum"),
                    Contratos=("Volumen", "count"),
                    Volumen_Total=("Volumen", "sum"),
                ).sort_values("Prima_Total", ascending=False).reset_index()

                display_pp = prima_puts_venc.copy()
                display_pp["Prima_Total"] = display_pp["Prima_Total"].apply(_fmt_dolar)
                display_pp["Volumen_Total"] = display_pp["Volumen_Total"].apply(_fmt_entero)
                st.markdown(
                    render_pro_table(display_pp, title="📋 PUTs por Vencimiento"),
                    unsafe_allow_html=True,
                )
            else:
                st.info("Sin datos de PUTs.")

        # Top strikes donde se concentra el dinero
        st.markdown("#### 🎯 Top 15 Strikes con Mayor Prima Total Ejecutada")
        df_prima_strike = df_analisis.copy()
        prima_cols = ["Tipo", "Strike", "Vencimiento", "Volumen", "Prima_Vol", "IV", "Ultimo", "Lado"]
        top_prima = df_prima_strike.nlargest(15, "Prima_Vol")[
            [c for c in prima_cols if c in df_prima_strike.columns]
        ].reset_index(drop=True)

        top_prima_display = top_prima.copy()
        top_prima_display = top_prima_display.rename(columns={"Prima_Vol": "Prima Total"})
        if "Tipo" in top_prima_display.columns and "Lado" in top_prima_display.columns:
            top_prima_display.insert(0, "Sentimiento", top_prima_display.apply(
                lambda row: _sentiment_badge(row["Tipo"], row.get("Lado", "N/A")), axis=1
            ))
        if "Tipo" in top_prima_display.columns:
            top_prima_display["Tipo"] = top_prima_display["Tipo"].apply(_type_badge)
        top_prima_display["Prima Total"] = top_prima_display["Prima Total"].apply(_fmt_dolar)
        top_prima_display["Volumen"] = top_prima_display["Volumen"].apply(_fmt_entero)
        top_prima_display["IV"] = top_prima_display["IV"].apply(_fmt_iv)
        top_prima_display["Ultimo"] = top_prima_display["Ultimo"].apply(_fmt_precio)
        top_prima_display["Strike"] = top_prima_display["Strike"].apply(lambda x: f"${x:,.1f}")
        if "Lado" in top_prima_display.columns:
            top_prima_display["Lado"] = top_prima_display["Lado"].apply(_fmt_lado)

        st.markdown(
            render_pro_table(top_prima_display, title="🎯 Top 15 Mayor Prima Ejecutada", badge_count="15"),
            unsafe_allow_html=True,
        )

        # Grífica de prima por strike
        st.markdown("#### 📊 Flujo de Prima por Strike (CALL vs PUT)")
        pivot_prima = df_analisis.pivot_table(
            index="Strike", columns="Tipo",
            values="Prima_Vol", aggfunc="sum", fill_value=0,
        )
        pivot_prima = pivot_prima[pivot_prima.sum(axis=1) > 0]
        if not pivot_prima.empty:
            pivot_prima = pivot_prima.nlargest(30, pivot_prima.columns.tolist()[0] if len(pivot_prima.columns) > 0 else pivot_prima.index).sort_index()
            st.bar_chart(pivot_prima)
        st.caption("Prima por Volumen distribuida por strike — muestra dónde se concentran las apuestas mís grandes")


# ============================================================================
#   ⭐ FAVORITES — CONTRATOS FAVORITOS
# ============================================================================
elif pagina == "⭐ Favorites":
    st.markdown("### ⭐ Contratos Favoritos")

    favoritos = st.session_state.get("favoritos", [])

    if not favoritos:
        st.info("No hay contratos en favoritos. Ejecutí un escaneo y usí el botón ☆ **Guardar en Favoritos** en cualquier alerta.")
    else:
        # Mîtricas rípidas
        n_calls_fav = sum(1 for f in favoritos if f.get("Tipo_Opcion") == "CALL")
        n_puts_fav = sum(1 for f in favoritos if f.get("Tipo_Opcion") == "PUT")
        prima_total_fav = sum(f.get("Prima_Volumen", 0) for f in favoritos)
        st.markdown(render_metric_row([
            render_metric_card("Total Favoritos", f"{len(favoritos)}"),
            render_metric_card("Calls", f"{n_calls_fav}"),
            render_metric_card("Puts", f"{n_puts_fav}"),
            render_metric_card("Prima Total", _fmt_monto(prima_total_fav)),
        ]), unsafe_allow_html=True)

        # Tabla resumen
        fav_df = pd.DataFrame(favoritos)
        cols_tabla_fav = ["Contrato", "Ticker", "Tipo_Opcion", "Strike", "Vencimiento", 
                          "Volumen", "OI", "Ask", "Bid", "Ultimo", "Lado", "Prima_Volumen"]
        cols_disp_fav = [c for c in cols_tabla_fav if c in fav_df.columns]
        display_fav_df = fav_df[cols_disp_fav].copy()
        if "Tipo_Opcion" in display_fav_df.columns and "Lado" in display_fav_df.columns:
            display_fav_df.insert(0, "Sentimiento", display_fav_df.apply(
                lambda row: _sentiment_badge(row["Tipo_Opcion"], row.get("Lado", "N/A")), axis=1
            ))
        if "Tipo_Opcion" in display_fav_df.columns:
            display_fav_df["Tipo_Opcion"] = display_fav_df["Tipo_Opcion"].apply(_type_badge)
        if "Lado" in display_fav_df.columns:
            display_fav_df["Lado"] = display_fav_df["Lado"].apply(_fmt_lado)
        if "Prima_Volumen" in display_fav_df.columns:
            display_fav_df = display_fav_df.rename(columns={"Prima_Volumen": "Prima Total"})
            display_fav_df["Prima Total"] = display_fav_df["Prima Total"].apply(_fmt_monto)
        st.markdown(
            render_pro_table(display_fav_df, title="⭐ Favoritos", badge_count=f"{len(favoritos)}"),
            unsafe_allow_html=True,
        )

        st.markdown("---")

        # Detalle individual de cada favorito
        st.markdown("#### 🔍 Detalle de Contratos")
        for idx_fav, fav in enumerate(favoritos):
            fav_sym = fav.get("Contrato", "N/A")
            fav_tipo = fav.get("Tipo_Opcion", "N/A")
            fav_strike = fav.get("Strike", 0)
            fav_venc = fav.get("Vencimiento", "N/A")
            fav_prima = fav.get("Prima_Volumen", 0)

            # Calcular días para vencimiento
            try:
                dias_venc = (datetime.strptime(fav_venc, "%Y-%m-%d") - datetime.now()).days
                dias_str = f"{dias_venc}d" if dias_venc >= 0 else "EXPIRADO"
            except Exception:
                dias_str = "N/A"

            fav_label = (
                f"⭐ {fav_tipo} ${fav_strike} | Venc: {fav_venc} ({dias_str}) | "
                f"Prima: ${fav_prima:,.0f} | {fav_sym}"
            )

            with st.expander(fav_label, expanded=False):
                col_fav_info, col_fav_chart = st.columns([1, 2])

                with col_fav_info:
                    st.markdown("**📄 Información del Contrato**")
                    st.markdown(f"- **Símbolo:** `{fav_sym}`")
                    st.markdown(f"- **Ticker:** {fav.get('Ticker', 'N/A')}")
                    st.markdown(f"- **Tipo:** {fav_tipo}")
                    st.markdown(f"- **Strike:** ${fav_strike}")
                    st.markdown(f"- **Vencimiento:** {fav_venc} ({dias_str})")
                    st.markdown(f"- **Volumen:** {fav.get('Volumen', 0):,}")
                    st.markdown(f"- **OI:** {fav.get('OI', 0):,}")
                    oi_chg_val = fav.get('OI_Chg', 0)
                    st.markdown(f"- **OI Chg:** {_fmt_oi_chg(oi_chg_val)}")
                    st.markdown(f"- **Ask:** ${fav.get('Ask', 0)}")
                    st.markdown(f"- **Bid:** ${fav.get('Bid', 0)}")
                    st.markdown(f"- **Último:** ${fav.get('Ultimo', 0)}")
                    st.markdown(f"- **Lado:** {_fmt_lado(fav.get('Lado', 'N/A'))}")
                    iv_fav = fav.get('IV', 0)
                    st.markdown(f"- **IV:** {iv_fav:.1f}%" if iv_fav > 0 else "- **IV:** N/A")
                    st.markdown(f"- **Prima Total:** {_fmt_monto(fav.get('Prima_Volumen', 0))}")
                    st.markdown(f"- **Tipo Alerta:** {fav.get('Tipo_Alerta', 'N/A')}")
                    st.markdown(f"- **Guardado:** {fav.get('Guardado_En', 'N/A')}")

                    # Botón eliminar
                    if st.button(f"🗑️ Eliminar de Favoritos", key=f"del_fav_{idx_fav}_{fav_sym}", use_container_width=True):
                        _eliminar_favorito(fav_sym)
                        st.success(f"🗑️ {fav_sym} eliminado de Favoritos")
                        st.rerun()

                with col_fav_chart:
                    if fav_sym and fav_sym != "N/A":
                        with st.spinner("Cargando grífica del contrato..."):
                            hist_fav, err_fav = obtener_historial_contrato(fav_sym)

                        if err_fav:
                            st.warning(f"⚠️ Error al cargar historial: {err_fav}")
                        elif hist_fav.empty:
                            st.info("ℹ️ No hay datos históricos disponibles.")
                        else:
                            st.markdown(f"**Precio del contrato** — `{fav_sym}`")
                            chart_fav_price = hist_fav[["Close"]].copy()
                            chart_fav_price.columns = ["Precio"]
                            st.line_chart(chart_fav_price, height=280)

                            if "Volume" in hist_fav.columns:
                                chart_fav_vol = hist_fav[["Volume"]].copy()
                                chart_fav_vol.columns = ["Volumen"]
                                st.bar_chart(chart_fav_vol, height=160)

        # Botón limpiar todos
        st.markdown("---")
        col_limpiar, _ = st.columns([1, 3])
        with col_limpiar:
            if st.button("🗑️ Limpiar todos los favoritos", use_container_width=True, type="secondary"):
                st.session_state.favoritos = []
                _guardar_favoritos([])
                st.success("Se eliminaron todos los favoritos")
                st.rerun()


# ============================================================================
#   📐 RANGE — RANGO ESPERADO
# ============================================================================
elif pagina == "📐 Range":
    st.markdown("### 📐 Rango Esperado de Movimiento (1σ)")

    rango_delta = st.slider(
        "Delta objetivo (σ)", min_value=0.01, max_value=1.00, value=st.session_state.rango_delta, step=0.01,
        help="0.16 ≈ 1σ (68%). 0.05 ≈ 2σ (95%). Menor delta = rango mís amplio.", key="rango_delta_slider"
    )
    st.session_state.rango_delta = rango_delta

    st.markdown("")

    # Reusar fechas ya cargadas del último escaneo para evitar llamada extra
    fechas_exp_disponibles = list(st.session_state.get("fechas_escaneadas", []))
    if not fechas_exp_disponibles:
        try:
            session_rango, _ = crear_sesion_nueva()
            ticker_rango = yf.Ticker(ticker_symbol, session=session_rango)
            fechas_exp_disponibles = list(ticker_rango.options)
        except Exception as e:
            logger.warning("Error obteniendo fechas de expiración para rango: %s", e)

    col_r1, col_r2, col_r3 = st.columns([2, 2, 1])
    with col_r1:
        rango_symbol = st.text_input(
            "Símbolo", value=ticker_symbol, max_chars=10,
            key="rango_symbol", help="Ticker de la acción (ej: SPY, META, AAPL)"
        ).upper()
    with col_r2:
        if fechas_exp_disponibles:
            rango_exp_date = st.selectbox(
                "Fecha de Expiración",
                fechas_exp_disponibles,
                key="rango_exp",
                help="Fechas de expiración disponibles para este ticker",
            )
        else:
            rango_exp_date = st.text_input(
                "Fecha de Expiración (YYYY-MM-DD)",
                value=(datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d"),
                key="rango_exp",
                help="Fecha de expiración de las opciones a analizar",
            )
    with col_r3:
        st.markdown("<br>", unsafe_allow_html=True)
        calc_btn = st.button("📐 Calcular Rango", type="primary", use_container_width=True)

    if calc_btn:
        st.session_state.scanning_active = True
        with st.spinner(f"Calculando rango esperado para {rango_symbol} al {rango_exp_date}..."):
            resultado, error = calcular_rango_esperado(
                rango_symbol, rango_exp_date,
                target_delta=rango_delta,
            )
        st.session_state.rango_resultado = resultado
        st.session_state.rango_error = error
        st.session_state.scanning_active = False

    if st.session_state.rango_error:
        st.error(f"❌ {st.session_state.rango_error}")

    if st.session_state.rango_resultado:
        r = st.session_state.rango_resultado

        total_range = r["downside_points"] + r["upside_points"]
        if total_range > 0:
            down_pct_bar = (r["downside_points"] / total_range) * 100
            up_pct_bar = (r["upside_points"] / total_range) * 100
        else:
            down_pct_bar = 50
            up_pct_bar = 50

        full_range = r["expected_range_high"] - r["expected_range_low"]
        if full_range > 0:
            precio_pos = ((r["underlying_price"] - r["expected_range_low"]) / full_range) * 100
        else:
            precio_pos = 50

        dias_str = f" ({r['dias_restantes']} días)" if r['dias_restantes'] is not None else ""

        st.markdown(f"#### 📐 {r['symbol']} — Rango Esperado 1σ")
        st.caption(f"Expiración: {r['expiration']}{dias_str}")

        col_r1, col_r2, col_r3, col_r4 = st.columns(4)
        with col_r1:
            st.metric("💵 Precio Actual", f"${r['underlying_price']:,.2f}")
        with col_r2:
            st.metric("📈 Subida Esperada", f"+${r['upside_points']:,.2f}", f"+{r['upside_percent']:.2f}%")
        with col_r3:
            st.metric("📉 Bajada Esperada", f"-${r['downside_points']:,.2f}", f"-{r['downside_percent']:.2f}%", delta_color="inverse")
        with col_r4:
            st.metric("↔️ Rango Total", f"${r['total_range_points']:,.2f}", f"{r['total_range_pct']:.2f}%", delta_color="off")

        st.markdown("")
        bar_col1, bar_col2, bar_col3 = st.columns([1, 6, 1])
        with bar_col1:
            st.markdown(f"**▼ ${r['expected_range_low']:,.2f}**")
        with bar_col2:
            progress_val = max(0.0, min(1.0, precio_pos / 100.0))
            st.progress(progress_val, text=f"● Precio actual: ${r['underlying_price']:,.2f}  —  Rango: ${r['expected_range_low']:,.2f} a ${r['expected_range_high']:,.2f}")
        with bar_col3:
            st.markdown(f"**▲ ${r['expected_range_high']:,.2f}**")

        st.divider()

        st.markdown("#### 🎯 Contratos Usados para el Cílculo")
        col_d1, col_d2 = st.columns(2)

        with col_d1:
            st.success(f"""
**📈 CALL (límite superior)**
- Strike: **${r['call_strike']}**
- Delta: **{r['call_delta']}**
- IV: **{r['call_iv']:.1f}%**
- _Precio debe superar ${r['call_strike']} para salir del rango_
""")

        with col_d2:
            st.error(f"""
**📉 PUT (límite inferior)**
- Strike: **${r['put_strike']}**
- Delta: **{r['put_delta']}**
- IV: **{r['put_iv']:.1f}%**
- _Precio debe caer bajo ${r['put_strike']} para salir del rango_
""")

        with st.expander("📋 Ver datos completos del cílculo"):
            resumen_data = {
                "Campo": [
                    "Símbolo", "Precio Actual", "Expiración", "Días Restantes",
                    "Delta Objetivo", "Subida Esperada (pts)", "Subida Esperada (%)",
                    "Bajada Esperada (pts)", "Bajada Esperada (%)", "Rango Inferior",
                    "Rango Superior", "Rango Total (pts)", "Rango Total (%)",
                    "Call Strike", "Call Delta", "Call IV",
                    "Put Strike", "Put Delta", "Put IV",
                    "Calls Analizadas", "Puts Analizadas",
                ],
                "Valor": [
                    r["symbol"], f"${r['underlying_price']:,.2f}", r["expiration"],
                    r["dias_restantes"] if r["dias_restantes"] else "N/A",
                    f"±{r['target_delta']}",
                    f"+${r['upside_points']:,.2f}", f"+{r['upside_percent']:.2f}%",
                    f"-${r['downside_points']:,.2f}", f"-{r['downside_percent']:.2f}%",
                    f"${r['expected_range_low']:,.2f}", f"${r['expected_range_high']:,.2f}",
                    f"${r['total_range_points']:,.2f}", f"{r['total_range_pct']:.2f}%",
                    f"${r['call_strike']}", f"{r['call_delta']}", f"{r['call_iv']:.1f}%",
                    f"${r['put_strike']}", f"{r['put_delta']}", f"{r['put_iv']:.1f}%",
                    r["n_calls"], r["n_puts"],
                ]
            }
            st.markdown(
                render_pro_table(pd.DataFrame(resumen_data), title="📋 Datos del Cílculo"),
                unsafe_allow_html=True,
            )

        st.markdown(
            f"""
            <div class="rango-info">
                🧠 <b>Interpretación:</b> El mercado de opciones estima que <b>{r['symbol']}</b>
                se moverí entre <b>${r['expected_range_low']:,.2f}</b> y <b>${r['expected_range_high']:,.2f}</b>
                (un rango de <b>${r['total_range_points']:,.2f}</b> / <b>{r['total_range_pct']:.2f}%</b>)
                hasta el <b>{r['expiration']}</b> con ~68% de probabilidad.
                Esto equivale a ±1 desviación estíndar implícita del mercado.
            </div>
            """,
            unsafe_allow_html=True,
        )


# ============================================================================
#   🏢 IMPORTANT COMPANIES
# ============================================================================
elif pagina == "🏢 Important Companies":
    st.markdown("### 🏢 Proyecciones de Crecimiento a 10 Años")

    # ==============================================================
    #  SECCIÓN 1: EMPRESAS CONSOLIDADAS
    # ==============================================================
    st.markdown("---")
    st.markdown("## 🏢 Empresas Consolidadas — Top Corporations")
    st.caption("Grandes corporaciones con historial probado y proyección de crecimiento sostenido a 10 años.")

    col_btn_c, col_info_c = st.columns([1, 3])
    with col_btn_c:
        analizar_consol_btn = st.button(
            "📊 Analizar Consolidadas en Vivo",
            type="primary",
            use_container_width=True,
            key="btn_analizar_consolidadas",
        )
    with col_info_c:
        if "proyecciones_resultados" in st.session_state and st.session_state.proyecciones_resultados:
            st.success(f"✅ Datos en vivo cargados — {len(st.session_state.proyecciones_resultados)} empresas analizadas")

    if analizar_consol_btn:
        st.session_state.scanning_active = True
        analizar_watchlist(WATCHLIST_EMPRESAS, "proyecciones_resultados", "consolidadas")
        st.session_state.scanning_active = False

    if "proyecciones_resultados" in st.session_state and st.session_state.proyecciones_resultados:
        resultados = st.session_state.proyecciones_resultados

        col_s1, col_s2, col_s3 = st.columns(3)
        alta_count = sum(1 for r in resultados if r["clasificacion"] == "ALTA")
        media_count = sum(1 for r in resultados if r["clasificacion"] == "MEDIA")
        baja_count = sum(1 for r in resultados if r["clasificacion"] == "BAJA")
        st.markdown(render_metric_row([
            render_metric_card("Proyección Alta", f"{alta_count}"),
            render_metric_card("Proyección Media", f"{media_count}", color_override="#f59e0b"),
            render_metric_card("Proyección Baja", f"{baja_count}", color_override="#ef4444"),
        ]), unsafe_allow_html=True)

        for r in resultados:
            info_emp = WATCHLIST_EMPRESAS.get(r["symbol"])
            st.html(render_empresa_card(r, info_emp, WATCHLIST_EMPRESAS))

        st.markdown("#### 📋 Tabla Comparativa")
        df_tabla = render_tabla_comparativa(resultados)
        st.markdown(
            render_pro_table(df_tabla, title="📋 Tabla Comparativa Consolidadas", badge_count=f"{len(df_tabla)}"),
            unsafe_allow_html=True,
        )
        # Botón CSV eliminado según solicitud

    else:
        st.markdown("#### 🏛️ Top Empresas Consolidadas")
        render_watchlist_preview(WATCHLIST_EMPRESAS)

    if "proyecciones_resultados" in st.session_state and st.session_state.proyecciones_resultados:
        with st.expander("📊 Anílisis de las Empresas Consolidadas", expanded=False):
            render_analisis_completo(st.session_state.proyecciones_resultados, WATCHLIST_EMPRESAS)

    # ==============================================================
    #  SECCIÓN 2: EMPRESAS EMERGENTES
    # ==============================================================
    st.markdown("---")
    st.markdown("## 🚀 Empresas Emergentes — Futuras Transnacionales")
    st.caption("Empresas de menor capitalización con tecnologías disruptivas y potencial de convertirse en gigantes. Mayor riesgo, mayor recompensa.")

    col_btn_e, col_info_e = st.columns([1, 3])
    with col_btn_e:
        analizar_emerg_btn = st.button(
            "🚀 Analizar Emergentes en Vivo",
            type="primary",
            use_container_width=True,
            key="btn_analizar_emergentes",
        )
    with col_info_e:
        if "emergentes_resultados" in st.session_state and st.session_state.emergentes_resultados:
            st.success(f"✅ Datos en vivo cargados — {len(st.session_state.emergentes_resultados)} empresas analizadas")

    if analizar_emerg_btn:
        st.session_state.scanning_active = True
        analizar_watchlist(WATCHLIST_EMERGENTES, "emergentes_resultados", "emergentes")
        st.session_state.scanning_active = False

    if "emergentes_resultados" in st.session_state and st.session_state.emergentes_resultados:
        resultados_em = st.session_state.emergentes_resultados

        col_e1, col_e2, col_e3 = st.columns(3)
        alta_em = sum(1 for r in resultados_em if r["clasificacion"] == "ALTA")
        media_em = sum(1 for r in resultados_em if r["clasificacion"] == "MEDIA")
        baja_em = sum(1 for r in resultados_em if r["clasificacion"] == "BAJA")
        st.markdown(render_metric_row([
            render_metric_card("Proyección Alta", f"{alta_em}"),
            render_metric_card("Proyección Media", f"{media_em}", color_override="#f59e0b"),
            render_metric_card("Proyección Baja", f"{baja_em}", color_override="#ef4444"),
        ]), unsafe_allow_html=True)

        for r in resultados_em:
            info_emp = WATCHLIST_EMERGENTES.get(r["symbol"])
            st.html(render_empresa_card(r, info_emp, WATCHLIST_EMERGENTES, es_emergente=True))

        st.markdown("#### 📋 Tabla Comparativa Emergentes")
        df_emerg = render_tabla_comparativa(resultados_em, es_emergente=True)
        st.markdown(
            render_pro_table(df_emerg, title="📋 Tabla Comparativa Emergentes", badge_count=f"{len(df_emerg)}"),
            unsafe_allow_html=True,
        )
        # Botón CSV eliminado según solicitud

    else:
        st.markdown("#### 🚀 Top Empresas Emergentes")
        render_watchlist_preview(WATCHLIST_EMERGENTES)

    if "emergentes_resultados" in st.session_state and st.session_state.emergentes_resultados:
        with st.expander("📊 Anílisis de las Empresas Emergentes", expanded=False):
            render_analisis_completo(st.session_state.emergentes_resultados, WATCHLIST_EMERGENTES, es_emergente=True)


# ============================================================================
#   📰 NEWS & CALENDAR — NOTICIAS
# ============================================================================
elif pagina == "📰 News & Calendar":
    st.markdown("### 📰 Noticias Financieras en Tiempo Real")
    st.markdown(
        """
        <div class="watchlist-info">
            📡 <b>Centro de Noticias</b> — Noticias financieras de
            Yahoo Finance, MarketWatch, CNBC, Reuters e Investing.com.
            Filtra por relevancia, tendencia mundial o categoría. 🆓 100% gratuito vía RSS.
        </div>
        """,
        unsafe_allow_html=True,
    )

    # --- CONTROLES SIEMPRE VISIBLES ---
    col_load, col_refresh, col_auto = st.columns([1.5, 1.5, 2])

    with col_load:
        cargar_noticias_btn = st.button(
            "📡 Cargar Noticias" if not st.session_state.noticias_data else "📡 Recargar Todo",
            type="primary",
            use_container_width=True,
            key="btn_cargar_noticias_main",
        )
    with col_refresh:
        refresh_noticias_btn = st.button(
            "🔄 Refrescar",
            use_container_width=True,
            key="btn_refresh_noticias",
            disabled=not st.session_state.noticias_data,
        )
    with col_auto:
        auto_refresh_noticias = st.checkbox(
            "⏱️ Auto-refresco cada 5 min",
            value=st.session_state.noticias_auto_refresh,
            key="chk_auto_refresh_noticias",
            help="Actualiza las noticias automíticamente cada 5 minutos",
        )
        st.session_state.noticias_auto_refresh = auto_refresh_noticias

    # --- FILTROS ---
    col_filtro1, col_filtro2 = st.columns([3, 2])
    with col_filtro1:
        filtro_noticias = st.selectbox(
            "🏷️ Filtrar por:",
            [
                "Todas",
                "🔥 Mís relevantes",
                "🌍 Mís vistas a nivel mundial",
                "Mís relevantes para trading",
                "Top Stories",
                "Earnings",
                "Fed / Tasas",
                "Economía",
                "Trading",
                "Crypto",
                "Commodities",
                "Geopolítica",
            ],
            index=0,
            key="sel_filtro_noticias",
        )
    with col_filtro2:
        ordenar_por = st.selectbox(
            "📊 Ordenar por:",
            ["Mís recientes", "Mís relevantes primero"],
            index=0,
            key="sel_orden_noticias",
        )

    # --- CARGAR / REFRESCAR ---
    necesita_refresh = False
    if auto_refresh_noticias:
        last = st.session_state.noticias_last_refresh
        if last is None:
            necesita_refresh = True
        else:
            elapsed = (datetime.now() - last).total_seconds()
            if elapsed >= AUTO_REFRESH_INTERVAL:
                necesita_refresh = True

    if cargar_noticias_btn or refresh_noticias_btn or necesita_refresh:
        with st.spinner("📡 Obteniendo noticias de múltiples fuentes..."):
            noticias = obtener_noticias_financieras()
            if noticias:
                st.session_state.noticias_data = noticias
                st.session_state.noticias_last_refresh = datetime.now()
                if cargar_noticias_btn or refresh_noticias_btn:
                    st.rerun()

    # --- AUTO-REFRESH COUNTDOWN ---
    if auto_refresh_noticias and st.session_state.noticias_last_refresh:
        elapsed = (datetime.now() - st.session_state.noticias_last_refresh).total_seconds()
        remaining = max(0, AUTO_REFRESH_INTERVAL - elapsed)
        mins_left = int(remaining // 60)
        secs_left = int(remaining % 60)
        st.caption(
            f"🔄 Auto-refresco activo — Próxima actualización en **{mins_left}:{secs_left:02d}** · "
            f"Último: **{st.session_state.noticias_last_refresh.strftime('%H:%M:%S')}**"
        )

    # --- CONTENIDO ---
    if not st.session_state.noticias_data:
        st.info(
            "👆 Presiona **Cargar Noticias** para obtener las últimas noticias financieras "
            "de Yahoo Finance, MarketWatch, CNBC, Reuters e Investing.com."
        )
    else:
        # Mîtricas
        col_status1, col_status2, col_status3 = st.columns(3)
        with col_status1:
            st.metric("🕐 Última actualización", st.session_state.noticias_last_refresh.strftime('%H:%M:%S'))
        with col_status2:
            st.metric("📰 Total noticias", len(st.session_state.noticias_data))
        with col_status3:
            st.metric("🏷️ Filtro activo", filtro_noticias)

        # Distribución por categoría
        cat_counts = {}
        for n in st.session_state.noticias_data:
            cat = n["categoria"]
            cat_counts[cat] = cat_counts.get(cat, 0) + 1

        top_cats = sorted(cat_counts.items(), key=lambda x: x[1], reverse=True)[:6]
        if top_cats:
            stat_cols = st.columns(len(top_cats))
            for i, (cat_name, cat_count) in enumerate(top_cats):
                with stat_cols[i]:
                    st.metric(cat_name, cat_count)

        st.divider()

        # Filtrar y ordenar
        noticias_filtradas = filtrar_noticias(st.session_state.noticias_data, filtro_noticias)

        if ordenar_por == "Mís relevantes primero" and filtro_noticias not in ("🔥 Mís relevantes", "🌍 Mís vistas a nivel mundial"):
            from core.news import calcular_relevancia
            noticias_filtradas = sorted(noticias_filtradas, key=calcular_relevancia, reverse=True)

        if not noticias_filtradas:
            st.info(f"No hay noticias para el filtro '{filtro_noticias}'. Prueba con 'Todas'.")
        else:
            st.markdown(f"#### 📋 {len(noticias_filtradas)} noticias — {filtro_noticias}")

            cat_emoji_map = {
                "Earnings": "💰",
                "Fed / Tasas": "🏛️",
                "Economía": "📊",
                "Trading": "📈",
                "Crypto": "₿",
                "Commodities": "🛢️",
                "Geopolítica": "🌍",
                "Top Stories": "⭐",
                "Mercados": "📈",
            }

            for n in noticias_filtradas:
                cat = n["categoria"]
                emoji = cat_emoji_map.get(cat, "📰")

                with st.container():
                    col_noticia, col_cat = st.columns([5, 1])
                    with col_noticia:
                        if n["url"]:
                            st.markdown(f"**[{n['titulo']}]({n['url']})**")
                        else:
                            st.markdown(f"**{n['titulo']}**")

                        if n["descripcion"]:
                            st.caption(n["descripcion"])

                        meta_parts = []
                        if n["fuente"]:
                            meta_parts.append(f"📰 {n['fuente']}")
                        if n["tiempo"]:
                            meta_parts.append(f"🕐 {n['tiempo']}")
                        if meta_parts:
                            st.caption(" · ".join(meta_parts))

                    with col_cat:
                        st.markdown(f"**{emoji} {cat}**")

                    st.divider()

        # Auto-rerun si toca
        if auto_refresh_noticias and st.session_state.noticias_last_refresh:
            elapsed = (datetime.now() - st.session_state.noticias_last_refresh).total_seconds()
            if elapsed >= AUTO_REFRESH_INTERVAL:
                st.rerun()


    # ============================================================================
    #   CALENDARIO FINANCIERO — Sub-sección de News & Calendar
    # ============================================================================
    st.markdown("---")
    from ui.tabs.calendar_tab import render_calendar_tab
    render_calendar_tab()

# ============================================================================
#                    FOOTER
# ============================================================================
st.markdown(
    """
    <div class="footer-pro">
        <div>👑 OPTIONS<span style="color: #00ff88;">KING</span> Analytics v5.0 — Datos de Yahoo Finance</div>
        <div class="footer-badges">
            <span class="footer-badge">🔒 curl_cffi TLS</span>
            <span class="footer-badge">📊 Yahoo Finance</span>
            <span class="footer-badge">📐 Black-Scholes</span>
            <span class="footer-badge">📰 RSS Feeds</span>
            <span class="footer-badge">🎨 Streamlit</span>
            <span class="footer-badge">🐍 Python</span>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)
